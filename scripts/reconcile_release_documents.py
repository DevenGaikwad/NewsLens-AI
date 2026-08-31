"""Reconcile the four public Word documents with the hardened release evidence.

The original document generators remain useful for rebuilding the academic
layout. This script is the final release reconciliation layer: it replaces
superseded figures with current artifacts, updates measured values, documents
calibration/editorial review/privacy boundaries, and keeps the operation
idempotent so the final audit result can be recorded after execution.
"""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DIAGRAMS = ROOT / "reports" / "diagrams"
FIGURES = ROOT / "reports" / "figures"
SCREENSHOTS = ROOT / "reports" / "screenshots"
ASSETS = ROOT / "assets" / "github"

PUBLIC_AUTHOR = "Deven Gaikwad"
PUBLIC_COPYRIGHT = "© 2026 Deven Gaikwad. All rights reserved."
INTENDED_REPOSITORY = "https://github.com/DevenGaikwad/NewsLens-AI"
UPDATED_DATE = "25 August 2026"


def body_paragraphs(doc: Document) -> list:
    return list(doc.paragraphs)


def iter_all_paragraphs(doc: Document) -> Iterable:
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs
        yield from section.first_page_header.paragraphs
        yield from section.first_page_footer.paragraphs
        yield from section.even_page_header.paragraphs
        yield from section.even_page_footer.paragraphs


def replace_runs(doc: Document, old: str, new: str) -> None:
    for paragraph in iter_all_paragraphs(doc):
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)


def normalised(value: str) -> str:
    return " ".join(value.split())


def find_paragraph(
    doc: Document,
    *,
    exact: str | None = None,
    startswith: str | None = None,
    required: bool = True,
):
    for paragraph in body_paragraphs(doc):
        text = normalised(paragraph.text)
        if exact is not None and text == exact:
            return paragraph
        if startswith is not None and text.startswith(startswith):
            return paragraph
    if required:
        query = exact if exact is not None else startswith
        raise ValueError(f"Required paragraph not found: {query!r}")
    return None


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_paragraph(paragraph, text: str, *, style: str | None = None) -> None:
    clear_paragraph(paragraph)
    if style:
        paragraph.style = style
    paragraph.add_run(text)


def set_by_prefix(doc: Document, prefix: str, text: str, *, style: str | None = None) -> None:
    paragraph = find_paragraph(doc, startswith=prefix, required=False)
    if paragraph is not None:
        set_paragraph(paragraph, text, style=style)
        return
    # Idempotence matters because the same documents are reconciled once before
    # the audit and again only if the comprehensive audit actually passes.
    existing = find_paragraph(doc, exact=normalised(text), required=False)
    if existing is None:
        raise ValueError(f"Required paragraph not found and replacement absent: {prefix!r}")
    if style:
        existing.style = style


def insert_before(anchor, text: str = "", style: str | None = None):
    paragraph = anchor.insert_paragraph_before(text)
    if style:
        paragraph.style = style
    return paragraph


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:tblHeader"))
    if existing is None:
        existing = OxmlElement("w:tblHeader")
        tr_pr.append(existing)
    existing.set(qn("w:val"), "true")


def format_table(table) -> None:
    if not table.rows:
        return
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8.2)
                    if row_index == 0:
                        run.bold = True


def set_table(table, rows: Sequence[Sequence[str]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for row, values in zip(table.rows, rows, strict=True):
        if len(row.cells) != len(values):
            raise ValueError("Table column count does not match replacement data")
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = str(value)
    format_table(table)


def append_table_row(table, values: Sequence[str]) -> None:
    existing = {tuple(normalised(cell.text) for cell in row.cells) for row in table.rows}
    normalised_values = tuple(normalised(value) for value in values)
    if normalised_values in existing:
        return
    cells = table.add_row().cells
    if len(cells) != len(values):
        raise ValueError("Table column count does not match appended data")
    for cell, value in zip(cells, values, strict=True):
        cell.text = value
    format_table(table)


def upsert_table_row_by_key(table, values: Sequence[str]) -> None:
    """Update a glossary/reference row by first-column key and remove duplicates."""

    if not values:
        return
    matches = [
        row for row in table.rows[1:] if normalised(row.cells[0].text) == normalised(values[0])
    ]
    target = matches[0] if matches else table.add_row()
    if len(target.cells) != len(values):
        raise ValueError("Table column count does not match upsert data")
    for cell, value in zip(target.cells, values, strict=True):
        cell.text = value
    for duplicate in matches[1:]:
        table._tbl.remove(duplicate._tr)
    format_table(table)


def update_common_metadata(doc: Document) -> None:
    replace_runs(doc, "Deven Sachin Gaikwad", PUBLIC_AUTHOR)
    replace_runs(doc, "DEVEN SACHIN GAIKWAD", PUBLIC_AUTHOR.upper())
    replace_runs(doc, "© 2026 Deven Gaikwad. All Rights Reserved.", PUBLIC_COPYRIGHT)
    replace_runs(doc, "47 packaged checks", "56 packaged checks")
    replace_runs(doc, "all 47 checks", "all 56 checks")
    replace_runs(doc, "All 47", "All 56")
    replace_runs(doc, "Passing 47", "Passing 56")
    replace_runs(
        doc,
        "https://huggingface.co/datasets/EdinburghNLP/xsum",
        "https://github.com/EdinburghNLP/XSum",
    )
    doc.core_properties.author = PUBLIC_AUTHOR
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = (
        "Sanitized public-release documentation; publication and functional hosting remain "
        "blocked until the documented owner and model-redistribution gates clear."
    )
    doc.core_properties.keywords = (
        "NewsLens AI, NLP, linguistic credibility risk, summarization, calibration, "
        "editorial review, Streamlit"
    )
    for table in doc.tables:
        format_table(table)


def update_cover_metadata(table) -> None:
    values = {
        "Document status": "GitHub-ready public documentation; publication pending owner action",
        "Author and developer": PUBLIC_AUTHOR,
        "Copyright": PUBLIC_COPYRIGHT,
        "Repository": f"Intended canonical repository: {INTENDED_REPOSITORY} (not yet created)",
        "Document updated": UPDATED_DATE,
    }
    for row in table.rows[1:]:
        key = normalised(row.cells[0].text)
        if key in values:
            row.cells[1].text = values[key]
    format_table(table)


def update_page_reference_table(table, page_map: dict[str, str]) -> None:
    for row in table.rows[1:]:
        key = normalised(row.cells[0].text)
        if key in page_map:
            row.cells[1].text = page_map[key]
    format_table(table)


def picture_size(path: Path, *, width: float, max_height: float) -> tuple[Inches | None, Inches | None]:
    with Image.open(path) as source:
        pixel_width, pixel_height = source.size
    rendered_height = width * pixel_height / pixel_width
    if rendered_height > max_height:
        return None, Inches(max_height)
    return Inches(width), None


def add_picture_to_paragraph(
    paragraph,
    path: Path,
    *,
    alt_text: str,
    width: float = 6.2,
    max_height: float = 7.0,
) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    picture_width, picture_height = picture_size(path, width=width, max_height=max_height)
    run = paragraph.add_run()
    inline = (
        run.add_picture(str(path), width=picture_width)
        if picture_width is not None
        else run.add_picture(str(path), height=picture_height)
    )
    inline._inline.docPr.set("descr", alt_text)
    inline._inline.docPr.set("title", alt_text)


def image_paragraph_before(caption_paragraph):
    element = caption_paragraph._p.getprevious()
    while element is not None:
        if element.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            candidate = Paragraph(element, caption_paragraph._parent)
            if candidate._p.xpath(".//a:blip"):
                return candidate
            if normalised(candidate.text):
                break
        element = element.getprevious()
    raise ValueError(f"No image paragraph precedes caption: {caption_paragraph.text}")


def replace_picture_by_caption(
    doc: Document,
    caption_prefix: str,
    path: Path,
    *,
    alt_text: str,
    width: float = 6.2,
    max_height: float = 7.0,
) -> None:
    caption = find_paragraph(doc, startswith=caption_prefix)
    paragraph = image_paragraph_before(caption)
    old_relationships = [
        blip.get(qn("r:embed")) for blip in paragraph._p.xpath(".//a:blip")
    ]
    clear_paragraph(paragraph)
    add_picture_to_paragraph(
        paragraph,
        path,
        alt_text=alt_text,
        width=width,
        max_height=max_height,
    )
    live_relationships = {
        blip.get(qn("r:embed")) for blip in doc.element.body.xpath(".//a:blip")
    }
    for relationship in old_relationships:
        if relationship and relationship not in live_relationships:
            doc.part.drop_rel(relationship)


def add_figure_before(
    anchor,
    path: Path,
    caption: str,
    explanation: str,
    *,
    width: float = 6.2,
    max_height: float = 7.0,
    page_break: bool = True,
) -> None:
    picture = insert_before(anchor)
    picture.paragraph_format.page_break_before = page_break
    add_picture_to_paragraph(
        picture,
        path,
        alt_text=caption,
        width=width,
        max_height=max_height,
    )
    caption_paragraph = insert_before(anchor, caption, "Caption")
    caption_paragraph.paragraph_format.keep_with_next = True
    insert_before(anchor, explanation, "Normal")


def add_alt_text_from_captions(doc: Document) -> None:
    paragraphs = body_paragraphs(doc)
    for index, paragraph in enumerate(paragraphs):
        drawings = paragraph._p.xpath(".//wp:inline")
        if not drawings:
            continue
        caption = "NewsLens AI figure"
        for next_paragraph in paragraphs[index + 1 : index + 4]:
            if next_paragraph.style.name == "Caption" and normalised(next_paragraph.text):
                caption = normalised(next_paragraph.text)
                break
        for drawing in drawings:
            doc_properties = drawing.find(qn("wp:docPr"))
            if doc_properties is not None:
                doc_properties.set("descr", caption)
                doc_properties.set("title", caption)


def common_document_finish(doc: Document, output: Path) -> None:
    add_alt_text_from_captions(doc)
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    temporary = output.with_suffix(".reconciled.docx")
    doc.save(temporary)
    temporary.replace(output)


def reconcile_report(*, audit_passed: bool) -> Path:
    output = DOCS / "NewsLens_AI_Project_Report.docx"
    doc = Document(output)
    update_common_metadata(doc)
    update_cover_metadata(doc.tables[0])
    update_page_reference_table(
        doc.tables[1],
        {
            "Chapter 1: Introduction": "10",
            "Chapter 2: Literature Survey": "12",
            "Chapter 3: Requirements and Feasibility": "16",
            "Chapter 4: System Analysis and Design": "18",
            "Chapter 5: Dataset and Data Preprocessing": "29",
            "Chapter 6: Methodology and Algorithms": "40",
            "Chapter 7: Implementation": "43",
            "Chapter 8: Testing and Results": "45",
            "Chapter 9: Graphical User Interface": "54",
            "Chapter 10: Project Management": "71",
            "Chapter 11: Limitations, Ethics and Future Scope": "72",
            "Chapter 12: Conclusion": "73",
            "References": "74",
            "Appendices A-E": "76",
        },
    )
    set_by_prefix(doc, "Document updated: 16 August 2026", f"Document updated: {UPDATED_DATE}")

    set_by_prefix(
        doc,
        "NewsLens AI was designed and developed by",
        f"NewsLens AI was designed and developed by {PUBLIC_AUTHOR}. {PUBLIC_COPYRIGHT} "
        f"The intended canonical repository is {INTENDED_REPOSITORY}; it has not yet been "
        "created. Functional public hosting remains blocked until documentary permission or "
        "an explicit applicable licence confirms that the ISOT-derived model and calibration "
        "artefacts may be redistributed and hosted.",
    )
    set_by_prefix(
        doc,
        "The lightweight summarizer uses TF-IDF centroid",
        "The lightweight summarizer uses TF-IDF centroid sentence ranking with lead-position "
        "and information-density bonuses; CPU-compatible DistilBART remains optional. The "
        "packaged Logistic Regression artifact isot-tfidf-lr-v1.0.0 was reproduced against its "
        "established 4,800-row holdout and then kept byte-for-byte unchanged. A controlled "
        "private benchmark used 19,200 training rows, 1,199 calibration rows, 1,200 validation-"
        "policy rows and an untouched 2,399-row final test after quarantining two near-duplicate "
        "holdout rows. On that final test, Logistic Regression achieved 0.992080 accuracy, "
        "0.992080 macro-F1, 0.999481 ROC-AUC and 0.999423 PR-AUC.",
    )
    set_by_prefix(
        doc,
        "The extractive summarizer was evaluated",
        "The extractive summarizer was evaluated on 150 fixed-seed XSum test articles, producing "
        "ROUGE-1/2/L F1 of 0.153559/0.027824/0.102530, 70.6572% mean compression "
        "and 3.368 ms mean latency. The application adds Platt-calibrated confidence, a "
        "validation-policy-derived editorial-review threshold, input diagnostics, local feature "
        "contributions, a session-isolated review/archive workflow, privacy-safe newsroom "
        "analytics, drift-readiness indicators, JSON/PDF/CSV exports, 15 current responsive "
        "interface screenshots and 56 packaged checks, including the 29 established checks.",
    )
    set_by_prefix(
        doc,
        "Technical feasibility is demonstrated by",
        "Technical feasibility is demonstrated by a 0.8 MB packaged classical model, measured "
        "sub-millisecond inference, 56 packaged checks (the 29 established checks plus 27 "
        "hardening checks), and browser/a11y audit workflows. Operational feasibility comes "
        "from a guided six-section interface, calibrated abstention, human review, synthetic "
        "samples and visitor-isolated archives. Hosting and third-party service terms may still "
        "impose costs. Ethical feasibility depends on maintaining the evidence boundary and "
        "blocking deployment while the derived-model redistribution basis is unresolved.",
    )
    set_by_prefix(doc, "Figure 4.10.", "Figure 4.10. Target public deployment architecture.", style="Caption")
    set_by_prefix(
        doc,
        "GitHub is canonical;",
        f"After owner publication, {INTENDED_REPOSITORY} is intended to be canonical for code, "
        "documentation, issues, releases and deployment history. Vercel is only the presentation "
        "shell and Streamlit Community Cloud would run app.py with temporary visitor-isolated "
        "SQLite. This architecture is a target, not a live-deployment claim; functional hosting "
        "is blocked by the model-redistribution decision record.",
    )
    set_by_prefix(
        doc,
        "Exact deduplication precedes",
        "Exact deduplication and a deterministic word-five-gram near-duplicate screen precede "
        "partition use. The screen compared 41,994 candidates, found 17 near-duplicate pairs, "
        "quarantined two holdout rows that crossed the training boundary and left zero detected "
        "cross-partition pairs. The 2,399-row final test is excluded from model fitting, Platt "
        "calibration, model-retention policy and review-threshold selection.",
    )
    split_replacements = {
        "Deduplicate before train/test separation.": "Remove exact duplicates and group detected near-duplicates before partition use.",
        "Use seed 42 and stratified 80/20 split.": "Use seed 42, stratification and group-aware validation/test partitioning.",
        "Keep the 4,800-row test split unseen during tuning.": "Use 19,200 training rows; keep the 1,199 calibration and 1,200 policy rows outside fitting.",
        "Place TF-IDF and classifier in one scikit-learn Pipeline.": "Keep the 2,399-row final test untouched until one reporting pass.",
        "Exclude source, subject and date.": "Place TF-IDF and classifier in one scikit-learn Pipeline and exclude source, subject and date.",
        "Apply identical source-marker neutralisation in training and inference.": "Apply identical source-marker neutralisation in training and inference; verify the packaged Joblib SHA-256.",
        "Report remaining outlet/topic/time artefacts as a limitation.": "Report remaining outlet/topic/time artefacts and non-exhaustive near-duplicate screening as limitations.",
    }
    for old, new in split_replacements.items():
        paragraph = find_paragraph(doc, exact=old, required=False)
        if paragraph is not None:
            set_paragraph(paragraph, new)
    set_by_prefix(
        doc,
        "Multinomial Naive Bayes provides",
        "The controlled comparison fixes the established deployment-compatible settings: "
        "Multinomial Naive Bayes alpha 0.1, Linear SVC C=1.0 and the unchanged production "
        "Logistic Regression C=2.0. A Platt mapping is fitted for each candidate on 1,199 "
        "calibration rows. On the 1,200-row validation-policy partition, Linear SVC's macro-F1 "
        "advantage over Logistic Regression is 0.002500, below the predeclared 0.01 tolerance; "
        "Logistic Regression is retained for direct coefficient explanations, compact deployment "
        "and preservation of the verified production artifact.",
    )
    set_by_prefix(
        doc,
        "The maximum class probability is mapped",
        "The production decision score is transformed by a private Platt calibration artifact "
        "whose model SHA-256 must match the active Joblib. A calibrated confidence below 0.59, "
        "inadequate input, unsupported-language hint or domain-mismatch heuristic returns "
        "'Editorial review required'. Otherwise the exact outcomes are 'Lower misleading-content "
        "risk indicated' or 'Higher misleading-content risk indicated'. Local TF-IDF coefficient "
        "contributions describe learned correlations, never evidence that an article is true or false.",
    )
    test_status = (
        "All 56 packaged pytest checks passed in the final Python 3.12 audit, including the 29 "
        "established checks and 27 calibration, review, privacy, analytics and release-hardening checks."
        if audit_passed
        else "The release package defines 56 packaged pytest checks, including the 29 established "
        "checks and 27 calibration, review, privacy, analytics and release-hardening checks. The "
        "final execution record is written only after the comprehensive audit completes."
    )
    set_by_prefix(
        doc,
        "Verification is layered.",
        f"Verification is layered. {test_status} Browser automation covers all six sections, "
        "direct routes, refresh, back/forward, keyboard activation, same-tab navigation, text "
        "analysis, summarization, classification, calibrated confidence, explanations, review "
        "updates, downloads, analytics, drift readiness and cross-visitor archive isolation at "
        "360, 390, 768, 1366 and 1920 pixel widths.",
    )
    set_by_prefix(
        doc,
        "Linear SVM had the highest",
        "On the 1,200-row validation-policy partition, Linear SVC reached 0.997500 macro-F1 and "
        "Logistic Regression reached 0.995000, a 0.002500 advantage below the predeclared 0.01 "
        "retention tolerance. Logistic Regression therefore remains the production family for "
        "direct signed explanations, compact deployment and an unchanged verified artifact. The "
        "untouched final test is used for reporting, not this decision.",
    )
    set_by_prefix(
        doc,
        "Final-test Logistic Regression:",
        "Final-test Logistic Regression: 2,399 rows; accuracy 0.992080; macro-F1 0.992080; "
        "ROC-AUC 0.999481; PR-AUC 0.999423; calibrated Brier score 0.006292; ten-bin ECE "
        "0.005295; median-run inference 0.503 ms/article; model size 819,447 bytes.",
    )
    set_by_prefix(doc, "Figure 8.1.", "Figure 8.1. Controlled candidate confusion matrices on the untouched final test.", style="Caption")
    set_by_prefix(
        doc,
        "All three classical baselines perform strongly",
        "All candidates use the same 19,200 training rows and private Platt-calibration protocol. "
        "Final-test differences are reported for transparency and do not retroactively tune the "
        "selection or review policy.",
    )
    set_by_prefix(doc, "Figure 8.2.", "Figure 8.2. Production Logistic Regression final-test confusion matrix.", style="Caption")
    set_by_prefix(
        doc,
        "On 2,400 examples per class",
        "Among 1,200 reliable-labelled and 1,199 misleading-labelled final-test rows, the "
        "production model produced 1,196 true negatives, 1,184 true positives, four false "
        "positives and 15 false negatives.",
    )
    set_by_prefix(doc, "Figure 8.3.", "Figure 8.3. Calibrated ROC and precision-recall curves on the final test.", style="Caption")
    set_by_prefix(
        doc,
        "The synthetic demonstration produced",
        "The synthetic higher-risk demonstration produced a 28-word summary, 80.42% compression "
        "and 99.29% calibrated confidence in the higher-risk direction. It is a workflow sample, "
        "not a factual adjudication or benchmark row.",
    )
    set_by_prefix(
        doc,
        "Screenshot set:",
        "Screenshot set: Figures 9.2-9.16 document the current warm beige Streamlit interface, "
        "including calibrated abstention, model benchmarking, newsroom analytics, drift readiness "
        "and human editorial review. Dimensions and SHA-256 hashes are recorded in the screenshot manifest.",
    )
    set_by_prefix(
        doc,
        "Capture record:",
        "Capture record: The 15 interface images are genuine local Streamlit captures produced "
        "with the current ui/theme.py design and populated synthetic/session-safe evidence. The "
        "manifest records filename, dimensions, capture time and SHA-256; no image was repainted "
        "to simulate application output.",
    )
    set_by_prefix(
        doc,
        "Misinformation labels can amplify",
        "Dataset labels and model outcomes can amplify political, regional and publisher bias. "
        "False positives may damage legitimate reporting; false negatives may create false "
        "reassurance; calibrated percentages may be mistaken for factual probabilities; and "
        "feature contributions may be mistaken for evidence. NewsLens AI therefore uses bounded "
        "risk language, supports abstention and documented human review, keeps each public visitor's "
        "archive isolated, and requires independent verification for consequential claims.",
    )
    future_calibration = find_paragraph(doc, exact="Probability calibration on genuinely out-of-domain validation data.", required=False)
    if future_calibration is not None:
        set_paragraph(
            future_calibration,
            "External-domain calibration, publisher/time-group evaluation and threshold revalidation without touching the existing final test.",
        )
    conclusion_status = (
        "All 56 packaged checks passed in the final Python 3.12 audit"
        if audit_passed
        else "The final audit is configured to execute all 56 packaged checks"
    )
    set_by_prefix(
        doc,
        "NewsLens AI demonstrates an end-to-end",
        "NewsLens AI demonstrates an end-to-end editorial NLP system rather than two disconnected "
        "snippets. Summarization and linguistic credibility-risk classification consume the "
        "original cleaned article independently. The Streamlit layer adds calibrated confidence, "
        "explainability, abstention, session-private review/history, privacy-safe analytics, drift "
        f"readiness and exports without runtime retraining. {conclusion_status}; same-tab routing, "
        "responsive layouts, browser history, downloads and visitor isolation are included in the "
        "browser audit boundary.",
    )
    set_by_prefix(
        doc,
        "The held-out ISOT performance is technically strong",
        "The leakage-controlled final-test ISOT performance is technically strong but scientifically "
        "bounded by dataset artefacts. The low extractive XSum ROUGE is also reported rather than "
        "hidden. The central conclusion is not that the system determines truth: it supplies an "
        "inspectable linguistic credibility-risk signal and concise summary while preserving the "
        "need for evidence, external validation and human judgement.",
    )
    appendix_instruction = find_paragraph(doc, exact="Read the summary, probability, confidence band, local terms and disclaimer.", required=False)
    if appendix_instruction is not None:
        set_paragraph(
            appendix_instruction,
            "Read the summary, calibrated probabilities, outcome/review state, diagnostics, local term contributions and disclaimer.",
        )
    set_by_prefix(
        doc,
        "Run python -m pytest -q and confirm",
        "Run python -m pytest -q and confirm all 56 checks pass, including the 29 established checks.",
    )

    set_table(
        doc.tables[10],
        [
            ["Model", "Accuracy", "Macro-F1", "ROC-AUC", "PR-AUC", "Brier / ECE"],
            ["Linear SVC", "0.994581", "0.994581", "0.999851", "0.999848", "0.004059 / 0.004451"],
            ["Logistic Regression", "0.992080", "0.992080", "0.999481", "0.999423", "0.006292 / 0.005295"],
            ["Multinomial Naive Bayes", "0.960817", "0.960815", "0.991564", "0.992097", "0.029562 / 0.009859"],
        ],
    )
    set_table(
        doc.tables[11],
        [
            ["Metric", "Final-test value"],
            ["Accuracy", "0.992080"],
            ["Macro-F1", "0.992080"],
            ["ROC-AUC", "0.999481"],
            ["PR-AUC", "0.999423"],
            ["Calibrated Brier score", "0.006292"],
            ["Ten-bin ECE", "0.005295"],
            ["Editorial-review threshold", "0.59 calibrated confidence"],
            ["Final-test rows", "2,399"],
        ],
    )
    set_table(
        doc.tables[13],
        [
            ["Synthetic sample", "Words", "Summary", "Compression", "Displayed outcome", "Calibrated misleading p", "Band"],
            ["misleading_style_article.txt", "143", "28", "80.42%", "Higher misleading-content risk indicated", "99.29%", "High"],
            ["reliable_style_article.txt", "158", "41", "74.05%", "Lower misleading-content risk indicated", "15.38%", "Moderate"],
            ["uncertain_style_article.txt", "135", "38", "71.85%", "Editorial review required", "49.89%", "Review"],
        ],
    )
    doc.tables[5].rows[2].cells[1].text = "3.12"
    doc.tables[5].rows[2].cells[2].text = "Tested CPython 3.12, 64-bit"
    doc.tables[9].rows[6].cells[1].text = (
        "Verified model/calibration loading, calibrated probabilities, abstention outcomes, diagnostics and PredictionResult"
    )
    doc.tables[15].rows[-1].cells[-1].text = (
        "56 checks (including 29 established), responsive screenshots, diagrams, report and guides"
    )
    format_table(doc.tables[15])
    doc.tables[14].rows[22].cells[2].text = "All 29 established checks remain included"
    doc.tables[14].rows[22].cells[3].text = (
        "29 established plus 27 hardening checks passed"
        if audit_passed
        else "29 established plus 27 hardening checks collected; final run pending"
    )
    doc.tables[14].rows[22].cells[4].text = "Pass" if audit_passed else "Pending final audit"
    doc.tables[14].rows[23].cells[3].text = "Macro-F1 0.992080; ROC-AUC 0.999481; Brier 0.006292; ECE 0.005295"
    doc.tables[14].rows[6].cells[2].text = "Load without fitting; valid calibrated probabilities, outcome/review state and explanation"
    doc.tables[14].rows[29].cells[2].text = "Bounded risk outcomes, editorial-review state and verification warning remain explicit"
    format_table(doc.tables[14])

    report_picture_map = {
        "Figure 4.1.": (ASSETS / "system-architecture.png", "Layered NewsLens AI system architecture"),
        "Figure 4.8.": (DIAGRAMS / "08_component_module_diagram.png", "NewsLens AI component and module dependency diagram"),
        "Figure 4.9.": (DIAGRAMS / "11_sqlite_er_diagram.png", "Visitor-scoped SQLite analyses entity diagram"),
        "Figure 4.10.": (DIAGRAMS / "13_deployment_diagram.png", "Target GitHub, Vercel and Streamlit deployment architecture with privacy and licensing gates"),
        "Figure 5.1.": (DIAGRAMS / "09_ml_training_pipeline.png", "Leakage-controlled benchmark, calibration and untouched final-test workflow"),
        "Figure 6.1.": (DIAGRAMS / "10_combined_inference_pipeline.png", "Independent summarization and calibrated risk inference workflow"),
        "Figure 8.1.": (FIGURES / "model_benchmark_confusion_matrices.png", "Final-test confusion matrices for three controlled candidate models"),
        "Figure 8.2.": (FIGURES / "confusion_matrix.png", "Production Logistic Regression confusion matrix on 2,399 final-test rows"),
        "Figure 8.3.": (FIGURES / "roc_pr_curves.png", "Calibrated ROC and precision-recall curves for the production model"),
        "Figure 9.1.": (DIAGRAMS / "12_streamlit_navigation_diagram.png", "Native same-tab Streamlit navigation across six product sections"),
        "Figure 9.2.": (SCREENSHOTS / "01_home.png", "Current NewsLens AI News Desk desktop interface"),
        "Figure 9.3.": (SCREENSHOTS / "02_analysis_input.png", "Current article-analysis input controls"),
        "Figure 9.4.": (SCREENSHOTS / "03_summary_and_risk_results.png", "Current summary and calibrated linguistic-risk result"),
        "Figure 9.5.": (SCREENSHOTS / "04_explainability_and_downloads.png", "Current explainability and download controls"),
        "Figure 9.6.": (SCREENSHOTS / "06_model_accountability.png", "Current Model Accountability interface"),
        "Figure 9.7.": (SCREENSHOTS / "09_dataset_analysis.png", "Current Dataset Analysis interface"),
        "Figure 9.8.": (SCREENSHOTS / "10_newsroom_analytics.png", "Current session-private Editorial Archive and newsroom analytics interface"),
        "Figure 9.9.": (SCREENSHOTS / "13_research_about.png", "Current Research and About interface"),
        "Figure 9.10.": (SCREENSHOTS / "14_home_mobile.png", "Current mobile News Desk at 390 pixels"),
        "Figure 9.11.": (SCREENSHOTS / "15_analysis_mobile.png", "Current mobile analysis workflow at 390 pixels"),
    }
    for caption, (path, alt_text) in report_picture_map.items():
        replace_picture_by_caption(doc, caption, path, alt_text=alt_text)

    old_summarization_heading = find_paragraph(doc, exact="8.3 Summarization evaluation", required=False)
    if old_summarization_heading is not None:
        insert_before(old_summarization_heading, "8.3 Leakage-controlled evidence protocol", "Heading 2")
        insert_before(
            old_summarization_heading,
            "The unchanged production artifact first reproduced the established 4,800-row holdout "
            "accuracy and macro-F1 exactly. For the new study, 41,994 near-duplicate candidates were "
            "screened; 17 pairs met the 0.90 five-gram Jaccard rule; two contaminated holdout rows "
            "were quarantined; and zero detected near-duplicate pairs crossed the final train, "
            "validation and test boundaries. Candidate generation is approximate and does not claim "
            "semantic-duplicate exhaustiveness.",
            "Normal",
        )
        insert_before(old_summarization_heading, "8.4 Calibration and editorial-review policy", "Heading 2")
        insert_before(
            old_summarization_heading,
            "Platt scaling is fitted on 1,199 calibration rows. The 1,200-row validation-policy "
            "partition applies the predeclared 0.01 model-retention tolerance and selects 0.59 as "
            "the lowest confidence threshold meeting at least 80% coverage and a 95% Wilson lower "
            "accuracy bound of at least 99%. On the untouched final test, calibrated Brier score is "
            "0.006292 and ten-bin ECE is 0.005295, compared with 0.010464 and 0.044799 for native "
            "Logistic Regression probabilities. The review rate is 0.167%, automatic-decision "
            "coverage 99.833%, selective accuracy 99.290% and its 95% Wilson lower bound 98.866%. "
            "These values are relative to ISOT labels and do not verify facts.",
            "Normal",
        )
        add_figure_before(
            old_summarization_heading,
            FIGURES / "calibration_reliability.png",
            "Figure 8.5. Production-model reliability before and after Platt calibration.",
            "Calibration substantially reduces Brier score and expected calibration error on the "
            "untouched final test, while remaining a dataset-relative confidence study rather than "
            "a truth-probability claim.",
            width=5.9,
            max_height=6.2,
            page_break=False,
        )
        set_paragraph(old_summarization_heading, "8.5 Summarization evaluation", style="Heading 2")
        set_paragraph(find_paragraph(doc, exact="8.4 Packaged demonstration analyses"), "8.6 Packaged demonstration analyses", style="Heading 2")
        set_paragraph(find_paragraph(doc, exact="8.5 Test-case summary"), "8.7 Test-case summary", style="Heading 2")
        set_paragraph(find_paragraph(doc, exact="8.6 Error analysis"), "8.8 Error analysis", style="Heading 2")

    chapter_ten = find_paragraph(doc, exact="Chapter 10: Project Management")
    if find_paragraph(doc, startswith="Figure 9.12.", required=False) is None:
        additions = [
            (
                SCREENSHOTS / "05_editorial_review_required.png",
                "Figure 9.12. Calibrated abstention and editorial-review-required state.",
                "Low calibrated confidence and supported mismatch conditions route the article to review instead of forcing a binary risk outcome.",
            ),
            (
                SCREENSHOTS / "07_model_benchmarking.png",
                "Figure 9.13. Controlled model benchmarking evidence.",
                "Candidate metrics, fixed partition counts and the unchanged production-model decision are visible in Model Accountability.",
            ),
            (
                SCREENSHOTS / "08_calibration_reliability.png",
                "Figure 9.14. Calibration reliability and abstention-policy evidence.",
                "Brier score, ECE, reliability curves and the dataset-relative interpretation accompany the calibrated confidence display.",
            ),
            (
                SCREENSHOTS / "11_drift_readiness.png",
                "Figure 9.15. Privacy-safe drift-readiness indicators.",
                "Session aggregates are compared with reference ranges only after the minimum observation count; warnings indicate change, not automatic failure.",
            ),
            (
                SCREENSHOTS / "12_editorial_review_workflow.png",
                "Figure 9.16. Human editorial-review workflow.",
                "A visitor can record status, bounded notes, public supporting URLs and a final assessment inside the same isolated session archive.",
            ),
        ]
        for path, caption, explanation in additions:
            add_figure_before(chapter_ten, path, caption, explanation)

    list_of_tables = find_paragraph(doc, exact="List of Tables")
    if find_paragraph(doc, exact="Figure 8.5 Calibration reliability", required=False) is None:
        insert_before(find_paragraph(doc, exact="Figure 9.1 Streamlit navigation"), "Figure 8.5 Calibration reliability", "List Bullet")
    set_by_prefix(
        doc,
        "Figures 9.2-9.11",
        "Figures 9.2-9.16 Current desktop/mobile application, review, benchmark, analytics and drift screenshots",
        style="List Bullet",
    )
    _ = list_of_tables

    common_document_finish(doc, output)
    return output


def reconcile_developer_guide(*, audit_passed: bool) -> Path:
    output = DOCS / "NewsLens_AI_Code_Explanation_and_Developer_Guide.docx"
    doc = Document(output)
    update_common_metadata(doc)
    update_cover_metadata(doc.tables[0])
    update_page_reference_table(
        doc.tables[1],
        {
            "1-2. Architecture and folder map": "3",
            "3-4. Configuration and ingestion": "5",
            "5-7. Preprocessing and summarizers": "6",
            "8. Offline training pipeline": "7",
            "9. Prediction and explainability": "8",
            "10-11. SQLite and Streamlit orchestration": "10",
            "12-13. Visualisation, exports and tests": "12",
            "14-16. Modification, deployment and errors": "13",
            "17-18. Viva and acceptance checklist": "16",
        },
    )

    set_by_prefix(
        doc,
        "src/config.py derives",
        "src/config.py derives PROJECT_ROOT from its own file, so paths remain portable. "
        "MODEL_PATH and CALIBRATION_PATH may be overridden privately; the default public history "
        "mode is fail-closed per-session SQLite. Configuration records the 40-word minimum, bounded "
        "uploads/network requests, model version isot-tfidf-lr-v1.0.0, 56 expected packaged checks "
        "and exact responsible outcome labels. Runtime code reads fitted artifacts and never calls fit().",
    )
    set_by_prefix(
        doc,
        "MODEL_PATH = Path(os.getenv",
        "MODEL_PATH = Path(os.getenv('NEWSLENS_MODEL_PATH', MODELS_DIR / 'fake_news_pipeline.joblib'))\n"
        "CALIBRATION_PATH = Path(os.getenv('NEWSLENS_CALIBRATION_PATH', MODELS_DIR / 'confidence_calibration.json'))\n"
        "EXPECTED_PACKAGED_CHECKS = 56\n"
        "# NEWSLENS_HISTORY_MODE defaults to session; runtime loads artifacts and never retrains.",
        style="Code Block",
    )
    set_by_prefix(
        doc,
        "load_isot_dataset assigns",
        "load_isot_dataset assigns labels from the private checksum-verified CSV filenames, "
        "normalises text, rejects short/conflicting rows, removes exact duplicate model-text hashes "
        "and draws a balanced seed-42 sample. benchmark_models.py then applies a deterministic "
        "word-five-gram candidate screen with exact Jaccard verification so detected near-duplicate "
        "groups remain within one partition.",
    )
    set_by_prefix(
        doc,
        "base_pipeline defines",
        "The controlled release benchmark uses 19,200 training rows, 1,199 Platt-calibration rows, "
        "1,200 validation-policy rows and an untouched 2,399-row final test. Logistic Regression "
        "C=2.0, LinearSVC C=1.0 and MultinomialNB alpha=0.1 share the same TF-IDF configuration. "
        "The validation-policy macro-F1 difference (0.002500) is below the predeclared 0.01 "
        "retention tolerance, so the verified packaged Logistic Regression remains unchanged.",
    )
    old_persistence_heading = find_paragraph(doc, exact="8.3 Champion persistence", required=False)
    if old_persistence_heading is not None:
        set_paragraph(old_persistence_heading, "8.3 Production-artifact verification and private calibration", style="Heading 2")
    set_by_prefix(
        doc,
        "joblib.dump(champion",
        "assert sha256(MODEL_PATH) == accepted_model_sha256\n"
        "production = joblib.load(MODEL_PATH)  # unchanged verified artifact\n"
        "# Platt parameters are fitted offline and stored in private confidence_calibration.json.\n"
        "# Streamlit loads both matching artifacts; it never calls fit().",
        style="Code Block",
    )
    set_by_prefix(
        doc,
        "Metadata records the model artifact ID",
        "The benchmark summary records source/model checksums, exact partition hashes, leakage-screen "
        "counts, candidate metrics, policy selection, calibration evidence and final-test results. "
        "The private calibration artifact binds its coefficient/intercept and 0.59 review threshold "
        "to the exact production-model SHA-256.",
    )
    old_probability_heading = find_paragraph(doc, exact="9.2 Probability and band", required=False)
    if old_probability_heading is not None:
        set_paragraph(old_probability_heading, "9.2 Calibrated probability and outcome policy", style="Heading 2")
    set_by_prefix(
        doc,
        "_probabilities maps",
        "predict_credibility converts the production decision_function through the verified Platt "
        "mapping. It returns calibrated reliable/misleading probabilities, calibrated confidence, "
        "a band and review reasons. Confidence below 0.59 or supported mismatch/quality conditions "
        "yield Editorial review required; otherwise the result uses the exact lower- or higher-risk outcome text.",
    )
    xai_heading = find_paragraph(doc, exact="9.3 Local linear contribution", required=False)
    if xai_heading is not None:
        insert_before(xai_heading, "9.3 Calibration, abstention and artifact binding", "Heading 2")
        insert_before(
            xai_heading,
            "src/calibration.py refuses missing, incomplete, out-of-range or model-hash-mismatched "
            "calibration data. Platt parameters transform a one-dimensional decision score only; they "
            "do not alter classifier coefficients or retrain the model. The review threshold is a "
            "communication/abstention policy derived from separate validation-policy rows.",
        "Normal",
        )
        set_paragraph(xai_heading, "9.4 Local linear contribution", style="Heading 2")
    set_by_prefix(
        doc,
        "src/calibration.py refuses",
        "src/calibration.py refuses missing, incomplete, out-of-range or model-hash-mismatched "
        "calibration data. Platt parameters transform a one-dimensional decision score only; they "
        "do not alter classifier coefficients or retrain the model. The review threshold is a "
        "communication/abstention policy derived from separate validation-policy rows.",
    )
    chapter_ten = find_paragraph(doc, exact="10. SQLite operations")
    if find_paragraph(doc, exact="9.5 Input diagnostics and editorial review", required=False) is None:
        insert_before(chapter_ten, "9.5 Input diagnostics and editorial review", "Heading 2")
        insert_before(
            chapter_ten,
            "src/model_diagnostics.py compares article length and unigram vocabulary coverage with "
            "saved reference quantiles and adds lightweight language/domain mismatch warnings. "
            "src/editorial_review.py validates bounded notes, statuses and up to eight public HTTP(S) "
            "supporting URLs. The workflow records a human assessment; it never presents the model "
            "result as evidence.",
            "Normal",
        )
    set_by_prefix(
        doc,
        "The unique article_hash performs",
        "The unique article_hash performs duplicate detection without retaining the complete source "
        "article. Structured rows include calibrated confidence, review status/reasons, public "
        "supporting-source URLs and aggregate-safe coverage/mismatch fields. In public mode the "
        "database file remains unguessable and visitor-scoped.",
    )
    set_by_prefix(
        doc,
        "confidence_gauge maps",
        "confidence_gauge and risk displays use calibrated probabilities and exact bounded outcome "
        "language. The contribution chart remains a signed correlation view. Newsroom analytics "
        "export only aggregate counts/rates; review/history exports apply spreadsheet/PDF escaping. "
        "All views use the shared warm editorial palette and visible text labels.",
    )
    set_by_prefix(
        doc,
        "python -m pytest -q",
        "python -m pytest -q  # Expected collection: 56 checks, including 29 established checks",
        style="Code Block",
    )
    set_by_prefix(doc, "Figure D.8.", "Figure D.8. Target public deployment architecture.", style="Caption")
    set_by_prefix(
        doc,
        "Vercel provides the editorial",
        "Vercel is the intended editorial presentation shell while Streamlit Community Cloud would "
        "run the unchanged Python/ML application from app.py. Neither service is claimed live in "
        "this staging document.",
    )
    set_by_prefix(
        doc,
        "The public Streamlit UI defaults",
        f"The public Streamlit UI defaults to temporary per-session SQLite and makes no durable-"
        f"history promise. {INTENDED_REPOSITORY} is the intended canonical repository, but it has "
        "not yet been created. Public push and functional hosting remain blocked because the "
        "ISOT-derived Joblib and matching calibration redistribution basis is unresolved. The public "
        "staging archive excludes both artifacts; no runtime retraining or mock classifier is used.",
    )
    set_by_prefix(
        doc,
        "It stayed within 0.01 CV macro-F1",
        "On the validation-policy partition, Linear SVC's macro-F1 advantage was 0.002500, below the "
        "predeclared 0.01 tolerance. The unchanged Logistic Regression remains for direct signed "
        "explanations, compact deployment and a verified artifact hash; its confidence is separately Platt-calibrated.",
    )
    old_accuracy_heading = find_paragraph(doc, exact="Why is 99.35% accuracy suspicious?", required=False)
    if old_accuracy_heading is not None:
        set_paragraph(old_accuracy_heading, "Why is 99.21% final-test accuracy not universal?", style="Heading 2")
    set_by_prefix(
        doc,
        "ISOT labels correlate with outlet",
        "The 99.21% value is accuracy against 2,399 final-test ISOT labels after leakage controls; "
        "ISOT still carries outlet, topic, time and style correlations, so this is not an external-domain truth score.",
    )
    test_status = (
        "The final audit executed all 56 checks successfully."
        if audit_passed
        else "The final audit will execute all 56 checks and write the authoritative JUnit record."
    )
    checklist_anchor = find_paragraph(doc, exact="The final ZIP excludes raw copyrighted/large datasets and transient caches.")
    if find_paragraph(doc, exact="Calibration loads only when its model SHA-256 matches the active Joblib.", required=False) is None:
        insert_before(checklist_anchor, "Calibration loads only when its model SHA-256 matches the active Joblib.", "List Bullet")
        insert_before(checklist_anchor, "Review rows, notes and analytics remain isolated to one visitor session in public mode.", "List Bullet")
        insert_before(checklist_anchor, "Drift readiness waits for at least 20 observations and never retrains the model.", "List Bullet")
        insert_before(checklist_anchor, test_status, "List Bullet")

    doc.tables[3].rows[3].cells[1].text = (
        "raw/calibrated probabilities, calibrated confidence/band, exact outcome, review state/reasons, diagnostics, artifact ID and time"
    )
    append_table_row(
        doc.tables[4],
        [
            "test_placement_enhancements.py",
            "calibration/hash binding, abstention outcomes, diagnostics, editorial review, analytics, drift readiness and privacy isolation",
        ],
    )

    developer_picture_map = {
        "Figure D.1.": (DIAGRAMS / "01_overall_system_architecture.png", "Layered NewsLens AI source architecture"),
        "Figure D.2.": (DIAGRAMS / "08_component_module_diagram.png", "Current NewsLens AI component and module dependencies"),
        "Figure D.3.": (DIAGRAMS / "09_ml_training_pipeline.png", "Leakage-controlled offline benchmark and calibration pipeline"),
        "Figure D.4.": (DIAGRAMS / "10_combined_inference_pipeline.png", "Runtime summarization, calibrated risk and review contracts"),
        "Figure D.5.": (DIAGRAMS / "11_sqlite_er_diagram.png", "Session-isolated analyses and review schema"),
        "Figure D.6.": (DIAGRAMS / "12_streamlit_navigation_diagram.png", "Native same-tab Streamlit navigation"),
        "Figure D.7.": (SCREENSHOTS / "04_explainability_and_downloads.png", "Current explanation and export surface"),
        "Figure D.8.": (DIAGRAMS / "13_deployment_diagram.png", "Target public deployment architecture and blocking gates"),
    }
    for caption, (path, alt_text) in developer_picture_map.items():
        replace_picture_by_caption(doc, caption, path, alt_text=alt_text)

    common_document_finish(doc, output)
    return output


def reconcile_concepts_guide(*, audit_passed: bool) -> Path:
    output = DOCS / "NewsLens_AI_Complete_Concepts_Methodologies_and_Terminology_Guide.docx"
    doc = Document(output)
    update_common_metadata(doc)
    set_by_prefix(
        doc,
        "Prepared as a companion to the project report and source package · 16 August 2026",
        f"Prepared as a companion to the project report and source package · {UPDATED_DATE}",
    )
    for row in doc.tables[0].rows:
        key = normalised(row.cells[0].text)
        if key == "Author and developer":
            row.cells[1].text = PUBLIC_AUTHOR
        elif key == "Copyright":
            row.cells[1].text = PUBLIC_COPYRIGHT
        elif key == "Evidence basis":
            row.cells[1].text = "Current source, verified model/calibration hashes, benchmark CSV/JSON, 15 screenshots and 56 checks"
    format_table(doc.tables[0])

    set_by_prefix(
        doc,
        "The 24,000-row sample is split",
        "The balanced 24,000-row sample keeps 19,200 training rows. The cleaned 4,800-row original "
        "holdout is used first to reproduce the accepted model checkpoint; two detected cross-"
        "boundary near-duplicate rows are then quarantined. The remaining holdout is split into "
        "2,399 validation rows (1,199 calibration and 1,200 policy) and an untouched 2,399-row "
        "final test. The final test does not fit a pipeline, Platt mapping, retention rule or threshold.",
    )
    set_by_prefix(
        doc,
        "Held-out test set. Examples",
        "Untouched final test set. The 2,399 examples excluded from fitting, calibration and policy "
        "selection and used once for final reporting against dataset labels.",
    )
    set_by_prefix(
        doc,
        "The held-out test set follows",
        "The untouched final test follows a locked path. Candidate pipelines fit on 19,200 rows; "
        "Platt mappings fit on 1,199 calibration rows; model-retention and review-threshold policies "
        "use 1,200 separate validation-policy rows; the final test is reported once afterward.",
    )
    set_by_prefix(
        doc,
        "In 3-fold CV, the 19,200",
        "The original model-development run used three-fold GridSearchCV only inside the 19,200 "
        "training rows. The controlled release benchmark reuses the established selected settings "
        "for three candidate families rather than performing a new exhaustive search. Neither path "
        "uses the 2,399-row final test for fitting or policy selection.",
    )
    set_by_prefix(
        doc,
        "The default objective is the highest mean",
        "The predeclared release rule retains Logistic Regression when Linear SVC's macro-F1 "
        "advantage on the 1,200-row validation-policy partition is below 0.01 and the production "
        "artifact remains hash-verified. The measured policy macro-F1 values are 0.995000 for "
        "Logistic Regression and 0.997500 for Linear SVC, a 0.002500 advantage. The final test is "
        "reserved for reporting after this decision.",
    )
    set_by_prefix(
        doc,
        "In one sentence: The classifier learned from cleaned",
        "In one sentence: The classifier learned from cleaned, deduplicated and balanced ISOT "
        "articles; release evidence separates training, calibration, validation-policy decisions "
        "and an untouched final test.",
    )
    set_by_prefix(doc, "Figure 8.2.", "Figure 8.2. Controlled candidate comparison on the untouched final test.", style="Caption")
    set_by_prefix(doc, "Figure 9.1.", "Figure 9.1. Production Logistic Regression final-test confusion matrix.", style="Caption")
    set_by_prefix(
        doc,
        "All three classical models perform",
        "All three classical models perform strongly on the final-test ISOT distribution. These "
        "reported differences do not tune the already fixed model-retention or review policy.",
    )
    metric_updates = {
        "A confusion matrix counts every combination": "A confusion matrix counts every combination of dataset label and predicted direction. The final test has 1,200 reliable-labelled and 1,199 misleading-labelled rows.",
        "True positive (TP).": "True positive (TP). A misleading-labelled article assigned the higher-risk direction. Count: 1,184.",
        "True negative (TN).": "True negative (TN). A reliable-labelled article assigned the lower-risk direction. Count: 1,196.",
        "False positive (FP).": "False positive (FP). A reliable-labelled article assigned the higher-risk direction. Count: 4. This can unfairly cast doubt on credible reporting.",
        "False negative (FN).": "False negative (FN). A misleading-labelled article assigned the lower-risk direction. Count: 15. This can create false reassurance.",
        "The diagonal cells are correct predictions": "The diagonal cells are agreements with dataset labels; off-diagonal cells are errors. There are 19 errors among 2,399 final-test rows.",
        "The measured accuracy is": "The measured final-test accuracy is 0.992080: 2,380 of 2,399 dataset labels agree with the predicted direction.",
        "Measured precision is": "Measured misleading-class precision is 0.996633. Among rows assigned the higher-risk direction, 99.66% carry the misleading dataset label in this final test.",
        "Measured recall is": "Measured misleading-class recall is 0.987490. The model assigns the higher-risk direction to 98.75% of misleading-labelled final-test rows.",
        "Measured positive-class F1 is": "Measured misleading-class F1 is 0.992040; it balances precision and recall through their harmonic mean.",
        "Support. The number": "Support is the number of true examples used for a class metric: 1,200 reliable-labelled and 1,199 misleading-labelled rows.",
        "Macro-F1. Calculate": "Macro-F1 calculates F1 for each class and takes an equal average. The production final-test macro-F1 is 0.992080.",
        "Weighted-F1. Average": "Weighted-F1 weights each class by support. Because final-test support is almost balanced, it is approximately 0.992080 and close to macro-F1.",
        "ROC-AUC. Area": "ROC-AUC measures ranking discrimination across thresholds. The calibrated production final-test ROC-AUC is 0.999481.",
        "PR-AUC / Average Precision.": "PR-AUC / Average Precision summarises the precision-recall trade-off. The calibrated production final-test value is 0.999423.",
        "Near-perfect curves show": "Near-perfect final-test discrimination shows strong ISOT ranking. It does not establish factual truth, fairness or external-domain generalisation; calibration is assessed separately.",
        "CV mean. The average": "Brier score is the mean squared error of a probability against a binary dataset label. Platt calibration reduces the production final-test Brier score from 0.010464 to 0.006292.",
        "CV standard deviation.": "Expected calibration error (ECE) summarises confidence-frequency gaps across bins. Ten-bin ECE decreases from 0.044799 for native probabilities to 0.005295 after Platt scaling.",
        "Latency is elapsed processing time.": "Latency is elapsed processing time. The controlled run measured a median-run mean of 0.503 ms per final-test article for production prediction plus calibration.",
        "Training time. Time": "Model size is 819,447 bytes. Runtime loads this existing artifact and matching calibration parameters; it never performs training.",
        "Metric discipline:": "Metric discipline: state ISOT, the exact partition, label mapping, row count, calibration/policy separation and limitations. '99.21% final-test accuracy' alone is incomplete.",
    }
    for prefix, text in metric_updates.items():
        set_by_prefix(doc, prefix, text)
    old_metric_heading = find_paragraph(doc, exact="9.5 Cross-validation statistics and latency", required=False)
    if old_metric_heading is not None:
        set_paragraph(old_metric_heading, "9.5 Calibration error, abstention and latency", style="Heading 2")
    set_by_prefix(
        doc,
        "The raw predicted class uses",
        "A 0.50 calibrated misleading-probability boundary determines direction. Calibrated "
        "confidence is the larger class probability. If confidence is below the validated 0.59 "
        "review threshold—or input quality/language/domain diagnostics require caution—the interface "
        "abstains and displays Editorial review required instead of forcing a binary outcome.",
    )
    set_by_prefix(doc, "Confidence. The largest", "Calibrated confidence. The larger probability after applying the verified Platt mapping.")
    set_by_prefix(doc, "Confidence band.", "Confidence band. Review below 0.59, Moderate from 0.59 to below 0.80, and High from 0.80 upward.")
    set_by_prefix(
        doc,
        "Probability caveat:",
        "Calibration study: Platt scaling is fitted on 1,199 calibration rows and evaluated only "
        "after policy selection on the untouched final test. Brier score and ECE improve materially. "
        "The percentage remains agreement-oriented confidence against ISOT labels, not a verified "
        "chance that an article is factually true or false.",
    )
    xai_heading = find_paragraph(doc, exact="10.2 Explainable AI (XAI)", required=False)
    if xai_heading is not None:
        insert_before(xai_heading, "10.2 Platt calibration, abstention and human review", "Heading 2")
        insert_before(
            xai_heading,
            "Platt scaling fits a one-dimensional logistic mapping from the production decision score "
            "to the misleading-label probability. Abstention means declining to make an automatic "
            "lower/higher-risk communication when calibrated confidence or supported input diagnostics "
            "do not justify it. A human reviewer can record a bounded status, notes, public supporting "
            "URLs and a final assessment; the human record is not used to retrain the model.",
            "Normal",
        )
        set_paragraph(xai_heading, "10.3 Explainable AI (XAI)", style="Heading 2")
    old_explanation_heading = find_paragraph(doc, exact="10.3 What the explanation does not mean", required=False)
    if old_explanation_heading is not None:
        set_paragraph(old_explanation_heading, "10.4 What the explanation does not mean", style="Heading 2")
    chapter_eleven = find_paragraph(doc, startswith="Chapter 11:")
    if find_paragraph(doc, exact="10.5 Diagnostics, newsroom analytics and drift readiness", required=False) is None:
        insert_before(chapter_eleven, "10.5 Diagnostics, newsroom analytics and drift readiness", "Heading 2")
        insert_before(
            chapter_eleven,
            "Input diagnostics report word count, fitted-vocabulary coverage, out-of-vocabulary rate, "
            "a lightweight language hint and transparent length/coverage mismatch heuristics. "
            "Newsroom analytics aggregate only the current visitor's session. Drift readiness "
            "compares aggregate length, coverage, class share, confidence and mismatch rates with "
            "saved reference ranges only after at least 20 observations. A warning means the input "
            "distribution changed; it neither proves model failure nor triggers retraining.",
            "Normal",
        )
    set_by_prefix(
        doc,
        "Accessibility principle:",
        "Accessibility principle: Colour is reinforced by the exact text outcomes Lower misleading-content "
        "risk indicated, Higher misleading-content risk indicated and Editorial review required. "
        "The result never depends on red/green colour alone.",
    )
    test_status = (
        "All 56 passed in the final tested Python 3.12 environment."
        if audit_passed
        else "The final comprehensive audit is configured to execute all 56 and record the result."
    )
    set_by_prefix(
        doc,
        "The package defines",
        "The package defines 56 checks: the 29 established checks remain, and 27 hardening checks "
        "cover calibration/hash binding, abstention, diagnostics, human review, privacy-safe "
        f"analytics, drift readiness, visitor isolation and release policy. {test_status} Chromium "
        "audits cover all six sections, same-tab routes, direct navigation, refresh, back/forward, "
        "keyboard activation, analysis/download/review flows and five required responsive widths.",
    )
    set_by_prefix(
        doc,
        "Verification evidence:",
        "Verification evidence: reports/results/project_verification.json and pytest_results.xml are "
        "regenerated by the final audit; screenshots carry dimensions and SHA-256 hashes. Recorded "
        "evidence is never replaced by an unexecuted claim.",
    )
    set_by_prefix(
        doc,
        "Streamlit Community Cloud. A service",
        "Streamlit Community Cloud is the intended service for branch main and app.py after the "
        "repository exists and release gates clear. NewsLens AI is not claimed deployed in this guide.",
    )
    set_by_prefix(
        doc,
        "The packaged Logistic Regression pipeline correctly classified",
        "On the untouched 2,399-row final test, the packaged Logistic Regression agreed with 2,380 "
        "ISOT labels: 99.21% accuracy and 0.992080 macro-F1, with four false positives and 15 false "
        "negatives. Calibrated ROC-AUC is 0.999481, PR-AUC 0.999423, Brier 0.006292 and ECE 0.005295.",
    )
    set_by_prefix(
        doc,
        "Linear SVM achieved the highest measured macro-F1",
        "On the validation-policy partition, Linear SVC reached 0.997500 macro-F1 and Logistic "
        "Regression 0.995000. The 0.002500 advantage is below the predeclared 0.01 tolerance, so the "
        "unchanged, hash-verified Logistic Regression remains for direct coefficient explanations "
        "and compact deployment. Final-test metrics are reported only after this decision.",
    )
    set_by_prefix(
        doc,
        "The saved Logistic Regression pipeline performs",
        "The unchanged Logistic Regression pipeline performs strongly on a leakage-controlled, "
        "same-source ISOT final test and has separately measured calibration evidence.",
    )
    old_worldwide = find_paragraph(doc, exact="The model is 99.35% accurate on all news worldwide.", required=False)
    if old_worldwide is not None:
        set_paragraph(old_worldwide, "The model's 99.21% final-test accuracy applies to all news worldwide.")
    set_by_prefix(
        doc,
        "It remained within 0.01 CV macro-F1",
        "It remained within the predeclared 0.01 validation-policy macro-F1 tolerance, preserves the "
        "verified artifact and supports direct signed coefficients; a separate Platt mapping supplies calibrated confidence.",
    )
    old_question = find_paragraph(doc, exact="Why is 99.35% accuracy not enough?", required=False)
    if old_question is not None:
        set_paragraph(old_question, "Why is 99.21% final-test accuracy not enough?", style="Heading 2")
    set_by_prefix(
        doc,
        "ISOT contains publisher, topic",
        "ISOT contains publisher, topic, time and style correlations. Leakage controls reduce exact "
        "and detected near-duplicate shortcuts, but a same-source final test remains easier than "
        "genuinely new-domain news and does not measure factual truth.",
    )

    doc.tables[3].rows[1].cells[2].text = "Edit configuration/code; the editorial threshold comes from a validated private calibration artifact"
    doc.tables[5].rows[4].cells[2].text = "class/direction, calibrated probabilities/confidence/band, review state/reasons, diagnostics, artifact ID, time and explanation"
    doc.tables[13].rows[2].cells[2].text = "Selected production family: unchanged verified artifact, Platt-calibrated confidence and direct coefficients"
    set_table(
        doc.tables[15],
        [
            ["Candidate", "Accuracy", "Macro-F1", "ROC-AUC", "Brier / ECE"],
            ["Linear SVC", "0.994581", "0.994581", "0.999851", "0.004059 / 0.004451"],
            ["Logistic Regression", "0.992080", "0.992080", "0.999481", "0.006292 / 0.005295"],
            ["Multinomial Naive Bayes", "0.960817", "0.960815", "0.991564", "0.029562 / 0.009859"],
        ],
    )
    set_table(
        doc.tables[16],
        [
            ["Dataset label", "Predicted lower risk", "Predicted higher risk"],
            ["Reliable", "1,196 true negatives", "4 false positives"],
            ["Misleading", "15 false negatives", "1,184 true positives"],
        ],
    )
    set_table(
        doc.tables[17],
        [
            ["Condition", "Confidence band", "Displayed outcome"],
            ["Calibrated confidence below 0.59 or supported diagnostic caution", "Review", "Editorial review required"],
            ["Misleading p below 0.50 and confidence at least 0.59", "Moderate / High", "Lower misleading-content risk indicated"],
            ["Misleading p at least 0.50 and confidence at least 0.59", "Moderate / High", "Higher misleading-content risk indicated"],
        ],
    )
    doc.tables[20].rows[5].cells[1].text = "Session-private search/reopen/review/export/delete, privacy-safe analytics and drift readiness"
    doc.tables[22].rows[2].cells[2].text = "Never treat a lower-risk outcome as proof; verify consequential claims independently"
    append_table_row(doc.tables[21], ["Calibration/review", "model-hash binding, Platt mapping, 0.59 abstention, bounded review updates and public-URL validation"])
    append_table_row(doc.tables[21], ["Analytics/drift", "aggregate-only export, 20-observation minimum, reference comparisons and no runtime retraining"])
    append_table_row(doc.tables[24], ["src/calibration.py", "Private Platt-parameter loading, model-hash verification and calibrated probability mapping"])
    append_table_row(doc.tables[24], ["src/model_diagnostics.py", "Vocabulary/length/language diagnostics and aggregate drift-readiness checks"])
    append_table_row(doc.tables[24], ["src/editorial_review.py", "Bounded human review status, notes, assessments and public supporting-source URL validation"])
    append_table_row(doc.tables[24], ["src/newsroom_analytics.py", "Privacy-safe session aggregate metrics and CSV export"])
    glossary_additions = [
        (25, ["Abstention", "Withholding an automatic lower/higher-risk outcome and requiring editorial review."]),
        (25, ["Brier score", "Mean squared error between a probability and the binary dataset label; lower is better."]),
        (25, ["Calibration", "Agreement between predicted probabilities and observed dataset-label frequencies."]),
        (25, ["ECE", "Expected calibration error; a binned summary of confidence-frequency gaps."]),
        (25, ["Editorial review", "Human assessment recorded separately from the model outcome and not used for runtime retraining."]),
        (25, ["Drift readiness", "Aggregate reference checks that flag distributional change after sufficient observations."]),
        (27, ["Platt scaling", "A one-dimensional logistic mapping from model score to calibrated probability."]),
        (27, ["Selective accuracy", "Accuracy among examples the policy decides automatically rather than sends to review."]),
        (28, ["Vocabulary coverage", "Share of article unigram tokens found in the fitted TF-IDF vocabulary."]),
        (28, ["Wilson lower bound", "A conservative confidence lower bound for an observed binomial proportion."]),
    ]
    for table_index, row in glossary_additions:
        upsert_table_row_by_key(doc.tables[table_index], row)

    concepts_picture_map = {
        "Figure 1.1.": (DIAGRAMS / "10_combined_inference_pipeline.png", "Independent summary, calibrated risk and editorial-review branches"),
        "Figure 3.1.": (DIAGRAMS / "02_end_to_end_data_flow.png", "Complete NewsLens AI runtime data flow"),
        "Figure 4.1.": (FIGURES / "class_distribution.png", "Balanced 24,000-row modelling sample"),
        "Figure 8.1.": (DIAGRAMS / "09_ml_training_pipeline.png", "Leakage-controlled training, calibration, policy and final-test pipeline"),
        "Figure 8.2.": (FIGURES / "model_benchmark_confusion_matrices.png", "Final-test candidate confusion matrices"),
        "Figure 9.1.": (FIGURES / "confusion_matrix.png", "Production final-test confusion matrix"),
        "Figure 9.2.": (FIGURES / "roc_pr_curves.png", "Calibrated production ROC and precision-recall curves"),
        "Figure 10.1.": (FIGURES / "feature_importance.png", "Global production-model coefficient inspection"),
        "Figure 10.2.": (SCREENSHOTS / "04_explainability_and_downloads.png", "Current explanation, responsible-use and export interface"),
        "Figure 11.1.": (ASSETS / "system-architecture.png", "Current NewsLens AI layered architecture"),
        "Figure 13.1.": (DIAGRAMS / "11_sqlite_er_diagram.png", "Current session-scoped SQLite and editorial-review schema"),
        "Figure 14.1.": (DIAGRAMS / "12_streamlit_navigation_diagram.png", "Current native same-tab six-section Streamlit navigation"),
    }
    for caption, (path, alt_text) in concepts_picture_map.items():
        replace_picture_by_caption(doc, caption, path, alt_text=alt_text)

    common_document_finish(doc, output)
    return output


def reconcile_setup_guide(*, audit_passed: bool) -> Path:
    output = DOCS / "NewsLens_AI_Setup_and_Run_Guide.docx"
    doc = Document(output)
    update_common_metadata(doc)
    update_cover_metadata(doc.tables[0])
    update_page_reference_table(
        doc.tables[1],
        {
            "1-3. Prerequisites, Python and VS Code": "3",
            "4-6. Extract, virtual environment and interpreter": "4",
            "7-9. Dependencies and Streamlit startup": "5",
            "10. First sample analysis": "6",
            "11-15. Operations, datasets and retraining": "10",
            "16-17. Tests and public deployment": "12",
            "18. Troubleshooting": "12",
            "19-20. Demo checklist and commands": "14",
        },
    )

    set_by_prefix(
        doc,
        "The application runs immediately with the packaged model",
        "The private/local application runs with the packaged Joblib and its hash-matched private "
        "confidence_calibration.json. Runtime never retrains. Raw datasets are needed only for "
        "authorised offline reproduction and must remain Git-ignored/private.",
    )
    set_by_prefix(
        doc,
        "Expected: Python reports 3.11.x.",
        "Expected: Python reports 3.12.x. Do not use a Microsoft Store alias if VS Code points to a different interpreter.",
    )
    set_by_prefix(
        doc,
        "python3.11 -m venv .venv",
        "python3.12 -m venv .venv\nsource .venv/bin/activate",
        style="Code Block",
    )
    old_training_heading = find_paragraph(doc, exact="13. Train the fake-news model", required=False)
    if old_training_heading is not None:
        set_paragraph(old_training_heading, "13. Optional offline training and controlled benchmark reproduction", style="Heading 1")
    set_by_prefix(
        doc,
        "python training/train_fake_news_models.py",
        "# Offline development only; never called by Streamlit\n"
        "python training/train_fake_news_models.py\n"
        "# Controlled evidence run; CSVs stay outside the public repository\n"
        "python training/benchmark_models.py --raw-dir /absolute/private/isot-directory",
        style="Code Block",
    )
    set_by_prefix(
        doc,
        "The script cleans and deduplicates data",
        "The original training command remains an offline research tool. The controlled benchmark "
        "first verifies the accepted production-model SHA-256 and 4,800-row reproduction evidence, "
        "then screens near-duplicates, preserves 19,200 training rows, uses 1,199 calibration and "
        "1,200 policy rows, and reports once on a 2,399-row final test. It does not overwrite or "
        "retrain the packaged production model.",
    )
    old_retraining_heading = find_paragraph(doc, exact="15. Retrain with new data", required=False)
    if old_retraining_heading is not None:
        set_paragraph(old_retraining_heading, "15. Optional offline retraining with newly licensed data", style="Heading 1")
    test_status = (
        "All 56 packaged checks passed in the final tested Python 3.12 environment."
        if audit_passed
        else "The release defines 56 packaged checks; the final audit writes the authoritative pass/fail record."
    )
    set_by_prefix(
        doc,
        "Verification evidence:",
        f"Verification evidence: {test_status} The 29 established checks remain included; 27 "
        "additional checks cover calibration, review, analytics, drift readiness, privacy isolation "
        "and release policy. Browser audits cover six sections, same-tab routing, direct routes, "
        "refresh, back/forward, keyboard use, analysis/review/download flows and five viewport widths.",
    )
    old_deployment_heading = find_paragraph(doc, exact="17. Public deployment", required=False)
    if old_deployment_heading is not None:
        set_paragraph(old_deployment_heading, "17. Target public deployment (currently blocked)", style="Heading 1")
    set_by_prefix(
        doc,
        "GitHub is the canonical public source",
        f"After publication, {INTENDED_REPOSITORY} is intended to be canonical. The target Streamlit "
        "Community Cloud configuration is branch main, entrypoint app.py and root requirements.txt; "
        "Vercel uses web/ and exposes only NEXT_PUBLIC_STREAMLIT_APP_URL. The session history default "
        "is temporary visitor-isolated SQLite. No live deployment is claimed, and functional hosting "
        "must not proceed until the derived-model redistribution basis is documented.",
    )
    set_by_prefix(
        doc,
        "Publication gate:",
        "Publication gate: do not publish secrets, visitor data, generated databases, raw datasets, "
        "the private Joblib or confidence_calibration.json. The public staging package intentionally "
        "cannot run classification; do not retrain, mock or replace the model to bypass this gate. "
        "Create the empty public repository and update live links only after the owner completes the "
        "documented repository and rights steps.",
    )
    old_probs = find_paragraph(doc, exact="Probabilities sum to approximately 100%.", required=False)
    if old_probs is not None:
        set_paragraph(old_probs, "Calibrated probabilities sum to approximately 100%, and the three exact outcome states are reachable.", style="List Bullet")
    set_by_prefix(
        doc,
        "python -m pytest -q passes",
        "python -m pytest -q passes all 56 checks, including the 29 established checks.",
        style="List Bullet",
    )

    setup_picture_map = {
        "Figure S.1.": (SCREENSHOTS / "01_home.png", "Current NewsLens AI News Desk"),
        "Figure S.2.": (SCREENSHOTS / "02_analysis_input.png", "Current populated article input workflow"),
        "Figure S.3.": (SCREENSHOTS / "03_summary_and_risk_results.png", "Current summary and calibrated risk result"),
        "Figure S.4.": (SCREENSHOTS / "04_explainability_and_downloads.png", "Current explainability and export workflow"),
    }
    for caption, (path, alt_text) in setup_picture_map.items():
        replace_picture_by_caption(doc, caption, path, alt_text=alt_text)

    common_document_finish(doc, output)
    return output


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--audit-passed",
        action="store_true",
        help="Record a successful comprehensive audit only after that audit has actually run.",
    )
    args = parser.parse_args()
    outputs = [
        reconcile_report(audit_passed=args.audit_passed),
        reconcile_developer_guide(audit_passed=args.audit_passed),
        reconcile_concepts_guide(audit_passed=args.audit_passed),
        reconcile_setup_guide(audit_passed=args.audit_passed),
    ]
    for output in outputs:
        print(output.relative_to(ROOT))


if __name__ == "__main__":
    main()
