"""Build the three professionally formatted Word deliverables.

Design resolution:
- Academic report: narrative_proposal preset + editorial_cover first page.
- Setup guide and developer guide: compact_reference_guide preset + editorial_cover.
- US Letter, 1-inch margins, Calibri, fixed 9360-DXA tables, quiet headers,
  page fields, real Word headings/lists, and report-grade figure captions.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Iterable, Sequence

import pandas as pd
from PIL import Image
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIGURES = ROOT / "reports" / "figures"
DIAGRAMS = ROOT / "reports" / "diagrams"
SCREENSHOTS = ROOT / "reports" / "screenshots"
RESULTS = ROOT / "reports" / "results"

NAVY = "2A241F"
BLUE = "6D5947"
CYAN = "9B8066"
TEAL = "496454"
VIOLET = "806A5A"
AMBER = "8A693D"
RED = "813F39"
INK = "1A1917"
MUTED = "6F685F"
PALE = "F3F0E8"
LIGHT = "EAE4D8"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
PROJECT_AUTHOR = "Deven Sachin Gaikwad"
COPYRIGHT_NOTICE = "© 2026 Deven Sachin Gaikwad. All Rights Reserved."


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, value: object, *, bold: bool = False, color: str = INK, size: float = 9.0, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(value))
    set_run(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, *, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_in: Sequence[float], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    widths_dxa = [int(round(width * 1440)) for width in widths_in]
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", 80), ("left", 120), ("bottom", 80), ("right", 120)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def add_field(paragraph, instruction: str, fallback: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction_node, separate, text, end):
        run._r.append(element)


def set_paragraph_border(
    paragraph,
    edge: str,
    *,
    color: str = RED,
    size: int = 10,
    space: int = 4,
) -> None:
    """Add a restrained Word paragraph rule used by academic headers/footers."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = borders.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        borders.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def configure_document(title: str, preset: str, *, academic_report: bool = False) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8 if preset == "narrative_proposal" else 6)
    normal.paragraph_format.line_spacing = 1.333 if preset == "narrative_proposal" else 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if preset == "narrative_proposal" else WD_ALIGN_PARAGRAPH.LEFT

    heading_tokens = {
        "Heading 1": (16, BLUE, 18 if preset == "narrative_proposal" else 18, 10),
        "Heading 2": (13, BLUE, 12 if preset == "narrative_proposal" else 14, 7),
        "Heading 3": (12, "1F4D78", 8 if preset == "narrative_proposal" else 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375 if preset == "compact_reference_guide" else 0.5)
        style.paragraph_format.first_line_indent = Inches(-0.188 if preset == "compact_reference_guide" else -0.25)
        style.paragraph_format.space_after = Pt(4 if preset == "compact_reference_guide" else 6)
        style.paragraph_format.line_spacing = 1.25 if preset == "compact_reference_guide" else 1.208

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.5)
    code.font.color.rgb = rgb(NAVY)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.0

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    paragraph.paragraph_format.space_after = Pt(0)
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_p.paragraph_format.space_before = Pt(0)

    if academic_report:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("NEWSLENS AI · MINI PROJECT REPORT · DEVEN SACHIN GAIKWAD")
        set_run(run, size=8.5, color=NAVY, bold=True)
        set_paragraph_border(paragraph, "bottom", color=RED, size=10, space=4)

        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run(f"NewsLens AI · Designed and developed by {PROJECT_AUTHOR}")
        set_run(run, size=8.2, color=MUTED)
        run = footer_p.add_run(f"\n{COPYRIGHT_NOTICE} · Page ")
        set_run(run, size=8.2, color=MUTED)
        add_field(footer_p, "PAGE", "1")
        run = footer_p.add_run(" of ")
        set_run(run, size=8.2, color=MUTED)
        add_field(footer_p, "NUMPAGES", "1")
        set_paragraph_border(footer_p, "top", color=RED, size=10, space=4)
    else:
        run = paragraph.add_run(title)
        set_run(run, size=8.5, color=MUTED, bold=True)
        paragraph.add_run("\t")
        right = paragraph.add_run("DEVEN SACHIN GAIKWAD · PROPRIETARY DOCUMENTATION")
        set_run(right, size=8.5, color=TEAL)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run(f"NewsLens AI · Designed and developed by {PROJECT_AUTHOR}")
        set_run(run, size=8.0, color=MUTED)
        run = footer_p.add_run(f"\n{COPYRIGHT_NOTICE} · Page ")
        set_run(run, size=8.5, color=MUTED)
        add_field(footer_p, "PAGE", "1")
        run = footer_p.add_run(" of ")
        set_run(run, size=8.5, color=MUTED)
        add_field(footer_p, "NUMPAGES", "1")

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    doc.core_properties.title = title
    doc.core_properties.subject = "NewsLens AI academic software documentation and responsible-use record"
    doc.core_properties.author = PROJECT_AUTHOR
    doc.core_properties.last_modified_by = PROJECT_AUTHOR
    doc.core_properties.comments = COPYRIGHT_NOTICE
    doc.core_properties.keywords = "NLP, summarization, fake news, Streamlit, explainable AI"
    return doc


def add_cover(doc: Document, title: str, subtitle: str, document_type: str) -> None:
    institute = doc.add_paragraph()
    institute.alignment = WD_ALIGN_PARAGRAPH.CENTER
    institute.paragraph_format.space_after = Pt(8)
    run = institute.add_run("NEWSLENS AI")
    set_run(run, size=12, color=NAVY, bold=True)

    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.paragraph_format.left_indent = Inches(2.4)
    logo.paragraph_format.right_indent = Inches(2.4)
    logo.paragraph_format.space_before = Pt(4)
    logo.paragraph_format.space_after = Pt(4)
    for edge in ("top", "left", "bottom", "right"):
        set_paragraph_border(logo, edge, color="9AA9BC", size=8, space=6)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT)
    logo._p.get_or_add_pPr().append(shading)
    run = logo.add_run("PUBLIC DOCUMENTATION")
    set_run(run, size=8.2, color=MUTED, bold=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run(document_type.upper())
    set_run(run, size=11, color=TEAL, bold=True)
    kicker.paragraph_format.space_after = Pt(10)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(10)
    run = title_p.add_run(title)
    set_run(run, size=24, color=NAVY, bold=True)
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(16)
    run = sub_p.add_run(subtitle)
    set_run(run, size=11.5, color=BLUE, italic=True)
    meta = [
        ("Product", "NewsLens AI"),
        ("Document status", "Sanitized publication-staging documentation"),
        ("Author and developer", PROJECT_AUTHOR),
        ("Copyright", COPYRIGHT_NOTICE),
        ("Runtime", "Streamlit · Python 3.12 · app.py"),
        ("Model boundary", "Linguistic credibility-risk estimate; not a verified fact-check"),
        ("Repository", "Canonical public URL added only after owner approval"),
        ("Document updated", "16 August 2026"),
    ]
    add_table(doc, ["Field", "Value"], meta, [2.25, 4.25], font_size=8.4, header_fill=PALE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("CODE · METHODS · EVIDENCE · RESPONSIBLE USE")
    set_run(run, size=9.5, color=MUTED, bold=True)
    doc.add_page_break()


def add_title_page(
    doc: Document,
    heading: str,
    paragraphs: Iterable[str],
    signatures: Sequence[str] = (),
    *,
    context_lines: Sequence[str] = (),
    page_break: bool = True,
) -> None:
    for index, text in enumerate(context_lines):
        context = doc.add_paragraph()
        context.alignment = WD_ALIGN_PARAGRAPH.CENTER
        context.paragraph_format.space_after = Pt(2 if index < len(context_lines) - 1 else 12)
        run = context.add_run(text)
        set_run(run, size=10 if index else 11, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18 if context_lines else 30)
    p.paragraph_format.space_after = Pt(26)
    run = p.add_run(heading.upper())
    set_run(run, size=18, color=NAVY, bold=True)
    for text in paragraphs:
        body = doc.add_paragraph(text)
        body.paragraph_format.space_after = Pt(12)
        body.paragraph_format.line_spacing = 1.35
        body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if signatures:
        doc.add_paragraph()
        rows = [[value, "____________________________"] for value in signatures]
        add_table(doc, ["Role", "Signature / Date"], rows, [2.5, 4.0], font_size=9.5, header_fill=LIGHT)
    if page_break:
        doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run(lead, bold=True, color=NAVY)
        run = paragraph.add_run(text[len(bold_lead):])
        set_run(run)
    else:
        run = paragraph.add_run(text)
        set_run(run)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run(run)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        set_run(run)


def add_code(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph(style="Code Block")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "EEF3F8")
    paragraph._p.get_or_add_pPr().append(shading)
    run = paragraph.add_run(code)
    set_run(run, name="Consolas", size=8.5, color=NAVY)


def add_callout(doc: Document, label: str, text: str, *, fill: str = "EEF5FF", accent: str = BLUE) -> None:
    # A callout is prose, not row/column data. Paragraph furniture keeps the
    # document semantically cleaner for assistive technology than a layout table.
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.08
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    set_paragraph_border(p, "left", color=accent, size=20, space=8)
    set_paragraph_border(p, "top", color="D6CEC2", size=4, space=5)
    set_paragraph_border(p, "bottom", color="D6CEC2", size=4, space=5)
    set_paragraph_border(p, "right", color="D6CEC2", size=4, space=5)
    r = p.add_run(f"{label}: ")
    set_run(r, size=9.5, color=accent, bold=True)
    r = p.add_run(text)
    set_run(r, size=9.5, color=INK)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[object]], widths_in: Sequence[float], *, font_size: float = 8.5, header_fill: str = PALE) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_in)
    set_repeat_table_header(table.rows[0])
    for index, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], value, bold=True, color=NAVY, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[index], header_fill)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if index == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[index], value, size=font_size, align=align)
            if row_index % 2 == 1:
                set_cell_shading(cells[index], "FAFCFF")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.keep_with_next = False
    run = p.add_run(text)
    set_run(run, size=9, color=MUTED, italic=True)


def add_figure(
    doc: Document,
    path: Path,
    caption: str,
    explanation: str,
    *,
    width: float = 6.2,
    max_height: float = 7.0,
    page_break: bool = True,
) -> None:
    p = doc.add_paragraph()
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    with Image.open(path) as source_image:
        pixel_width, pixel_height = source_image.size
    scaled_height = width * pixel_height / pixel_width
    if scaled_height > max_height:
        inline = run.add_picture(str(path), height=Inches(max_height))
    else:
        inline = run.add_picture(str(path), width=Inches(width))
    try:
        inline._inline.docPr.set("descr", caption)
        inline._inline.docPr.set("title", caption)
    except Exception:
        pass
    add_caption(doc, caption)
    add_body(doc, explanation)


def add_toc(doc: Document, entries: Sequence[Sequence[object]]) -> None:
    add_heading(doc, "Table of Contents", 1)
    add_table(doc, ["Section", "Rendered page"], entries, [5.7, 0.8], font_size=9.3, header_fill="E8EEE8")
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    run = note.add_run("Page references match the packaged render; headings remain Word-native for navigation and future TOC updates.")
    set_run(run, size=8.5, color=MUTED, italic=True)
    doc.add_page_break()


def add_chapter(doc: Document, number: int, title: str):
    paragraph = doc.add_heading(f"Chapter {number}: {title}", level=1)
    paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def build_report() -> Path:
    metrics = json.loads((RESULTS / "model_metrics.json").read_text(encoding="utf-8"))
    profile = json.loads((RESULTS / "dataset_profile.json").read_text(encoding="utf-8"))
    summary_metrics = json.loads((RESULTS / "summarization_metrics.json").read_text(encoding="utf-8"))
    papers = json.loads((DOCS / "research_papers.json").read_text(encoding="utf-8"))
    comparison = pd.read_csv(RESULTS / "model_comparison.csv")
    samples = pd.read_csv(RESULTS / "packaged_sample_analyses.csv")
    tests = pd.read_csv(RESULTS / "test_cases.csv")

    doc = configure_document(
        "NewsLens AI",
        "narrative_proposal",
        academic_report=True,
    )
    add_cover(
        doc,
        "NewsLens AI",
        "AI-Based News Article Summarization and Fake News Detection System",
        "Mini Project Report",
    )
    add_title_page(
        doc,
        "Publication-Staging Note",
        [
            "This sanitized public report documents the NewsLens AI research application, its implementation, measured evidence, limitations and deployment architecture. Personal academic identifiers and signature fields are intentionally excluded.",
            f"NewsLens AI was designed and developed by {PROJECT_AUTHOR}. {COPYRIGHT_NOTICE} GitHub will become the canonical publication record only after owner approval. The packaged model must not be redistributed publicly until documentary permission or an explicit applicable licence is confirmed.",
        ],
        context_lines=[
            "NewsLens AI · Publication-staging documentation",
            "Document updated: 16 August 2026",
            "Runtime: Streamlit · Python 3.12 · app.py",
        ],
    )
    add_title_page(
        doc,
        "Responsible-Use Declaration",
        [
            "NewsLens AI estimates linguistic credibility risk from patterns learned in its training data. It does not retrieve evidence, determine objective truth, infer author intent or replace professional fact-checking.",
            "Sources, datasets, papers and tutorial references are acknowledged. Measured model and summarization results come from committed artifacts; limitations, uncertainty and the need for independent verification remain part of the product boundary.",
        ],
    )
    add_title_page(
        doc,
        "Acknowledgement",
        [
            f"NewsLens AI's original project components were authored and developed by {PROJECT_AUTHOR}. The project acknowledges the University of Victoria ISOT Research Group, the Edinburgh NLP Group, the authors represented in the research matrix, and the open-source Python community.",
            "The two supplied videos were treated as conceptual tutorial references only. Their baseline ideas were re-engineered into an original layered application with leakage controls, model persistence, local explanations, SQLite history, testing and responsible-AI boundaries.",
        ],
    )
    add_title_page(
        doc,
        "Abstract",
        [
            "The rapid volume of online news creates two related but scientifically distinct needs: concise reading and credibility assessment. This project implements NewsLens AI, a local-first Streamlit application that ingests direct text, public URLs, TXT files and text-based PDFs; cleans and characterises the article; generates a selectable summary; and independently estimates credibility risk from the original cleaned article.",
            f"The lightweight summarizer uses TF-IDF centroid sentence ranking with lead-position and information-density bonuses. A CPU-compatible DistilBART mode is optional. For classification, three leakage-controlled scikit-learn pipelines were compared on a balanced {profile['training_sample_rows']:,}-row ISOT sample after removing {profile['duplicates_removed']:,} exact duplicates, screening near-duplicates and neutralising source markers. Logistic Regression was retained because Linear SVC's validation-policy macro-F1 advantage was only 0.002500, below the predefined 0.01 tolerance, while the verified production artifact provides direct coefficient explanations. On the untouched {metrics['test_samples']:,}-row final test it achieved accuracy {metrics['accuracy']:.6f}, macro-F1 {metrics['macro_f1']:.6f}, ROC-AUC {metrics['roc_auc']:.6f} and PR-AUC {metrics['pr_auc']:.6f}.",
            f"The extractive summarizer was evaluated on {summary_metrics['sample_size']} fixed-seed XSum test articles, producing ROUGE-1/2/L F1 of {summary_metrics['rouge1_f1']:.6f}/{summary_metrics['rouge2_f1']:.6f}/{summary_metrics['rougeL_f1']:.6f}, {summary_metrics['mean_compression_ratio_pct']:.4f}% mean compression and {summary_metrics['mean_latency_ms']:.3f} ms mean latency. Low extractive ROUGE on highly abstractive XSum references is reported transparently. The application adds local feature contributions, confidence bands, session-isolated SQLite history, duplicate hashing, JSON/PDF/CSV export, 15 responsive interface screenshots and 56 packaged checks.",
        ],
        page_break=False,
    )
    add_callout(doc, "Keywords", "automatic summarization; credibility-risk classification; TF-IDF; Logistic Regression; DistilBART; explainable AI; Streamlit; SQLite; responsible AI", fill="E8EEE8", accent=TEAL)
    add_callout(doc, "Interface design", "The application and every frontend figure use the same warm ivory, beige, brown and charcoal editorial newsroom design.", fill="F3EBDD", accent=AMBER)
    doc.add_page_break()
    add_toc(doc, [
        ("Chapter 1: Introduction", 10),
        ("Chapter 2: Literature Survey", 12),
        ("Chapter 3: Requirements and Feasibility", 16),
        ("Chapter 4: System Analysis and Design", 18),
        ("Chapter 5: Dataset and Data Preprocessing", 29),
        ("Chapter 6: Methodology and Algorithms", 40),
        ("Chapter 7: Implementation", 43),
        ("Chapter 8: Testing and Results", 45),
        ("Chapter 9: Graphical User Interface", 53),
        ("Chapter 10: Project Management", 65),
        ("Chapter 11: Limitations, Ethics and Future Scope", 66),
        ("Chapter 12: Conclusion", 67),
        ("References", 68),
        ("Appendices A-E", 70),
    ])

    add_heading(doc, "List of Figures", 1)
    figures_list = [
        "Figure 4.1 Overall system architecture", "Figure 4.2 End-to-end data flow", "Figure 4.3 DFD Level 0", "Figure 4.4 DFD Level 1", "Figure 4.5 Use-case diagram", "Figure 4.6 Activity diagram", "Figure 4.7 Sequence diagram", "Figure 4.8 Component/module diagram", "Figure 4.9 SQLite ER diagram", "Figure 4.10 Deployment diagram", "Figure 5.1 Machine-learning training pipeline", "Figure 5.2 Class distribution", "Figure 5.3 Article-length distribution", "Figure 5.4 Title-length distribution", "Figure 5.5 Average sentence length", "Figure 5.6 Missing-value heatmap", "Figure 5.7 Duplicate-record analysis", "Figure 5.8 N-gram comparison", "Figure 5.9 Numerical-feature correlations", "Figure 5.10 Subject leakage risk", "Figure 6.1 Combined inference pipeline", "Figure 8.1 Model comparison", "Figure 8.2 Confusion matrix", "Figure 8.3 ROC and precision-recall curves", "Figure 8.4 Global feature coefficients", "Figure 9.1 Streamlit navigation", "Figures 9.2-9.16 Current desktop and mobile application screenshots",
    ]
    add_bullets(doc, figures_list)
    doc.add_page_break()
    add_heading(doc, "List of Tables", 1)
    add_bullets(doc, ["Table 2.1 Literature survey comparison", "Table 3.1 Functional requirements", "Table 3.2 Software and hardware", "Table 5.1 Dataset comparison", "Table 5.2 Cleaning audit", "Table 8.1 Model comparison", "Table 8.2 Champion metrics", "Table 8.3 Test cases", "Table 10.1 Project risks", "Table A.1 Dataset schema", "Table B.1 Function reference"])
    doc.add_page_break()
    add_heading(doc, "List of Abbreviations", 1)
    add_table(doc, ["Abbreviation", "Meaning"], [
        ("AI", "Artificial Intelligence"), ("NLP", "Natural Language Processing"), ("TF-IDF", "Term Frequency-Inverse Document Frequency"), ("LR", "Logistic Regression"), ("SVM", "Support Vector Machine"), ("XAI", "Explainable Artificial Intelligence"), ("ROUGE", "Recall-Oriented Understudy for Gisting Evaluation"), ("DFD", "Data-Flow Diagram"), ("CRUD", "Create, Read, Update and Delete"), ("SSRF", "Server-Side Request Forgery"), ("CPU", "Central Processing Unit"), ("UI", "User Interface"),
    ], [1.4, 5.1], font_size=9.5)

    add_chapter(doc, 1, "Introduction")
    sections = {
        "1.1 Background": "Online news readers face an attention problem and an evidence problem. Summarization can reduce reading time, but it does not establish truth. Text classification can detect learned linguistic patterns, but it cannot verify claims without external evidence. Treating these tasks as interchangeable is scientifically unsafe.",
        "1.2 Problem statement": "Tutorial implementations commonly demonstrate a single model in a notebook or a minimal interface. They seldom combine robust ingestion, long-text handling, independent AI branches, probability communication, local explanations, persistent history, reproducible evaluation and responsible-use warnings in one student-maintainable system.",
        "1.3 Motivation": "The project provides an accessible engineering baseline for students while exposing rather than hiding uncertainty. It demonstrates how a model moves from dataset audit to saved Pipeline, how leakage can inflate metrics, and how interface language influences ethical use.",
        "1.4 Research gap": "Existing work often optimises one benchmark component. A practical gap remains for an offline, CPU-friendly news intelligence application that summarises long articles and presents an interpretable credibility-risk estimate without misrepresenting it as fact verification.",
        "1.5 Aim": "To design, implement and verify an integrated NLP system for article summarization and responsible credibility-risk classification using modular Python, Streamlit and locally persisted artifacts.",
        "1.6 Scope": "The implemented scope is English news-style text, public web extraction where technically accessible, TXT/text-PDF ingestion, extractive and optional abstractive summarization, binary credibility-risk classification, local linear explanations, SQLite history, plots and downloads. Evidence retrieval, OCR, multilingual modelling and production moderation are explicitly out of scope.",
        "1.7 Innovation and expected outcomes": "The original contribution is the engineering integration: independent classification of original text, source-marker mitigation, a preloaded saved model, confidence abstention language, feature contributions, session-isolated duplicate-aware history, measured accountability views, complete tests, and documentation suitable for academic examination.",
    }
    for heading, text in sections.items():
        add_heading(doc, heading, 2); add_body(doc, text)
    add_heading(doc, "1.8 Objectives", 2)
    add_bullets(doc, ["Accept four input forms and fail gracefully.", "Generate short, medium and detailed summaries.", "Compare at least three classifiers using macro-F1 and more than accuracy.", "Persist the fitted vectorizer-classifier Pipeline outside Streamlit.", "Explain observed local term contributions and always show a verification disclaimer.", "Store compact structured history without retaining uploaded files.", "Calculate real classification, ROUGE, latency and test evidence.", "Deliver code, model, cards, notebooks, diagrams, screenshots, spreadsheet and Word guides in one ZIP."])
    add_heading(doc, "1.9 Applications", 2)
    add_bullets(doc, ["Academic demonstrations of classical NLP and transformers.", "News-reading assistance where users still verify claims independently.", "Model-risk and dataset-bias teaching through visible examples.", "A baseline for later evidence retrieval, multilingual and multimodal research."])

    add_chapter(doc, 2, "Literature Survey")
    add_heading(doc, "2.1 Review methodology", 2)
    add_body(doc, "The review searched primary IEEE Xplore and ACL records for summarization, transformer generation, fake-news detection, explainability, dataset bias and generalisation. Ten peer-reviewed papers were selected; six have IEEE Xplore records. DOI, author, year, venue and access status were checked. Public preprints or mirrors were used when available. Publisher-formatted copies requiring institutional access are not represented as fully reviewed.")
    add_heading(doc, "2.2 Paper-wise review", 2)
    for paper in papers:
        add_heading(doc, f"[{paper['Sr. No.']}] {paper['Paper title']} ({paper['Year']})", 3)
        add_body(doc, f"{paper['Authors']}. {paper['Method/model']}. Dataset: {paper['Dataset']}. {paper['Major findings']} Limitation: {paper['Limitations']} Relevance: {paper['Relevance to this project']}")
    add_heading(doc, "2.3 Literature-survey comparison", 2)
    literature_rows = [[p["Sr. No."], p["Year"], "IEEE" if p["IEEE Xplore link"] != "Not applicable" else "ACL", p["Paper title"], p["Research gap identified"]] for p in papers]
    add_table(doc, ["No.", "Year", "Venue", "Paper", "Gap carried forward"], literature_rows, [0.42, 0.55, 0.58, 2.25, 2.70], font_size=7.4)
    add_callout(doc, "Final gap", "Benchmark performance is not equivalent to reasoning or fact verification. This project therefore combines lightweight engineering with visible model limits and leaves evidence-based claim verification as future work.", fill="F3EBDD", accent=AMBER)

    add_chapter(doc, 3, "Requirements and Feasibility")
    add_heading(doc, "3.1 Functional requirements", 2)
    req_rows = [
        ("FR-01", "Accept direct text, public URL, TXT and text-based PDF"), ("FR-02", "Extract text/metadata with fallbacks and validation"), ("FR-03", "Generate selectable extractive/optional abstractive summary"), ("FR-04", "Classify original cleaned article and report probabilities"), ("FR-05", "Show confidence band, local terms, model artifact ID and disclaimer"), ("FR-06", "Store, search, export and delete local analysis history"), ("FR-07", "Export one analysis as JSON or PDF"), ("FR-08", "Expose real model, EDA, ROUGE and error-analysis artifacts"),
    ]
    add_table(doc, ["ID", "Requirement"], req_rows, [1.0, 5.5], font_size=9)
    add_heading(doc, "3.2 Non-functional requirements", 2)
    add_bullets(doc, ["CPU-first lightweight mode; the core workflow requires no paid API, while hosting, network, compute and third-party terms may still carry costs.", "Portable relative paths and Windows/VS Code compatibility.", "Deterministic seed and no automatic retraining at app startup.", "Clear accessibility contrast and color-independent text labels.", "User-correctable errors rather than stack-trace crashes.", "Local privacy: no permanent uploaded-file storage.", "Reproducible metrics and primary-source citations."])
    add_heading(doc, "3.3 Software and hardware", 2)
    add_table(doc, ["Category", "Minimum", "Recommended"], [
        ("Operating system", "Windows 10 / modern Linux / macOS", "Windows 11 or Ubuntu LTS"), ("Python", "3.11", "3.11 64-bit"), ("RAM", "4 GB for extractive mode", "8-16 GB for DistilBART"), ("CPU", "Dual-core", "Modern 4+ core CPU"), ("Storage", "1 GB lightweight", "5+ GB with caches/datasets"), ("Software", "VS Code, browser, pip", "Git and Jupyter optional"),
    ], [1.35, 2.55, 2.60], font_size=8.7)
    add_heading(doc, "3.4 Feasibility", 2)
    add_body(doc, "Technical feasibility is demonstrated by 56 packaged checks, including the 29 established checks, Chromium release audits, a packaged 0.8 MB classical model and sub-millisecond classifier inference. Operational feasibility comes from a guided multipage interface, synthetic samples and verified cross-visitor archive isolation. The core workflow requires no paid API, but hosting, network, compute and third-party terms may still carry costs. Ethical feasibility depends on maintaining uncertainty, disclaimers, error visibility and human verification.")

    add_chapter(doc, 4, "System Analysis and Design")
    add_heading(doc, "4.1 Existing versus proposed system", 2)
    add_table(doc, ["Aspect", "Tutorial-scale baseline", "Proposed NewsLens AI"], [
        ("Structure", "Single script/notebook", "Layered modules and six-page experience"), ("Inputs", "Pasted text or one URL", "Text, URL, TXT, text-PDF"), ("Models", "One summarizer or classifier", "Independent summary and risk branches"), ("Persistence", "Often none", "Saved Pipeline and SQLite history"), ("Evaluation", "Example output", "Held-out, CV, ROC/PR, ROUGE, latency, errors"), ("Responsibility", "Binary claim", "Risk labels, uncertainty, XAI and disclaimer"),
    ], [1.25, 2.25, 3.0], font_size=8.4)
    diagram_info = [
        ("01_overall_system_architecture.png", "Figure 4.1. Overall system architecture.", "Six layers isolate UI, ingestion, NLP, AI, persistence and evaluation so a student can test or replace one concern without rewriting the application."),
        ("02_end_to_end_data_flow.png", "Figure 4.2. End-to-end data-flow diagram.", "The original cleaned article fans out to the summary and classifier. Only their result objects merge in the final editorial report."),
        ("03_dfd_level_0.png", "Figure 4.3. DFD Level 0 context diagram.", "The user interacts with one system boundary; a public website, saved model and local SQLite store are explicit external data sources."),
        ("04_dfd_level_1.png", "Figure 4.4. DFD Level 1.", "Process 0 is decomposed into acquisition, cleaning, independent AI tasks, composition and persistence/export."),
        ("05_use_case_diagram.png", "Figure 4.5. Use-case diagram.", "The user can submit, configure, inspect, revisit and export. Network actors provide content/models but do not determine truth."),
        ("06_activity_diagram.png", "Figure 4.6. Activity diagram.", "Validation failure returns to input. Valid content reaches parallel AI activities, then a duplicate-aware storage decision."),
        ("07_sequence_diagram.png", "Figure 4.7. Sequence diagram.", "The sequence makes caching, independent calls and SQLite insertion order explicit."),
        ("08_component_module_diagram.png", "Figure 4.8. Component/module diagram.", "Python modules depend inward on reusable business logic; offline training writes identified artifacts consumed at runtime."),
        ("11_sqlite_er_diagram.png", "Figure 4.9. SQLite ER diagram.", "The single analyses entity uses a unique SHA-256 article hash and probability/time CHECK constraints."),
        ("13_deployment_diagram.png", "Figure 4.10. Public deployment architecture.", "GitHub is canonical; Vercel hosts the presentation shell; Streamlit Community Cloud runs app.py with temporary visitor-isolated SQLite."),
    ]
    for filename, caption, explanation in diagram_info:
        add_figure(doc, DIAGRAMS / filename, caption, explanation)

    add_chapter(doc, 5, "Dataset and Data Preprocessing")
    add_heading(doc, "5.1 Dataset comparison and selection", 2)
    add_table(doc, ["Dataset", "Size/content", "Strength", "Constraint", "Decision"], [
        ("ISOT", "44,898 full articles", "Binary, accessible, full text", "Severe outlet/topic leakage risk", "Selected with controls"), ("LIAR", "12.8K short claims", "Manual PolitiFact labels", "Six labels; not full articles", "Literature comparator"), ("FakeNewsNet", "News + social context", "Publisher and interactions", "Collection/API/reproducibility friction", "Future extension"), ("XSum", "226,711 BBC article-summary pairs", "Human one-sentence references", "Highly abstractive; single broadcaster", "Summary evaluation"),
    ], [0.75, 1.25, 1.45, 1.75, 1.30], font_size=7.8)
    add_heading(doc, "5.2 Schema and cleaning", 2)
    add_body(doc, "ISOT provides title, text, subject and date. The source filename defines the binary label. Title and body are joined for modelling; subject, date and source identity are excluded. Missing values are normalised, short rows removed, explicit URLs/control bytes cleaned, Reuters/byline markers neutralised, and exact hashes deduplicated before sampling or splitting. Stemming and aggressive lemmatisation are intentionally avoided because style and negation can be predictive and because display text must remain readable.")
    add_table(doc, ["Audit item", "Measured rows"], [
        ("Raw rows", f"{profile['raw_rows']:,}"), ("Eligible unique corpus", f"{profile['clean_rows']:,}"), ("Exact duplicates removed", f"{profile['duplicates_removed']:,}"), ("Short/empty removed", f"{profile['short_or_empty_rows_removed']:,}"), ("Balanced working sample", f"{profile['training_sample_rows']:,}"), ("Reliable / misleading", f"{profile['reliable_rows']:,} / {profile['misleading_rows']:,}"),
    ], [3.6, 2.9], font_size=9.2)
    add_figure(doc, DIAGRAMS / "09_ml_training_pipeline.png", "Figure 5.1. Machine-learning training pipeline.", "Exact deduplication precedes the fixed split. GridSearchCV fits TF-IDF within training folds; the held-out test set is used once after selection.")
    eda_figures = [
        ("class_distribution.png", "Figure 5.2. Class distribution after preprocessing.", "The 24,000-row working set is balanced by design, preventing a majority-class accuracy shortcut."),
        ("word_count_distribution.png", "Figure 5.3. Article word-count distribution by class.", "Length profiles differ, indicating collection-style effects that may persist after text cleaning."),
        ("title_length_distribution.png", "Figure 5.4. Title-length distribution.", "Misleading-labelled titles are substantially longer on average; title length is disclosed as a possible shortcut."),
        ("average_sentence_length.png", "Figure 5.5. Average sentence length.", "Sentence length differs across labels, consistent with outlet and editing-style variation."),
        ("missing_values_heatmap.png", "Figure 5.6. Missing-value heatmap.", "The raw CSV fields contain no direct NaN values in this run; separate short/blank-content rules still remove unusable rows."),
        ("duplicate_record_analysis.png", "Figure 5.7. Duplicate-record and eligibility audit.", "Removing 5,713 duplicates before splitting prevents identical text from appearing in train and test."),
        ("ngram_frequency_comparison.png", "Figure 5.8. Frequent unigram, bigram and trigram comparison.", "Terms such as wire-service and date patterns reveal residual collection signatures; explicit markers are neutralised but not every proxy is removable."),
        ("numerical_feature_correlation.png", "Figure 5.9. Correlation of engineered numerical features.", "Title length correlates with label (0.59), a strong warning that high random-split accuracy may not generalise."),
        ("subject_distribution.png", "Figure 5.10. Original subject distribution.", "Subject categories are almost label-specific and therefore excluded completely from model features."),
    ]
    for filename, caption, explanation in eda_figures:
        add_figure(doc, FIGURES / filename, caption, explanation)
    add_heading(doc, "5.3 Split and leakage prevention", 2)
    add_bullets(doc, ["Remove exact duplicates before partitioning and screen near-duplicate groups.", "Use seed 42, stratification and group-aware validation/final-test partitioning.", "Quarantine contaminated holdout rows before calibration, policy or test use.", "Keep the 2,399-row final test untouched until one reporting pass.", "Place TF-IDF and classifier in one scikit-learn Pipeline.", "Exclude source, subject and date; neutralise source markers consistently.", "Report remaining outlet/topic/time artefacts and approximate-screen limits."])

    add_chapter(doc, 6, "Methodology and Algorithms")
    add_heading(doc, "6.1 Extractive sentence ranking", 2)
    add_body(doc, "The article is segmented with a deterministic punctuation-aware splitter. Sentences shorter than four words are discarded. A word 1-2 gram TF-IDF matrix is fitted on sentences, its mean vector becomes the document centroid, and cosine similarity ranks each sentence. A 0.12 exponentially decaying lead bonus reflects news inverted-pyramid structure; a capped 0.04 density bonus favours informative sentence length. Top sentences are restored to source order.")
    add_code(doc, "score_i = cosine(tfidf(sentence_i), centroid)\n        + 0.12 * exp(-position_i / scale)\n        + 0.04 * clip(words_i / 24, 0, 1)\nsummary = join(sorted(top_k_sentence_indices))")
    add_heading(doc, "6.2 Abstractive summarization", 2)
    add_body(doc, "The optional model is sshleifer/distilbart-cnn-6-6, selected as a smaller DistilBART CNN/DailyMail checkpoint. Streamlit caches the pipeline. Long input is divided at sentence boundaries by tokenizer length, adjacent chunks overlap by one sentence, and chunk summaries are hierarchically summarised if their concatenation still exceeds the context window. Dependency or download failure returns users to extractive mode rather than breaking the core application.")
    add_heading(doc, "6.3 Compression and ROUGE", 2)
    add_code(doc, "Compression Ratio (%) = (1 - Summary Word Count / Original Word Count) * 100")
    add_body(doc, "ROUGE-1 and ROUGE-2 measure unigram and bigram overlap; ROUGE-L uses longest common subsequence. Precision, recall and F1 are reported separately. ROUGE cannot by itself measure factuality, coherence or usefulness.")
    add_heading(doc, "6.4 TF-IDF and classifiers", 2)
    add_code(doc, "TF-IDF(t,d) = TF(t,d) * log(N / (DF(t) + 1))\np(misleading|x) = 1 / (1 + exp(-(w^T x + b)))")
    add_body(doc, "Multinomial Naive Bayes provides a probabilistic count baseline. Linear SVC maximises a separating margin and achieved the highest validation-policy score. Logistic Regression models class log-odds and supplies signed coefficients. After every candidate receives the same private Platt calibration, the predefined rule retains Logistic Regression when Linear SVC's validation-policy macro-F1 advantage is below 0.01. The untouched final test is excluded from that decision.")
    add_heading(doc, "6.5 Evaluation formulas", 2)
    add_code(doc, "Precision = TP / (TP + FP)\nRecall = TP / (TP + FN)\nF1 = 2 * Precision * Recall / (Precision + Recall)\nAccuracy = (TP + TN) / (TP + TN + FP + FN)")
    add_body(doc, "Accuracy alone can remain high when a majority class dominates. Macro-F1 weights each class equally; PR-AUC emphasises positive-class retrieval under imbalance; class-wise recall exposes missed misleading articles; timing and interpretability matter for a student laptop and responsible interface.")
    add_heading(doc, "6.6 Confidence and local explanation", 2)
    add_body(doc, "The model score is transformed by a model-hash-bound Platt calibration artifact. Confidence below the validation-selected 0.59 threshold, inadequate input or a supported-scope warning returns 'Editorial review required'. Otherwise the exact outcomes are 'Lower misleading-content risk indicated' or 'Higher misleading-content risk indicated'. For each observed TF-IDF term, local contribution equals its TF-IDF value multiplied by the Logistic Regression coefficient. These correlations are not evidence that a claim is true or false.")
    add_figure(doc, DIAGRAMS / "10_combined_inference_pipeline.png", "Figure 6.1. Combined inference pipeline.", "The no-feedback-loop annotation records the core scientific requirement: the summary never replaces the original classifier input.")

    add_chapter(doc, 7, "Implementation")
    add_heading(doc, "7.1 Development environment", 2)
    add_body(doc, "Python 3.12 is the tested public-release runtime. The verified build used Streamlit 1.59.2, pandas 2.2.3, NumPy 2.3.5, scikit-learn 1.8.0, Joblib 1.5.3, Plotly 6.9.0, Matplotlib 3.10.8, Seaborn 0.13.2, Requests 2.34.2, BeautifulSoup 4.15.0, Trafilatura 2.1.0, pypdf 6.15.0, ReportLab 4.4.9 and pytest 9.0.3. Transformers is pinned below major release 5 for the summarization pipeline API.")
    add_heading(doc, "7.2 Folder structure", 2)
    add_code(doc, "NewsLens-AI/\n  app.py · pages/ · ui/ · assets/ · src/\n  training/ · models/ · data/sample/\n  notebooks/ · tests/ · database/\n  reports/{figures,diagrams,screenshots,results}/\n  docs/ · scripts/ · web/ · requirements*.txt")
    add_heading(doc, "7.3 Module responsibilities", 2)
    module_rows = [
        ("article_extractor.py", "Public URL validation, SSRF blocking, Trafilatura and BeautifulSoup fallback"), ("file_parser.py", "10 MB TXT/text-PDF parsing and friendly failures"), ("text_preprocessor.py", "Shared source-marker-aware cleaning and deterministic sentence split"), ("extractive_summarizer.py", "TF-IDF centroid ranking and compression metrics"), ("abstractive_summarizer.py", "Cached pipeline, tokenizer-aware chunks, overlap and hierarchy"), ("fake_news_predictor.py", "Saved model loading, probabilities, bands and PredictionResult"), ("explainability.py", "Local observed term contributions and global coefficients"), ("database.py", "Schema, duplicate-aware insert, search, read, delete and clear"), ("session_history.py", "Fail-closed visitor isolation for the public session archive"), ("report_exporter.py", "In-memory JSON and printable PDF exports"), ("visualizations.py", "Warm publication-style Plotly figures"), ("app.py + ui/*", "Native same-tab routing, editorial tokens and reusable page/result components"),
    ]
    add_table(doc, ["Module", "Responsibility"], module_rows, [2.0, 4.5], font_size=8.3)
    add_heading(doc, "7.4 Representative implementation", 2)
    add_code(doc, "@st.cache_resource\ndef cached_model():\n    return load_model()\n\nsummary = summarize_extractive(cleaned, summary_length)\nprediction = predict_credibility(cleaned, cached_model())")
    add_body(doc, "The model is loaded through resource caching and never retrained at page start. ArticleData, SummaryResult and PredictionResult dataclasses make interfaces explicit. Portable Path objects replace user-specific paths. Recoverable errors have user-facing exception classes; an unexpected exception is contained behind a generic message and optional technical detail.")
    add_heading(doc, "7.5 URL and file safety", 2)
    add_bullets(doc, ["Only bounded HTTP(S) URLs with resolvable public hosts are accepted.", "Credential-bearing, localhost, private, loopback, link-local, reserved, numeric and obfuscated hosts are rejected.", "Every redirect is requested manually and independently validated; loops and chains beyond five hops are blocked.", "Every DNS answer must be globally routable and the observed peer address must match the validated set where the transport exposes it.", "Timeout, status, HTML content type, declared length and a streamed 5 MB response cap are enforced.", "Upload names reject traversal; encrypted PDFs, PDFs over 200 pages and extracted text over two million characters are rejected.", "CSV user fields are formula-neutralised and ReportLab Paragraph content is escaped.", "No uploaded file is stored in SQLite."])
    add_heading(doc, "7.6 Caching and performance", 2)
    add_body(doc, "The Joblib classifier is small and loaded once. Plot/data artifacts are read from stable files. DistilBART is downloaded only when requested and then cached by Hugging Face and Streamlit. Extractive mode remains the dependable offline path. SQLite uses indexed timestamp and label columns, and history queries cap results.")

    add_chapter(doc, 8, "Testing and Results")
    add_heading(doc, "8.1 Test strategy", 2)
    add_body(doc, "Verification is layered. All 56 packaged pytest checks passed in the Python 3.12 release environment, including the 29 established checks. Browser automation then exercised all six sections, direct routes, refresh, browser back/forward, keyboard activation, same-tab navigation, text analysis, summary, classification, confidence, explanation, JSON/PDF/CSV exports and visitor-isolated archives. The responsive audit covered 360, 390, 768, 1366 and 1920 pixel widths.")
    add_callout(doc, "Verification boundary", "Passing checks demonstrates the specified behaviour in the tested environment; it does not prove factual correctness or cross-domain generalisation. Re-run python -m pytest -q and the two optional browser audit scripts after dependency, hosting or interface changes.", fill="F3EBDD", accent=AMBER)
    comparison_rows = []
    for _, row in comparison.iterrows():
        comparison_rows.append([row["model"], f"{row['accuracy']:.6f}", f"{row['macro_f1']:.6f}", f"{row['roc_auc']:.6f}", f"{row['pr_auc']:.6f}", f"{row['calibrated_brier_score']:.6f} / {row['calibrated_ece']:.6f}"])
    add_table(doc, ["Model", "Accuracy", "Macro-F1", "ROC-AUC", "PR-AUC", "Brier / ECE"], comparison_rows, [1.45, 0.85, 0.85, 0.85, 0.85, 1.65], font_size=7.5)
    add_heading(doc, "8.2 Champion selection", 2)
    add_body(doc, "On the 1,200-row validation-policy partition, Linear SVC reached 0.997500 macro-F1 and Logistic Regression reached 0.995000. The 0.002500 advantage was below the predefined 0.01 retention tolerance, so Logistic Regression remained selected for direct signed explanations, compact deployment and preservation of the verified production artifact. Final-test results are reported transparently but do not retroactively tune this decision.")
    core_metric_keys = {"accuracy", "precision", "recall", "f1", "macro_f1", "weighted_f1", "roc_auc", "pr_auc"}
    champion_rows = [(key.replace("_", " ").title(), value) for key, value in metrics.items() if key in core_metric_keys]
    add_table(doc, ["Metric", "Actual value"], champion_rows, [3.4, 3.1], font_size=8.4)
    add_body(doc, f"Training: {metrics['train_samples']:,} rows; validation: {metrics['validation_samples']:,} rows divided into calibration and policy subsets; untouched final test: {metrics['test_samples']:,} rows; calibrated Brier/ECE: {metrics['brier_score']:.6f}/{metrics['expected_calibration_error']:.6f}; median-run mean inference: {metrics['mean_inference_ms_per_article']:.3f} ms/article.")
    eval_figures = [
        ("model_comparison.png", "Figure 8.1. Held-out model comparison.", "All three classical baselines perform strongly, but differences in recall, native probability support and explanation support affect the champion decision."),
        ("confusion_matrix.png", "Figure 8.2. Production-model final-test confusion matrix.", "Among 1,200 reliable-labelled and 1,199 misleading-labelled rows, Logistic Regression produced 1,196 true negatives, 1,184 true positives, four false positives and 15 false negatives."),
        ("roc_pr_curves.png", "Figure 8.3. ROC and precision-recall curves.", "Near-perfect curves on ISOT should be interpreted with the disclosed source/style shortcut risk."),
        ("feature_importance.png", "Figure 8.4. Global Logistic Regression coefficients.", "Global coefficients reveal which dataset terms are associated with each class and help audit suspicious shortcuts."),
    ]
    for filename, caption, explanation in eval_figures:
        add_figure(doc, FIGURES / filename, caption, explanation)
    add_heading(doc, "8.3 Summarization evaluation", 2)
    add_table(doc, ["Metric", "Precision", "Recall", "F1"], [
        ("ROUGE-1", summary_metrics["rouge1_precision"], summary_metrics["rouge1_recall"], summary_metrics["rouge1_f1"]), ("ROUGE-2", summary_metrics["rouge2_precision"], summary_metrics["rouge2_recall"], summary_metrics["rouge2_f1"]), ("ROUGE-L", summary_metrics["rougeL_precision"], summary_metrics["rougeL_recall"], summary_metrics["rougeL_f1"]),
    ], [1.55, 1.65, 1.65, 1.65], font_size=9)
    add_body(doc, f"Evaluation used {summary_metrics['sample_size']} XSum test articles selected with seed {summary_metrics['random_seed']}. Mean compression was {summary_metrics['mean_compression_ratio_pct']:.4f}% and mean latency {summary_metrics['mean_latency_ms']:.3f} ms. XSum references are professionally written, single-sentence and highly abstractive; an extractive system can be coherent and faithful while receiving low lexical overlap. Qualitative consistency still requires human review.")
    add_heading(doc, "8.4 Packaged demonstration analyses", 2)
    sample_rows = [[row["sample_file"], row["original_words"], row["summary_words"], f"{row['compression_ratio_pct']:.2f}%", row["prediction_label"], f"{row['misleading_probability']:.2%}", row["confidence_band"]] for _, row in samples.iterrows()]
    add_table(doc, ["Synthetic sample", "Words", "Summary", "Compression", "Displayed label", "Misleading p", "Band"], sample_rows, [1.52, 0.52, 0.65, 0.78, 1.28, 0.85, 0.90], font_size=7.1)
    add_body(doc, "These are synthetic UI demonstrations, not ground-truth evaluation. They exercise the reliable, misleading and uncertainty display paths without republishing labelled news text.")
    add_heading(doc, "8.5 Test-case summary", 2)
    test_rows = [[row["Test ID"], row["Module"], row["Expected result"], row["Actual result"], row["Status"]] for _, row in tests.iterrows()]
    add_table(doc, ["ID", "Module", "Expected", "Actual", "Status"], test_rows, [0.55, 1.05, 1.95, 2.20, 0.75], font_size=6.9)
    add_heading(doc, "8.6 Error analysis", 2)
    add_body(doc, "False positives and false negatives are retained as hashes, probabilities and word counts in error_analysis.csv. Manual inspection should focus on satire, opinion, short event reports, outlet-specific phrasing and topics near the training boundary. A low numerical test error does not remove the need for cross-publisher and temporal validation.")

    add_chapter(doc, 9, "Graphical User Interface")
    add_body(doc, "NewsLens AI uses a calm newsroom system: warm ivory paper surfaces, charcoal text, muted brown rules, serif display headings, sans-serif reading text and restrained semantic colours. app.py registers all six sections with st.Page and st.navigation; native Streamlit links preserve a single browser tab, direct routes and browser history without duplicating business logic. Focus indicators, reduced-motion handling, colour-independent result labels and stacked narrow-screen layouts support accessibility.")
    add_callout(doc, "Screenshot set", "Figures 9.2-9.16 document the current beige editorial interface. Their filenames, dimensions and SHA-256 hashes are recorded in reports/results/ui_screenshot_manifest.json.", fill="E8EEE8", accent=TEAL)
    add_figure(doc, DIAGRAMS / "12_streamlit_navigation_diagram.png", "Figure 9.1. Streamlit navigation diagram.", "The responsive top navigation exposes News Desk, Analyse Article, Model Accountability, Dataset Analysis, Editorial Archive and Research & About.")
    screenshot_info = [
        ("01_home.png", "Figure 9.2. News Desk.", "The landing page states the scientific boundary before users see a prediction."),
        ("02_analysis_input.png", "Figure 9.3. Article input.", "Input method, title, full text and summary settings are visible. URL and file alternatives use the same validation pipeline."),
        ("03_summary_and_risk_results.png", "Figure 9.4. Summary and credibility-risk result.", "The synthetic demonstration produced a 28-word summary, 80.4% compression and 84.3% misleading probability."),
        ("04_explainability_and_downloads.png", "Figure 9.5. Explanation and exports.", "Observed TF-IDF contributions appear on both sides of zero; the warning and export controls remain adjacent."),
        ("05_editorial_review_required.png", "Figure 9.6. Editorial-review required state.", "The abstention state separates uncertain model output from a publishable editorial conclusion."),
        ("06_model_accountability.png", "Figure 9.7. Model Accountability.", "Cards and charts are populated from saved measured artifacts rather than hard-coded values."),
        ("07_model_benchmarking.png", "Figure 9.8. Controlled model benchmarking.", "Three fixed classical candidates are compared on identical partitions while model retention remains validation-based."),
        ("08_calibration_reliability.png", "Figure 9.9. Calibration reliability.", "Brier score, expected calibration error and the review policy are presented with dataset-relative limits."),
        ("09_dataset_analysis.png", "Figure 9.10. Dataset Analysis.", "Data quality, distributions and leakage disclosures are organised as a numbered research appendix."),
        ("10_newsroom_analytics.png", "Figure 9.11. Session-local newsroom analytics.", "Aggregate signals avoid article text and personal identifiers."),
        ("11_drift_readiness.png", "Figure 9.12. Drift readiness.", "The application refuses to overstate drift when the current private archive has too few observations."),
        ("12_editorial_review_workflow.png", "Figure 9.13. Editorial review workflow.", "A human can record evidence URLs, notes and a final assessment inside the current visitor's archive."),
        ("13_research_about.png", "Figure 9.14. Research & About.", "Verified paper metadata and responsible-use boundaries remain visible in the application."),
        ("14_home_mobile.png", "Figure 9.15. Mobile News Desk.", "The masthead, navigation and editorial hierarchy adapt to a 390-pixel viewport without horizontal clipping."),
        ("15_analysis_mobile.png", "Figure 9.16. Mobile article workflow.", "Input controls stack into a single readable column and retain accessible labels and touch-sized actions."),
    ]
    for filename, caption, explanation in screenshot_info:
        screenshot_width = 3.2 if filename.endswith("_mobile.png") else 6.3
        add_figure(doc, SCREENSHOTS / filename, caption, explanation, width=screenshot_width)
    add_callout(doc, "Capture record", "The 15 interface images are genuine local Streamlit captures using the current ui/theme.py CSS, shared component class names and session-safe measured/synthetic artifacts in Chromium. The capture manifest records each filename, dimension, capture date and SHA-256 hash. scripts/capture_streamlit_screenshots.py is the reproducible capture workflow.", fill="E8EEE8", accent=TEAL)

    add_chapter(doc, 10, "Project Management")
    add_heading(doc, "10.1 Work breakdown and timeline", 2)
    add_table(doc, ["Phase", "Weeks", "Outputs"], [
        ("Research and requirements", "1-2", "Scope, papers, dataset comparison, risk register"), ("Data engineering and EDA", "3-4", "Cleaning audit, visualisations, split protocol"), ("Model development", "5-6", "Three tuned classifiers, champion artifact, error analysis"), ("Summarization and ingestion", "7-8", "Extractive/optional abstractive modules, URL/file fallbacks"), ("Application and storage", "9-10", "Streamlit pages, SQLite, exports, XAI"), ("Testing and documentation", "11-12", "56 checks, responsive screenshots, diagrams, report and guides"),
    ], [2.0, 0.9, 3.6], font_size=8.7)
    add_heading(doc, "10.2 Risks and mitigation", 2)
    add_table(doc, ["Risk", "Impact", "Mitigation"], [
        ("Dataset source leakage", "Inflated accuracy", "Remove duplicates/metadata/markers; disclose residual risk; future group split"), ("Transformer download/RAM", "Slow or unavailable mode", "Extractive fallback; lazy cache; lighter checkpoint"), ("Paywall/anti-bot URL", "No extractable body", "Layered extractor and direct-paste guidance"), ("False misinformation label", "Reputational/ethical harm", "Risk language, abstention, XAI and permanent disclaimer"), ("Dependency drift", "Startup failure", "Pinned tested requirements and Python 3.12 guide"), ("Scanned PDF", "No text", "Explain OCR requirement; do not crash"),
    ], [1.7, 1.45, 3.35], font_size=8.1)
    add_heading(doc, "10.3 Cost and resources", 2)
    add_body(doc, "The core workflow requires no paid API. A suitable computer and network access for installation/downloads are still required; hosting, compute, bandwidth and third-party service terms may impose costs. Local operation avoids a mandatory production-platform dependency but is not represented as universally cost-free.")

    add_chapter(doc, 11, "Limitations, Ethics and Future Scope")
    add_heading(doc, "11.1 Technical and dataset limitations", 2)
    add_bullets(doc, ["ISOT is outlet-, topic-, period- and style-specific; random splitting may be optimistic.", "English-only TF-IDF does not model multilingual meaning or cross-lingual misinformation.", "The classifier does not retrieve evidence, reason over claims or inspect images/video.", "Satire, parody, opinion, clickbait and evolving events can cross the learned boundary.", "DistilBART may hallucinate or omit details; extractive summaries may be disfluent or redundant.", "URL extraction cannot defeat paywalls, consent walls or browser-only rendering reliably."])
    add_heading(doc, "11.2 Ethical considerations", 2)
    add_body(doc, "Misinformation labels can amplify political, regional and publisher bias. False positives may damage legitimate reporting; false negatives may increase trust in harmful content. Feature explanations can be mistaken for reasons or evidence. The system therefore avoids definitive 'true/false' language, supports uncertainty, explains that terms are correlations, keeps data local and asks users to verify important claims independently.")
    add_heading(doc, "11.3 Future enhancements", 2)
    add_bullets(doc, ["Claim segmentation with evidence retrieval from trusted primary sources.", "Publisher-, topic-, event- and time-grouped external evaluation.", "Probability calibration on genuinely out-of-domain validation data.", "Multilingual models and language-specific evaluation.", "OCR and multimodal image/video consistency analysis.", "Human-in-the-loop explanation usefulness and fairness audits.", "Factual-consistency metrics and human scoring for summaries.", "Authenticated team history with encryption and retention policy only if deployment requires it."])

    add_chapter(doc, 12, "Conclusion")
    add_body(doc, "NewsLens AI demonstrates an end-to-end academic AI system rather than two disconnected code snippets. The summarizer and credibility-risk classifier operate independently on the correct input, and the Streamlit presentation layer combines their outputs with metadata, confidence, explainability, a session-private archive and export. The packaged model loads without retraining, the lightweight path runs on a standard laptop, and all reported model results are calculated. All 56 packaged checks, including the 29 established checks, passed in the Python 3.12 release environment; same-tab routing, responsive layouts, browser history, exports and cross-visitor archive isolation were also exercised in Chromium.")
    add_body(doc, "The held-out ISOT performance is technically strong but scientifically bounded by dataset artefacts. The low extractive XSum ROUGE is also reported rather than hidden. The central conclusion is therefore not that the system can determine truth: it can provide a fast, inspectable credibility-risk baseline and a concise summary while teaching why evidence, generalisation and human judgement remain necessary.")

    doc.add_page_break()
    add_heading(doc, "References", 1)
    for paper in papers:
        venue = paper["Publisher/conference/journal"]
        add_body(doc, f"[{paper['Sr. No.']}] {paper['Authors']}, '{paper['Paper title']},' {venue}, {paper['Year']}, doi: {paper['DOI']}. {paper['Source URL']}")
    extra_refs = [
        "[11] University of Victoria ISOT Research Lab, 'Fake News Detection Datasets,' official dataset page, https://onlineacademiccommunity.uvic.ca/isot/2022/11/27/fake-news-detection-datasets/.",
        "[12] Edinburgh NLP Group, 'XSum Dataset,' official project repository, https://github.com/EdinburghNLP/XSum.",
        "[13] NeuralNine, 'Summarize News Articles with Machine Learning in Python,' YouTube conceptual reference, https://youtu.be/z4DQYprjPSs.",
        "[14] 'Fake News Detection in Python,' YouTube conceptual reference supplied with the project brief, https://youtu.be/ZE2DANLfBIs.",
        "[15] scikit-learn developers, 'Pipeline, TfidfVectorizer, model selection and metrics documentation,' https://scikit-learn.org/.",
        "[16] Streamlit, 'Streamlit documentation,' https://docs.streamlit.io/.",
    ]
    for ref in extra_refs:
        add_body(doc, ref)

    doc.add_page_break()
    add_heading(doc, "Appendix A: Installation and User Manual", 1)
    add_code(doc, "py -3.12 -m venv .venv\n.venv\\Scripts\\activate\npython -m pip install --upgrade pip\npip install -r requirements-lite.txt\nstreamlit run app.py")
    add_numbered(doc, ["Open Analyse Article from the top navigation.", "Choose Paste text, Public URL or File upload.", "Select summary method and length.", "Provide at least 40 words and select Analyse Article.", "Read the summary, probability, confidence band, local terms and disclaimer.", "Download JSON/PDF or reopen the item in Editorial Archive."])
    add_heading(doc, "Appendix B: Dataset Schema", 1)
    add_table(doc, ["Field", "Origin", "Use"], [("title", "ISOT CSV", "Combined with body"), ("text", "ISOT CSV", "Primary content"), ("subject", "ISOT CSV", "EDA only; excluded from model"), ("date", "ISOT CSV", "EDA only; excluded from model"), ("label", "True/Fake filename", "0 reliable; 1 misleading"), ("document", "XSum", "Summary input"), ("summary", "XSum", "ROUGE reference"), ("id", "XSum", "Stable sample identifier")], [1.2, 1.5, 3.8], font_size=8.5)
    add_heading(doc, "Appendix C: Function Reference", 1)
    add_table(doc, ["Function", "Input", "Output"], [("extract_article", "public URL", "ArticleData"), ("parse_uploaded_file", "filename + bytes", "clean text"), ("clean_article_text", "raw text", "display-clean text"), ("summarize_extractive", "text + length", "SummaryResult"), ("summarize_abstractive", "text + pipeline + length", "SummaryResult"), ("predict_credibility", "original text + Pipeline", "PredictionResult"), ("insert_analysis", "record dict", "ID + duplicate flag"), ("analysis_pdf_bytes", "analysis dict", "PDF bytes")], [2.15, 2.0, 2.35], font_size=8.4)
    add_heading(doc, "Appendix D: Research Matrix", 1)
    add_body(doc, "The complete 16-column matrix is supplied as NewsLens_AI_Research_Paper_Matrix.xlsx with Overview, Literature Matrix and Coverage Map sheets. It includes access status, dataset, method, metrics, major findings, limitations, relevance and identified gap for all ten papers.")
    add_heading(doc, "Appendix E: Reproduction Checklist", 1)
    add_bullets(doc, ["Verify Python 3.12 and activate the virtual environment.", "Install requirements-lite.txt for the core path.", "Run python -m pytest -q and confirm all 56 checks pass, including the 29 established checks.", "Launch streamlit run app.py and analyse all three packaged samples.", "For training, download official ISOT files and compare SHA-256 metadata.", "For ROUGE, download XSum and rerun the fixed 150-row seed-42 evaluation.", "Never report synthetic sample labels as truth or benchmark accuracy."])

    path = DOCS / "NewsLens_AI_Project_Report.docx"
    doc.save(path)
    return path


def build_setup_guide() -> Path:
    doc = configure_document("NewsLens AI · Setup and Run Guide", "compact_reference_guide")
    add_cover(doc, "NewsLens AI\nSetup and Run Guide", "Beginner-friendly installation, first run, retraining, testing and troubleshooting", "Practical Guide")
    add_toc(doc, [
        ("1-3. Prerequisites, Python and VS Code", 3),
        ("4-6. Extract, virtual environment and interpreter", 4),
        ("7-9. Dependencies and Streamlit startup", 5),
        ("10. First sample analysis", 6),
        ("11-15. Operations, datasets and retraining", 10),
        ("16-17. Tests and public deployment", 11),
        ("18. Troubleshooting", 11),
        ("19-20. Demo checklist and commands", 13),
    ])
    add_callout(doc, "Recommended path", "Use the tested Python 3.12 runtime and requirements-lite.txt first. It runs the complete extractive summary, saved classifier, explanations, session-private history and exports without PyTorch.", fill="E8EEE8", accent=TEAL)

    add_heading(doc, "1. What you need", 1)
    add_bullets(doc, ["A 64-bit Windows 10/11 computer (Linux/macOS also work).", "At least 4 GB RAM for lightweight mode; 8 GB or more for DistilBART.", "About 1 GB free space for lightweight mode; several GB for datasets and transformer caches.", "Internet for installation/downloads. Core analysis is local after installation.", "The core workflow requires no paid API; hosting, network, compute and third-party terms may still carry costs."])
    add_heading(doc, "2. Install Python", 1)
    add_numbered(doc, ["Visit https://www.python.org/downloads/ and choose Python 3.12.x for Windows.", "Run the installer and tick 'Add python.exe to PATH'.", "Choose Install Now.", "Open Command Prompt and run the verification command below."])
    add_code(doc, "py -3.12 --version")
    add_callout(doc, "Expected", "Python reports 3.11.x. Do not use a Microsoft Store alias if VS Code points to a different interpreter.")
    add_heading(doc, "3. Install VS Code and the Python extension", 1)
    add_numbered(doc, ["Download Visual Studio Code from https://code.visualstudio.com/.", "Install and open VS Code.", "Open Extensions (Ctrl+Shift+X).", "Install the Microsoft 'Python' extension.", "Optionally install Jupyter to open the packaged notebooks."])

    doc.add_page_break()
    add_heading(doc, "4. Extract and open the project", 1)
    add_numbered(doc, ["Clone the public NewsLens-AI repository or download its source archive.", "Extract it to a short path such as C:\\Projects when using a ZIP.", "In VS Code select File > Open Folder.", "Open the NewsLens-AI folder that directly contains app.py.", "Choose Terminal > New Terminal."])
    add_callout(doc, "Common folder mistake", "If the terminal cannot find app.py or requirements-lite.txt, you opened the parent extraction folder. Open the folder containing those files.", fill="F3EBDD", accent=AMBER)
    add_heading(doc, "5. Create and activate a virtual environment", 1)
    add_code(doc, "py -3.12 -m venv .venv")
    add_heading(doc, "5.1 Windows PowerShell", 2)
    add_code(doc, ".venv\\Scripts\\Activate.ps1")
    add_heading(doc, "5.2 Windows Command Prompt", 2)
    add_code(doc, ".venv\\Scripts\\activate.bat")
    add_heading(doc, "5.3 Linux or macOS", 2)
    add_code(doc, "python3.11 -m venv .venv\nsource .venv/bin/activate")
    add_body(doc, "A successful activation normally shows (.venv) at the left of the terminal prompt.")
    add_heading(doc, "6. Select the VS Code interpreter", 1)
    add_numbered(doc, ["Press Ctrl+Shift+P.", "Run 'Python: Select Interpreter'.", "Choose the interpreter inside .venv.", "Open a new terminal and confirm the path."])
    add_code(doc, "python --version\npython -c \"import sys; print(sys.executable)\"")

    doc.add_page_break()
    add_heading(doc, "7. Upgrade pip and install dependencies", 1)
    add_code(doc, "python -m pip install --upgrade pip\npip install -r requirements-lite.txt")
    add_body(doc, "The lightweight file installs Streamlit, pandas, NumPy, scikit-learn, Joblib, Plotly, Matplotlib, Seaborn, URL/PDF libraries, ReportLab and pytest. It excludes PyTorch and Transformers.")
    add_heading(doc, "8. Optional abstractive dependencies", 1)
    add_code(doc, "pip install -r requirements.txt")
    add_callout(doc, "First-use download", "Selecting Abstractive - DistilBART downloads sshleifer/distilbart-cnn-6-6. Keep Extractive selected if the download, disk, RAM or CPU time is unsuitable.", fill="EDE6DD", accent=VIOLET)
    add_heading(doc, "9. Start Streamlit", 1)
    add_code(doc, "streamlit run app.py")
    add_body(doc, "If the `streamlit` command is not recognised, run the module form below. Both commands start the same application.")
    add_code(doc, "python -m streamlit run app.py")
    add_body(doc, "Open http://localhost:8501 if the browser does not open automatically. Keep the terminal running while using the application.")
    add_figure(doc, SCREENSHOTS / "01_home.png", "Figure S.1. News Desk.", "The landing view presents the masthead, exact project boundary, primary/archive actions, original local artwork and technical information strip.", width=6.3)

    add_heading(doc, "10. Analyse the first sample article", 1)
    add_numbered(doc, ["Select Analyse Article in the top navigation.", "Keep Paste text selected.", "Select Load packaged sample or paste one of the files from data/sample.", "Choose Extractive - TF-IDF centroid and Short.", "Select Analyse Article.", "Read the summary and credibility verdict together.", "Inspect local terms and the disclaimer.", "Download JSON or PDF, then open Editorial Archive."])
    add_figure(doc, SCREENSHOTS / "02_analysis_input.png", "Figure S.2. Populated article input.", "Input method, article fields and summary controls share one responsive editorial workspace.", width=6.3)
    add_figure(doc, SCREENSHOTS / "03_summary_and_risk_results.png", "Figure S.3. Analysis result.", "Word counts, compression, executive summary, textual verdict, probability/confidence and responsible interpretation are presented as one editorial report.", width=6.3)
    add_figure(doc, SCREENSHOTS / "04_explainability_and_downloads.png", "Figure S.4. Explanation and exports.", "Muted green and oxblood bars oppose/support the misleading class; they are model correlations, not factual evidence.", width=6.3)

    doc.add_page_break()
    add_heading(doc, "11. Stop, restart and change the port", 1)
    add_body(doc, "Press Ctrl+C in the Streamlit terminal to stop the server. Restart with the same run command after reactivating .venv. If port 8501 is busy, choose another port.")
    add_code(doc, "python -m streamlit run app.py --server.port 8502")
    add_heading(doc, "12. Dataset download and placement", 1)
    add_body(doc, "The application runs immediately with the packaged model. Raw datasets are needed only to reproduce training/evaluation.")
    add_code(doc, "python training/download_data.py --dataset all")
    add_body(doc, "The script places True.csv, Fake.csv and xsum-test.parquet in data/raw. Manual users may place files there with the exact names. Dataset sources and redistribution cautions are documented in docs/DATASET_CARD.md.")
    add_heading(doc, "13. Train the fake-news model", 1)
    add_code(doc, "python training/train_fake_news_models.py")
    add_body(doc, "The script cleans and deduplicates data, samples 24,000 balanced rows, uses a seed-42 stratified split, runs three-fold grid search for three models, evaluates the held-out set, saves the champion Pipeline and regenerates real figures/results. It can take several minutes depending on CPU.")
    add_heading(doc, "14. Evaluate summarization", 1)
    add_code(doc, "python training/evaluate_summarizer.py --sample-size 150 --seed 42")
    add_body(doc, "This regenerates ROUGE-1/2/L, compression and latency on a fixed XSum test sample. Do not compare a different sample size without stating the change.")
    add_heading(doc, "15. Retrain with new data", 1)
    add_numbered(doc, ["Back up models/fake_news_pipeline.joblib and model_metadata.json.", "Map new trustworthy labels to 0 reliable and 1 misleading.", "Audit licence, duplicates, sources, topics, dates and annotation rules.", "Modify preparation only if the new schema is documented.", "Use group/time-aware validation where metadata permits.", "Run tests and inspect new error analysis before replacing the packaged model."])

    doc.add_page_break()
    add_heading(doc, "16. Run automated tests", 1)
    add_code(doc, "python -m pytest -q")
    add_callout(doc, "Verification evidence", "All 56 packaged checks passed in the tested Python 3.12 environment, including the 29 established checks. Browser audits also cover six sections, same-tab routing, direct routes, refresh, back/forward, keyboard activation, analysis outputs, exports, cross-visitor archive isolation and the five required viewport widths.", fill="E8EEE8", accent=TEAL)
    add_heading(doc, "17. Public deployment", 1)
    add_body(doc, "GitHub is the canonical public source. Deploy the Streamlit application from branch main with app.py as the entrypoint and the repository-root requirements.txt. The default NEWSLENS_HISTORY_MODE=session gives each visitor a temporary, isolated SQLite file; do not claim durable cloud history. Configure Vercel with web/ as its project root, set only NEXT_PUBLIC_STREAMLIT_APP_URL to the public Streamlit URL, and verify that /app adds ?embed=true with a same-tab fallback link.")
    add_code(doc, "# Local verification before either deployment\npython -m pytest -q\npython -m streamlit run app.py\n\n# Vercel presentation shell\ncd web\nnpm install\nnpm run build")
    add_callout(doc, "Publication gate", "Do not publish secrets, visitor data, generated databases, unlicensed datasets or a model artifact whose redistribution rights have not been confirmed. Update the README and web/app/site-config.ts only with the real GitHub, Streamlit and Vercel URLs.", fill="F3EBDD", accent=AMBER)

    add_heading(doc, "18. Troubleshooting", 1)
    troubleshooting = [
        ("'streamlit' is not recognized", "Activate .venv, then use `python -m streamlit run app.py`."),
        ("ModuleNotFoundError", "Confirm the selected interpreter is .venv and rerun pip install -r requirements-lite.txt."),
        ("Wrong Python interpreter", "Ctrl+Shift+P > Python: Select Interpreter > choose .venv."),
        ("PowerShell execution policy", "Use Command Prompt activation, or run `Set-ExecutionPolicy -Scope Process Bypass` only for the current window if institutional policy permits."),
        ("Torch installation problem", "Use requirements-lite.txt. Install the optional full file later; consult the official PyTorch selector for your CPU/OS."),
        ("Transformer download problem", "Check internet/disk, retry later, set HF_HOME to a writable folder, or use extractive mode."),
        ("NLTK resource error", "This project does not require downloaded NLTK tokenizers; ensure you are running the packaged code, not a tutorial copy."),
        ("Missing dataset", "Run training/download_data.py or place exact filenames under data/raw."),
        ("Joblib model not found", "Restore models/fake_news_pipeline.joblib or run the training script after downloading ISOT."),
        ("Port already in use", "Use --server.port 8502 or stop the old terminal with Ctrl+C."),
        ("URL extraction failure", "The site may be paywalled or blocking bots. Paste the article text directly."),
        ("Browser does not open", "Visit http://localhost:8501 manually."),
        ("Blank Streamlit page", "Hard refresh, inspect the terminal error, confirm project root, and clear only browser cache/session if needed."),
        ("Windows path problem", "Extract to a short local path without OneDrive permission conflicts; do not run directly inside the ZIP."),
        ("PDF has no text", "The PDF is scanned. Apply OCR outside the app, then upload a text-based PDF or TXT."),
    ]
    for issue, fix in troubleshooting:
        add_heading(doc, issue, 2); add_body(doc, fix)

    add_heading(doc, "19. Beginner demonstration checklist", 1)
    add_bullets(doc, ["News Desk loads without a red error.", "All six top-navigation destinations open.", "Desktop and mobile layouts avoid horizontal clipping.", "Reliable, misleading and uncertain synthetic samples run.", "Extractive summary changes with Short/Medium/Detailed.", "Probabilities sum to approximately 100%.", "The disclaimer is visible.", "JSON and PDF downloads open.", "Editorial Archive shows one row per unique article.", "Model Accountability and Dataset Analysis use real saved charts.", "python -m pytest -q passes."])
    add_heading(doc, "20. Exact quick-reference commands", 1)
    add_code(doc, "py -3.12 -m venv .venv\n.venv\\Scripts\\activate\npython -m pip install --upgrade pip\npip install -r requirements-lite.txt\npython -m pytest -q\npython -m streamlit run app.py")

    path = DOCS / "NewsLens_AI_Setup_and_Run_Guide.docx"
    doc.save(path)
    return path


def build_developer_guide() -> Path:
    doc = configure_document("NewsLens AI · Developer Guide", "compact_reference_guide")
    add_cover(doc, "NewsLens AI\nCode Explanation and Developer Guide", "Module-by-module architecture, algorithms, extension points, tests and viva preparation", "Developer Reference")
    add_toc(doc, [
        ("1-2. Architecture and folder map", 3),
        ("3-4. Configuration and ingestion", 5),
        ("5-7. Preprocessing and summarizers", 6),
        ("8. Offline training pipeline", 7),
        ("9. Prediction and explainability", 8),
        ("10-11. SQLite and Streamlit orchestration", 9),
        ("12-13. Visualisation, exports and tests", 11),
        ("14-16. Modification, deployment and errors", 12),
        ("17-18. Viva and acceptance checklist", 15),
    ])
    add_callout(doc, "Core invariant", "Summarization and classification are independent. `predict_credibility` always receives the original cleaned article, never only `SummaryResult.summary`.", fill="FFF1F4", accent=RED)

    add_heading(doc, "1. How to read the project", 1)
    add_body(doc, "Start with app.py and pages/ for user flow, then inspect src/ for reusable business logic. training/ is an offline pipeline that writes models/ and reports/results/. Streamlit reads those artifacts; it does not train. tests/ verify both isolated functions and page startup.")
    add_code(doc, "presentation -> ingestion -> NLP -> independent AI branches\n             -> explanation/result objects -> persistence/export\noffline training -> saved Joblib + JSON/CSV/PNG artifacts -> runtime loading")
    add_figure(doc, DIAGRAMS / "01_overall_system_architecture.png", "Figure D.1. Layered source architecture.", "Each layer exposes small typed functions; Streamlit orchestrates but does not own model logic.", page_break=False)
    add_figure(doc, DIAGRAMS / "08_component_module_diagram.png", "Figure D.2. Component/module dependencies.", "The offline component writes identified artifacts. Runtime modules read them through central configuration.", page_break=False)

    add_heading(doc, "2. Folder and file purpose", 1)
    file_rows = [
        ("app.py", "st.Page registry and native same-tab st.navigation router"), ("pages/00_News_Desk.py", "News Desk, measured headline cards and scientific boundary"), ("pages/01_Analyse_Article.py", "End-to-end input, independent AI calls, result cards, private history and exports"), ("pages/02_Model_Performance.py", "Measured Model Accountability and ROUGE views"), ("pages/03_Dataset_EDA.py", "Dataset Analysis, EDA and leakage disclosure"), ("pages/04_Analysis_History.py", "Session-private Editorial Archive search/filter/read/export/delete"), ("pages/05_Research_About.py", "Paper matrix and responsible-AI context"), ("src/session_history.py", "Safe session SQLite adapter; trusted-local persistent opt-in"), ("src/config.py", "Portable paths, thresholds, model name and disclaimer"), ("src/utils.py", "Hashing, word/reading statistics, JSON helpers"), ("web/", "Next.js presentation shell and responsive Streamlit iframe"), ("training/*.py", "Download, prepare, tune, evaluate and save"), ("tests/*.py", "Deterministic unit/integration/smoke tests"), ("docs/*.md", "Model, dataset, architecture, deployment and release records"), ("scripts/*.py", "Documentation, audit and screenshot reproducibility helpers"),
    ]
    add_table(doc, ["File / group", "Purpose"], file_rows, [2.3, 4.2], font_size=8.4)

    add_heading(doc, "3. Central configuration", 1)
    add_body(doc, "src/config.py derives PROJECT_ROOT from its own file, so paths work after extraction. Environment variables may override the classifier path, trusted-local database path and optional model. src/session_history.py separately interprets NEWSLENS_HISTORY_MODE and fails closed to per-session storage for the public UI. Constants define 40 minimum words, 10 MB upload size, 15-second request timeout and confidence thresholds.")
    add_code(doc, "MODEL_PATH = Path(os.getenv('NEWSLENS_MODEL_PATH', MODELS_DIR / 'fake_news_pipeline.joblib'))\nDATABASE_PATH = Path(os.getenv('NEWSLENS_DATABASE_PATH', DATABASE_DIR / 'analysis_history.db'))\n# NEWSLENS_HISTORY_MODE defaults to 'session'; only trusted local use may select 'persistent'.\nUNCERTAIN_THRESHOLD = 0.60\nHIGH_CONFIDENCE_THRESHOLD = 0.80")

    add_heading(doc, "4. Ingestion modules", 1)
    add_heading(doc, "4.1 Public URL extraction", 2)
    add_body(doc, "_validate_public_target parses and length-bounds the scheme/host, rejects credentials, malformed or obfuscated authorities and local names, and requires every DNS answer to be globally routable. _download_public_html disables environment proxies and automatic redirects, independently validates every redirect, caps the chain at five hops, blocks loops, compares an observable peer address with the validated DNS set, and streams at most 5 MB. extract_article then tries Trafilatura metadata-aware extraction before its BeautifulSoup fallback.")
    add_code(doc, "def extract_article(url: str) -> ArticleData:\n    html, final_url, headers = _download_public_html(url)\n    document = trafilatura.bare_extraction(html, with_metadata=True)\n    # Every requested hop was validated; response bytes were bounded.\n    # If extracted text is too short: BeautifulSoup article/main/p fallback")
    add_heading(doc, "4.2 File parser", 2)
    add_body(doc, "parse_uploaded_file accepts a filename and bytes so it is independent from Streamlit UploadedFile. TXT decoding tries UTF-8 variants, CP1252 and Latin-1. PDF uses pypdf page extraction and rejects protected, malformed, scanned/too-short or over-10-MB input with FileParseError.")
    add_heading(doc, "4.3 Data contracts", 2)
    add_table(doc, ["Dataclass", "Important fields"], [("ArticleData", "text, title, author, date, URL/domain, extractor, word/read counts"), ("SummaryResult", "summary, method, length, counts, compression, time, selected sentences"), ("PredictionResult", "class/display label, two probabilities, confidence/band, model artifact ID, time, explanation")], [1.55, 4.95], font_size=8.8)

    add_heading(doc, "5. Text preprocessing", 1)
    add_body(doc, "clean_article_text is conservative for display: HTML entities, controls, URLs, whitespace and optional source markers are handled without destroying punctuation. text_for_model lowercases, maps digits to a shared token and removes unsupported characters. The same model transformation runs during training and inference. split_sentences uses deterministic regex rules and needs no NLTK download.")
    add_code(doc, "display_text = clean_article_text(raw, remove_source_markers=False)\nmodel_text = text_for_model(display_text)  # marker mitigation + lowercase + number token")
    add_callout(doc, "Design choice", "No stemming/lemmatisation is applied. The linear model can use readable n-grams, negation and style; the UI can display original sentences.")

    add_heading(doc, "6. Extractive summarizer", 1)
    add_body(doc, "summarize_extractive filters sentences, builds a sentence-level TfidfVectorizer, calculates cosine similarity to the mean centroid, adds lead/density bonuses, chooses k from length-specific fraction/cap, sorts indices and returns a frozen SummaryResult. Empty and two-sentence inputs follow explicit low-complexity paths.")
    add_code(doc, "LENGTH_FRACTIONS = {'Short': .18, 'Medium': .30, 'Detailed': .45}\nMAX_SENTENCES = {'Short': 3, 'Medium': 6, 'Detailed': 10}\nchosen = sorted(np.argsort(scores)[-target:].tolist())")
    add_heading(doc, "6.1 Complexity", 2)
    add_body(doc, "For S sentences and V sentence vocabulary, sparse TF-IDF construction is approximately proportional to observed tokens. Centroid similarity is sparse matrix-vector multiplication. The implementation caps selected sentences but does not require a transformer context window.")

    add_heading(doc, "7. Abstractive summarizer", 1)
    add_body(doc, "load_transformer_pipeline imports optional dependencies lazily, then creates a Hugging Face summarization pipeline for sshleifer/distilbart-cnn-6-6. pages/01 wraps it in st.cache_resource. summarize_abstractive tokenizes full sentences, forms context-safe chunks with one-sentence overlap, summarizes each chunk and optionally summarizes the concatenated first pass.")
    add_code(doc, "chunks = sentence_aware_chunks(text, tokenizer, max_tokens, overlap_sentences=1)\nfirst_pass = [pipeline(chunk, ...)[0]['summary_text'] for chunk in chunks]\nif tokens(join(first_pass)) > max_tokens:\n    final = hierarchical_reduce(first_pass)")
    add_callout(doc, "Dependency compatibility", "requirements.txt keeps Transformers below major release 5 because the summarization pipeline API used here is not available in that release.", fill="EDE6DD", accent=VIOLET)

    add_heading(doc, "8. Training pipeline", 1)
    add_figure(doc, DIAGRAMS / "09_ml_training_pipeline.png", "Figure D.3. Offline model training.", "The Pipeline boundary is the leakage boundary: vectorizer fitting occurs only within training data/folds.", page_break=False)
    add_heading(doc, "8.1 Dataset preparation", 2)
    add_body(doc, "load_isot_dataset assigns labels from filenames, normalises text, removes rows below 40 words, hashes model text, removes conflicting hashes, removes exact duplicates and draws an equal seed-42 sample per class. It returns both a DataFrame and an auditable profile dict.")
    add_heading(doc, "8.2 Candidates and tuning", 2)
    add_body(doc, "base_pipeline defines TfidfVectorizer(stop_words='english', ngram_range=(1,2), min_df=3, max_df=.92, max_features=40000, sublinear_tf=True). GridSearchCV tunes C for Logistic Regression/LinearSVC and alpha for MultinomialNB using three-fold macro-F1. The held-out 20% is never a GridSearch input.")
    add_heading(doc, "8.3 Champion persistence", 2)
    add_code(doc, "joblib.dump(champion, MODEL_PATH, compress=3)\ndump_json(MODEL_METADATA_PATH, metadata)\n# Streamlit later calls joblib.load; it never calls fit().")
    add_body(doc, "Metadata records the model artifact ID, champion name, best parameters, dataset URLs/hashes, split sizes, label mapping, seed, leakage controls and limitations.")

    add_heading(doc, "9. Prediction and explainability", 1)
    add_figure(doc, DIAGRAMS / "10_combined_inference_pipeline.png", "Figure D.4. Runtime inference contracts.", "Both branches share the original cleaned text, but neither consumes the other's output.", page_break=False)
    add_heading(doc, "9.1 Model loading", 2)
    add_body(doc, "load_model checks existence, catches incompatible Joblib failures and verifies predict/named_steps. The page-level cached_model wrapper guarantees one load per Streamlit process.")
    add_heading(doc, "9.2 Probability and band", 2)
    add_body(doc, "_probabilities maps pipeline class order explicitly. It supports predict_proba and has a logistic conversion fallback for decision_function. predict_credibility applies communication thresholds and uses responsible display labels.")
    add_heading(doc, "9.3 Local linear contribution", 2)
    add_code(doc, "x = tfidf.transform([model_text])\ncontribution_j = x_j * classifier.coef_[0, j]\n# sort observed non-zero terms by signed contribution")
    add_body(doc, "Only terms present in the article are shown. Direction and magnitude describe the model's log-odds contribution, not semantic truth, evidence quality or causality.")

    add_heading(doc, "10. SQLite operations", 1)
    add_figure(doc, DIAGRAMS / "11_sqlite_er_diagram.png", "Figure D.5. Local analyses table.", "The unique article_hash performs duplicate detection without retaining the complete source article.", page_break=False)
    add_heading(doc, "10.1 Lifecycle", 2)
    add_numbered(doc, ["connect creates the parent folder and enables foreign keys.", "initialize_database executes idempotent schema/index statements.", "insert_analysis attempts one insert and catches only the unique-hash IntegrityError.", "list_analyses constructs parameterised search/label clauses and a capped limit.", "get_analysis, delete_analysis and clear_history use parameterised SQL."])
    add_heading(doc, "10.2 Privacy", 2)
    add_body(doc, "Stored fields include title, domain, summary and probabilities, but not the full article or uploaded file. Public UI sessions map an opaque random session token to a hashed temporary SQLite filename, so a second visitor cannot list the first visitor's records. A trusted single-user local runtime may explicitly set NEWSLENS_HISTORY_MODE=persistent. SQLite is unencrypted and Community Cloud local storage is temporary; neither is presented as a durable multi-user store.")

    add_heading(doc, "11. Streamlit orchestration", 1)
    add_figure(doc, DIAGRAMS / "12_streamlit_navigation_diagram.png", "Figure D.6. Native same-tab navigation.", "app.py registers all six source files with st.Page; st.navigation owns direct routes, active state and browser history.", page_break=False)
    add_heading(doc, "11.1 Session state", 2)
    add_body(doc, "loaded_sample carries the sample button across rerun; last_analysis retains the completed payload; last_duplicate controls the warning; and _newslens_private_history_id selects the visitor's non-guessable temporary SQLite path. Only serialisable result fields are stored. Changing a widget causes a normal Streamlit rerun, but cached models and the visitor's session analysis remain.")
    add_heading(doc, "11.2 Resource and data caching", 2)
    add_bullets(doc, ["st.cache_resource: Joblib classifier and optional transformer pipeline.", "Stable JSON/CSV/PNG result files: read directly because they are small and immutable during one demonstration.", "SQLite queries: not globally cached, so delete/insert changes are visible immediately."])
    add_heading(doc, "11.3 Editorial shell, CSS and accessibility", 2)
    add_body(doc, "ui/theme.py defines one maintainable warm paper palette, serif/sans typography pairing, spacing/radius tokens, focus-visible controls, reduced-motion handling and responsive grids. app.py and Streamlit's native navigation render the top routes; ui/navigation.py supplies only the branded masthead; ui/components.py uses st.page_link for same-tab calls to action and provides the remaining editorial components. Plotly charts use the same muted publication palette and retain text labels/directional legends.")
    add_heading(doc, "11.4 Error containment", 2)
    add_body(doc, "User-correctable exceptions (ArticleExtractionError, FileParseError, ModelLoadError, AbstractiveDependencyError and ValueError) are shown directly. An unexpected exception receives only a stable generic message in the public interface, avoiding path, stack or implementation-detail leakage while preventing a full page crash.")

    add_heading(doc, "12. Visualisation and exports", 1)
    add_body(doc, "confidence_gauge maps misleading probability to a 0-100 dial with fixed reliable/transition/misleading zones. feature_contribution_chart builds a diverging bar chart. model_comparison_chart reads measured CSV columns. analysis_json_bytes returns readable UTF-8 JSON; archive_csv_bytes prefixes dangerous spreadsheet-formula cells; and analysis_pdf_bytes escapes every user-controlled ReportLab Paragraph value before constructing the document in memory.")
    add_figure(doc, SCREENSHOTS / "04_explainability_and_downloads.png", "Figure D.7. Explanation and export surface.", "The UI keeps interpretability, limitation text and downloadable evidence in one viewport.", width=6.3, page_break=False)

    add_heading(doc, "13. Tests", 1)
    test_rows = [
        ("test_preprocessing.py", "cleaning, sentence split, stats, hash, language hint"), ("test_summarization.py", "compression, ordering, length, empty and long input"), ("test_prediction.py", "Joblib load, probabilities, labels, XAI shape, missing model"), ("test_database.py", "CRUD, search, duplicate hash and clear"), ("test_article_extractor.py", "obfuscated/private targets, public-to-private redirects, loops, DNS/peer mismatch, size cap and fallback"), ("test_file_parser.py", "TXT success, upload failures and path-traversal filenames"), ("test_exports.py", "JSON, valid escaped PDF and spreadsheet-formula-safe CSV"), ("test_app_smoke.py", "AppTest startup for all six Streamlit scripts"), ("test_editorial_ui.py", "assets, warm tokens, navigation, shared shell and responsible wording"), ("test_release_policy.py", "attribution, proprietary files, deployment gate, no runtime training and web headers/sandbox"),
    ]
    add_table(doc, ["Test file", "Coverage"], test_rows, [2.05, 4.45], font_size=8.4)
    add_code(doc, "python -m pytest -q\n# Expected collection: 56 packaged checks (29 established + 27 hardening)")
    add_body(doc, "For a new feature, add a unit test at the business-logic level and one integration/smoke assertion if it affects UI orchestration. Mock network and filesystem boundaries; do not make deterministic tests depend on a public news site.")

    add_heading(doc, "14. Safe modification recipes", 1)
    add_heading(doc, "14.1 Change UI colours or spacing", 2)
    add_numbered(doc, ["Open ui/theme.py for tokens and global CSS; use ui/components.py for shared presentation patterns.", "Edit shared tokens rather than page-local HTML.", "Keep result text labels; do not rely on colour alone.", "Run AppTest and capture every page at 1440 x 1000 plus the 390 x 844 mobile scenarios."])
    add_heading(doc, "14.2 Change the classifier", 2)
    add_numbered(doc, ["Add a candidate that supports sparse text in training/model_candidates().", "Define a small reproducible grid and macro-F1 scoring.", "If probabilities are absent, add validated calibration rather than treating margins as probabilities.", "Update explainability for the new estimator or disable it explicitly.", "Retrain, regenerate artifacts, update MODEL_CARD.md and rerun all tests."])
    add_heading(doc, "14.3 Add a dataset", 2)
    add_numbered(doc, ["Document source, licence, schema, label policy and annotation quality.", "Map to full article text and binary labels only with justified semantics.", "Deduplicate across datasets and group by publisher/event/time before splitting.", "Keep a test set from unseen groups and report domain-specific errors.", "Record hashes and artifact identifiers, then update DATASET_CARD.md."])
    add_heading(doc, "14.4 Add a summarizer", 2)
    add_numbered(doc, ["Return the existing SummaryResult contract.", "Respect sentence/token limits and CPU fallback.", "Cache resources, not per-user content.", "Evaluate on fixed reference data and add factual/coherence review.", "Never feed the new summary into the classifier unless a separate experiment is explicitly designed."])

    add_heading(doc, "15. Deployment guidance", 1)
    add_figure(doc, DIAGRAMS / "13_deployment_diagram.png", "Figure D.8. Public deployment architecture.", "Vercel provides the editorial presentation shell while Streamlit Community Cloud runs the unchanged Python/ML application from app.py.", page_break=False)
    add_body(doc, "The public Streamlit UI defaults to temporary, per-session SQLite and makes no durable-history promise. GitHub is canonical for code, documentation, issues, releases and deployment history. Vercel uses web/ as the project root and exposes only NEXT_PUBLIC_STREAMLIT_APP_URL; the /app iframe adds ?embed=true. Do not expose local paths, disable URL safety checks, place secrets in browser-visible variables, or publish visitor data. Confirm model and dataset redistribution rights before a public release.")

    add_heading(doc, "16. Common code errors", 1)
    errors = [
        ("Import works in notebook but not terminal", "Run from project root or preserve the project-root sys.path block used by training scripts."), ("Model predicts opposite class", "Inspect pipeline.classes_; never assume probability column order."), ("Leakage after refactor", "Keep TF-IDF inside Pipeline/GridSearch and deduplicate before split."), ("Streamlit retrains repeatedly", "Never call training functions from pages; load Joblib under cache_resource."), ("Empty Plotly explanation", "Estimator may not expose coef_ or article terms may be outside vocabulary; keep an empty-state chart."), ("Locked SQLite file", "Keep connections inside context managers and avoid long transactions."), ("Transformer cuts sentences", "Chunk by tokenizer length at sentence boundaries; reserve generation tokens."), ("Wrong relative path", "Import paths from src.config rather than constructing cwd-specific strings."),
    ]
    for issue, resolution in errors:
        add_heading(doc, issue, 2); add_body(doc, resolution)

    add_heading(doc, "17. Viva questions and concise answers", 1)
    viva = [
        ("Why combine the two tasks?", "They help readers triage news, but remain independent because compression and credibility estimation answer different questions."),
        ("Why classify original text?", "A summary may remove hedging, sourcing, contradictions or style features used by the classifier."),
        ("Why TF-IDF?", "It is fast, sparse, reproducible and interpretable on a CPU."),
        ("Why retain Logistic Regression over Linear SVC?", "Linear SVC's validation-policy macro-F1 advantage was only 0.002500, below the predefined 0.01 tolerance; Logistic Regression preserves the verified artifact and direct coefficient explanations."),
        ("What prevents data leakage?", "Deduplication before split, metadata exclusion, marker neutralisation and TF-IDF fitting inside training folds."),
        ("Why is 99.35% accuracy suspicious?", "ISOT labels correlate with outlet, subject, time and writing style; random rows are not an unseen-domain test."),
        ("What is macro-F1?", "The mean of per-class F1 values, giving equal weight to both labels."),
        ("What does ROC-AUC measure?", "Ranking discrimination across thresholds; it does not establish calibrated probability or fairness."),
        ("What is PR-AUC useful for?", "Positive-class precision/recall trade-offs, especially under imbalance."),
        ("What is local contribution?", "Observed TF-IDF value multiplied by a linear coefficient."),
        ("Is it a fact-checker?", "No. It predicts linguistic credibility risk and retrieves no evidence."),
        ("Why low ROUGE?", "XSum references are highly abstractive one-sentence summaries; extractive overlap is intentionally difficult."),
        ("How are long transformer inputs handled?", "Sentence-aware token chunks, one-sentence overlap and hierarchical reduction."),
        ("What is SSRF?", "A server being tricked into requesting private/internal addresses; the URL validator blocks those ranges."),
        ("Why SQLite?", "Portable, zero-admin local persistence suitable for a semester project."),
        ("How are duplicates detected?", "A SHA-256 hash of case/whitespace-normalised article text has a UNIQUE constraint."),
        ("What happens without PyTorch?", "The complete lightweight extractive/classification path still works."),
        ("How would you improve generalisation?", "Publisher/event/time group splits and external-domain evaluation."),
        ("How would you add real fact-checking?", "Claim extraction, trusted evidence retrieval, entailment and source/citation presentation with human review."),
        ("Main ethical risk?", "False or biased misinformation labels causing reputational or political harm."),
    ]
    for question, answer in viva:
        add_heading(doc, question, 2); add_body(doc, answer)

    add_heading(doc, "18. Developer acceptance checklist", 1)
    add_bullets(doc, ["No model fit call in Streamlit pages.", "Both AI branches receive original cleaned text.", "Every user input has a size/type/length or URL safety check.", "Every probability maps through explicit class order.", "All stored SQL uses parameters and closes promptly.", "New measured values come from scripts, not constants in UI.", "MODEL_CARD.md and DATASET_CARD.md match the packaged artifacts.", "python -m pytest -q passes and every page is visually inspected.", "The disclaimer and uncertainty state remain visible.", "The final ZIP excludes raw copyrighted/large datasets and transient caches."])

    path = DOCS / "NewsLens_AI_Code_Explanation_and_Developer_Guide.docx"
    doc.save(path)
    return path


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    outputs = [build_report(), build_setup_guide(), build_developer_guide()]
    for output in outputs:
        print(output.name, output.stat().st_size)


if __name__ == "__main__":
    main()
