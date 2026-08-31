"""Generate the complete beginner-facing NewsLens AI concepts guide.

The guide explains every project-specific concept in progressively deeper layers:
plain meaning, exact project implementation, measured evidence, and caveats.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.generate_word_documents import (  # noqa: E402
    AMBER,
    BLUE,
    CYAN,
    DIAGRAMS,
    DOCS,
    FIGURES,
    INK,
    LIGHT,
    MUTED,
    NAVY,
    PROJECT_AUTHOR,
    COPYRIGHT_NOTICE,
    RED,
    RESULTS,
    SCREENSHOTS,
    TEAL,
    VIOLET,
    add_body,
    add_bullets,
    add_chapter as _base_add_chapter,
    add_code,
    add_figure,
    add_heading,
    add_numbered,
    add_table,
    configure_document,
    set_paragraph_border,
    set_run,
)


OUTPUT = DOCS / "NewsLens_AI_Complete_Concepts_Methodologies_and_Terminology_Guide.docx"
DOCUMENT_DATE = "16 August 2026"


def add_chapter(doc, number: int, title: str):
    """Start a chapter on a new page without grouping it onto a blank page."""

    paragraph = _base_add_chapter(doc, number, title)
    paragraph.paragraph_format.keep_with_next = False
    if number == 12:
        # The preceding keep-together folder-map table naturally ends a page;
        # allowing normal pagination here avoids an otherwise blank page.
        paragraph.paragraph_format.page_break_before = False
    return paragraph


def add_note(doc, label: str, text: str, *, fill: str = "EEF5FF", accent: str = BLUE) -> None:
    """Add a shaded paragraph callout without using a layout table."""

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    set_paragraph_border(paragraph, "left", color=accent, size=18, space=7)
    label_run = paragraph.add_run(f"{label}: ")
    set_run(label_run, size=9.6, color=accent, bold=True)
    text_run = paragraph.add_run(text)
    set_run(text_run, size=9.6, color=INK)


def _new_decimal_numbering(doc) -> int:
    """Create one real, single-level Word list that restarts at 1."""

    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered(doc, items, *, space_after: float = 4, line_spacing: float = 1.25) -> None:
    """Add a compact-reference numbered list that restarts at 1 per call."""

    num_id = _new_decimal_numbering(doc)
    for item in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(space_after)
        paragraph.paragraph_format.line_spacing = line_spacing
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_node = OxmlElement("w:numId")
        num_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_node)
        p_pr.insert(0, num_pr)
        run = paragraph.add_run(str(item))
        set_run(run, size=11, color=INK)


def add_definition(doc, term: str, meaning: str, *, project_use: str | None = None, caution: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_together = True
    term_run = paragraph.add_run(f"{term}. ")
    set_run(term_run, size=10.6, color=NAVY, bold=True)
    meaning_run = paragraph.add_run(meaning)
    set_run(meaning_run, size=10.6, color=INK)
    if project_use:
        project_run = paragraph.add_run(f" Project use: {project_use}")
        set_run(project_run, size=10.2, color=TEAL)
    if caution:
        caution_run = paragraph.add_run(f" Important: {caution}")
        set_run(caution_run, size=10.2, color=RED)


def add_formula(doc, name: str, formula: str, explanation: str) -> None:
    add_heading(doc, name, 3)
    add_code(doc, formula)
    add_body(doc, explanation)


def add_cover(doc) -> None:
    for _ in range(1):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(16)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    run = kicker.add_run("BEGINNER LEARNING GUIDE")
    set_run(run, size=10.5, color=TEAL, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("NewsLens AI\nComplete Guide to the Project's Concepts,\nMethodologies and Terminology")
    set_run(run, size=28, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    run = subtitle.add_run(
        "Plain-language explanations of the data, NLP, summarization, machine learning, "
        "evaluation, security, storage, testing and responsible-use ideas used in the project"
    )
    set_run(run, size=13, color=BLUE, italic=True)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    set_paragraph_border(rule, "bottom", color=CYAN, size=16, space=5)

    add_table(
        doc,
        ["Guide detail", "Value"],
        [
            ("Intended reader", "A student with little or no prior knowledge of AI, ML, NLP or software terminology"),
            ("Project", "NewsLens AI — AI-based news article summarization and credibility-risk analysis"),
            ("Author and developer", PROJECT_AUTHOR),
            ("Copyright", COPYRIGHT_NOTICE),
            ("Packaged model", "ISOT TF-IDF + Logistic Regression; artifact ID isot-tfidf-lr-v1.0.0"),
            ("Interface design", "Warm beige editorial newsroom with responsive desktop and mobile layouts"),
            ("Evidence basis", "Actual source code, saved model metadata, measured CSV/JSON results, diagrams and tests"),
            ("Reading strategy", "Read Chapters 1-3 first; use later chapters and the A-Z glossary as a reference"),
        ],
        [1.65, 4.85],
        font_size=9.0,
        header_fill="E9FAF7",
    )

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    run = note.add_run(f"Prepared as a companion to the project report and source package · {DOCUMENT_DATE}")
    set_run(run, size=9.4, color=MUTED, italic=True)
    doc.add_page_break()


def add_guide_map(doc) -> None:
    add_heading(doc, "How to use this guide", 1)
    add_note(
        doc,
        "Reader promise",
        "No prior ML knowledge is assumed. Every major term is first explained in everyday language, then connected to the exact project implementation. Technical formulas are included only after their meaning is clear.",
        fill="E9FAF7",
        accent=TEAL,
    )
    add_body(
        doc,
        "The guide is deliberately layered. Chapters 1-3 give the complete story without mathematics. "
        "Chapters 4-17 explain each subsystem precisely. Chapters 18 and the appendices provide viva answers, "
        "formulas, file references and an A-Z terminology lookup. A term may therefore appear twice: once in context "
        "and once as a short glossary definition."
    )
    add_table(
        doc,
        ["Part", "Main question answered", "Recommended use"],
        [
            ("Chapters 1-3", "What does the project do and how does information move through it?", "Read first"),
            ("Chapters 4-5", "Where does the data come from and how is text prepared?", "Dataset and NLP foundation"),
            ("Chapters 6-7", "How are summaries and credibility-risk estimates produced?", "Core AI methods"),
            ("Chapters 8-10", "How was the model trained, measured and explained?", "ML methodology and results"),
            ("Chapters 11-14", "How do the software, inputs, database and interface work?", "System implementation"),
            ("Chapters 15-17", "How is quality checked and what can the results really support?", "Testing and responsible interpretation"),
            ("Chapter 18", "How should I answer common viva questions?", "Oral examination preparation"),
            ("Appendices", "What does a formula, file or unfamiliar term mean?", "Quick reference"),
        ],
        [1.05, 3.45, 2.0],
        font_size=8.7,
    )
    add_heading(doc, "Four labels used throughout", 2)
    add_definition(doc, "Plain meaning", "The shortest everyday-language definition of a concept.")
    add_definition(doc, "Project use", "The exact role that concept has in NewsLens AI.")
    add_definition(doc, "Measured fact", "A value read from the project's saved evaluation artifacts, not an invented example.")
    add_definition(doc, "Important limitation", "A boundary that prevents an apparently strong result from being overclaimed.")
    add_note(
        doc,
        "Most important boundary",
        "NewsLens AI is a linguistic credibility-risk classifier. It does not search the web for evidence, prove whether a claim is true, or replace human fact-checking.",
        fill="FFF1F4",
        accent=RED,
    )


def chapter_big_picture(doc) -> None:
    add_chapter(doc, 1, "The project in plain language")
    add_note(
        doc,
        "In one sentence",
        "A user supplies a news article; the application creates a shorter reading view, independently estimates whether the article's language resembles the dataset's reliable or misleading class, explains the influential words, and optionally saves or exports the result.",
        fill="E9FAF7",
        accent=TEAL,
    )
    add_heading(doc, "1.1 The problem being addressed", 2)
    add_body(
        doc,
        "News articles can be long, repetitive and difficult to assess quickly. The project combines two aids. "
        "Summarization reduces reading effort. Credibility-risk classification highlights articles whose wording resembles "
        "examples labelled misleading during training. The two aids are shown together because a reader may want both, "
        "but they answer different questions."
    )
    add_table(
        doc,
        ["Task", "Question it answers", "What it does not answer"],
        [
            ("Summarization", "What are the main points in fewer words?", "Whether those points are true"),
            ("Credibility-risk classification", "Which learned language pattern does this article resemble?", "Whether independent evidence confirms every claim"),
            ("Explainability", "Which observed terms pushed the linear model in either direction?", "Why the real-world event happened or which source is truthful"),
            ("History/export", "How can a completed analysis be revisited or shared?", "Long-term secure archival or collaborative case management"),
        ],
        [1.5, 2.45, 2.55],
        font_size=8.8,
    )
    add_heading(doc, "1.2 The user journey", 2)
    add_numbered(
        doc,
        [
            "Choose direct text, a public URL, a TXT file or a text-based PDF.",
            "Validate size, type, URL safety and minimum article length.",
            "Extract readable text and create display/model representations.",
            "Send the original cleaned article independently to the summarizer and classifier.",
            "Combine the summary, probabilities, confidence band, influential terms and timings in the interface.",
            "Optionally store one duplicate-aware row in the visitor's private session archive or download JSON/PDF.",
        ],
    )
    add_heading(doc, "1.3 The most important design rule", 2)
    add_body(
        doc,
        "The classifier never receives only the generated summary. A summary can remove uncertainty words, attribution, "
        "contradictions or stylistic signals. Classifying the original cleaned article prevents the summarizer from silently "
        "changing the classifier's evidence. The two outputs meet only in the presentation layer."
    )
    add_figure(
        doc,
        DIAGRAMS / "10_combined_inference_pipeline.png",
        "Figure 1.1. Independent summarization and classification branches.",
        "The red crossed annotation records the scientific invariant: SummaryResult is never used as PredictionResult input.",
        page_break=False,
    )
    add_heading(doc, "1.4 What 'fake news detection' means here", 2)
    add_body(
        doc,
        "The project title uses the familiar phrase 'fake news detection', but the responsible interpretation is narrower. "
        "The model learned statistical patterns from ISOT articles labelled reliable or fake. Its output is therefore a "
        "dataset-derived credibility-risk estimate. It is useful for triage and education, not as a verdict about a person, "
        "publisher or political claim."
    )
    add_note(
        doc,
        "Safe wording",
        "Say 'the model estimates an 84% misleading-risk probability from linguistic patterns.' Do not say 'the article is 84% fake' or 'the model proved the claim false.'",
        fill="FFF5E5",
        accent=AMBER,
    )


def chapter_foundations(doc) -> None:
    add_chapter(doc, 2, "Essential AI, ML and NLP foundations")
    add_note(doc, "In one sentence", "AI is the broad goal, machine learning is the data-driven method, and NLP is the language-focused application area.")
    definitions = [
        ("Artificial Intelligence (AI)", "The broad field of building computer systems that perform tasks associated with human intelligence, such as understanding language or making predictions.", "AI is the umbrella containing both summarization and classification."),
        ("Machine Learning (ML)", "A method in which a system learns useful patterns from examples instead of receiving a hand-written rule for every case.", "The classifier learns term weights from labelled ISOT articles."),
        ("Natural Language Processing (NLP)", "The area of computing concerned with analysing, transforming or generating human language.", "Cleaning, sentence splitting, TF-IDF, summarization and article classification are NLP operations."),
        ("Algorithm", "A defined sequence of steps for solving a problem.", "TF-IDF sentence ranking and Logistic Regression are algorithms."),
        ("Model", "A fitted mathematical object whose parameters were learned from training data.", "The packaged Joblib file contains a fitted vectorizer and Logistic Regression classifier."),
        ("Training", "The stage in which a model learns parameters from labelled examples.", "Training is performed offline by training/train_fake_news_models.py."),
        ("Inference", "Using an already-trained model to process a new input.", "The Streamlit app performs inference; it never retrains the classifier."),
        ("Feature", "A numerical input supplied to a model.", "TF-IDF values for unigrams and bigrams are classifier features."),
        ("Label or target", "The answer attached to a training example.", "ISOT rows are mapped to 0 = reliable and 1 = misleading."),
        ("Classification", "Choosing one category from a defined set.", "The binary classifier selects reliable or misleading before communication thresholds are applied."),
        ("Probability estimate", "A number from 0 to 1 representing the model's confidence under its learned assumptions.", "Reliable and misleading probabilities sum to approximately 1."),
        ("Pipeline", "A fixed chain of transformations that is fitted and used as one object.", "TF-IDF and the classifier stay together so preprocessing is identical during training and inference."),
    ]
    for term, meaning, use in definitions:
        add_definition(doc, term, meaning, project_use=use)
    add_heading(doc, "2.1 Rules versus learned behaviour", 2)
    add_body(
        doc,
        "The project mixes ordinary rules and learned behaviour. File limits, URL safety checks, sentence-count caps and "
        "database constraints are explicit rules. TF-IDF weights, Logistic Regression coefficients and DistilBART's generation "
        "behaviour are learned from data. Knowing which is which helps debugging: a 10 MB rejection is a rule; an incorrect "
        "credibility label is a model error."
    )
    add_table(
        doc,
        ["Type", "Examples in the project", "How it changes"],
        [
            ("Fixed rule", "40-word minimum, 10 MB limit, 15-second timeout, confidence thresholds", "Edit configuration/code and retest"),
            ("Learned parameter", "Logistic coefficients and TF-IDF vocabulary", "Retrain on data"),
            ("Saved artifact", "Joblib pipeline, metadata JSON, metric CSV/JSON, figures", "Regenerate from the responsible script"),
            ("User setting", "Input method, summary method and summary length", "Changed in the Streamlit interface"),
        ],
        [1.25, 3.45, 1.8],
        font_size=8.8,
    )
    add_heading(doc, "2.2 Supervised and unsupervised ideas", 2)
    add_definition(doc, "Supervised learning", "Learning from examples that already have labels.", project_use="Fake-news classification is supervised because every training row has label 0 or 1.")
    add_definition(doc, "Unsupervised-style ranking", "Finding structure without a supplied answer for each input.", project_use="The extractive summarizer ranks sentences inside the current article without a human-written summary label during runtime.")
    add_definition(doc, "Pretrained model", "A model learned earlier on a large external dataset and reused for another task.", project_use="Optional DistilBART is downloaded already trained; NewsLens AI does not train it.")
    add_note(
        doc,
        "Do not confuse",
        "The classifier is trained by this project on ISOT. DistilBART is only loaded as a pretrained optional summarizer. The extractive summarizer is a deterministic ranking algorithm fitted temporarily to sentences from one article.",
        fill="F3F0FF",
        accent=VIOLET,
    )


def chapter_system_flow(doc) -> None:
    add_chapter(doc, 3, "End-to-end methodology")
    add_figure(
        doc,
        DIAGRAMS / "02_end_to_end_data_flow.png",
        "Figure 3.1. Complete runtime data flow.",
        "Boxes are processing stages; arrows show data movement. The black fan-out dot sends the same original cleaned article to both AI branches.",
        page_break=False,
    )
    stages = [
        ("1. Acquire", "Receive text, URL or file bytes. No model work starts yet."),
        ("2. Validate", "Reject unsafe URLs, unsupported or oversized files, empty content and articles below 40 words."),
        ("3. Extract", "For URLs, download HTML and isolate article text; for files, decode TXT or extract PDF text."),
        ("4. Clean", "Create readable display text and a conservative model-specific representation."),
        ("5A. Summarize", "Select important original sentences or run optional DistilBART generation."),
        ("5B. Classify", "Convert the original cleaned text to TF-IDF features, estimate class probabilities and compute local term contributions."),
        ("6. Compose", "Place summary, statistics, probability gauge, label, explanation and disclaimer into one result."),
        ("7. Persist/export", "Use a SHA-256 article hash inside a session-isolated SQLite archive; create JSON or PDF only when requested."),
    ]
    for name, explanation in stages:
        add_definition(doc, name, explanation)
    add_heading(doc, "3.1 Online versus offline work", 2)
    add_table(
        doc,
        ["Offline development/training", "Online application/runtime"],
        [
            ("Download and audit datasets", "Accept one user article"),
            ("Clean, deduplicate and split thousands of labelled rows", "Apply the same saved text transformation"),
            ("Tune three candidate classifiers", "Load the selected Joblib pipeline once"),
            ("Calculate evaluation metrics and figures", "Predict one article and display saved benchmark context"),
            ("Write model and metadata artifacts", "Read artifacts; do not call fit()"),
        ],
        [3.25, 3.25],
        font_size=8.7,
    )
    add_heading(doc, "3.2 Data contracts", 2)
    add_body(doc, "A data contract is an agreed set of fields passed between modules. The project uses frozen dataclasses so each completed result has a predictable shape and cannot be accidentally modified in place.")
    add_table(
        doc,
        ["Contract", "Meaning", "Important fields"],
        [
            ("ArticleData", "Validated/extracted article", "text, title, author, date, source URL/domain, extractor, word/read counts"),
            ("SummaryResult", "Extractive summary response", "summary, method, length, word counts, compression, time, sentence count"),
            ("AbstractiveResult", "DistilBART response", "summary, method, length, word counts, compression, time, chunk count"),
            ("PredictionResult", "Credibility-risk response", "class, display label, two probabilities, confidence/band, artifact ID, time, explanation"),
        ],
        [1.25, 1.75, 3.5],
        font_size=8.3,
    )


def chapter_data(doc, profile: dict[str, object]) -> None:
    add_chapter(doc, 4, "Dataset and data-preparation methodology")
    add_note(doc, "In one sentence", "The classifier learned from cleaned, deduplicated and balanced ISOT articles, with strict separation between training and held-out testing.")
    add_heading(doc, "4.1 Dataset basics", 2)
    add_definition(doc, "Dataset", "An organised collection of examples used for analysis or model development.")
    add_definition(doc, "Row or record", "One example in a dataset.", project_use="One ISOT row represents one article.")
    add_definition(doc, "Column or field", "One property recorded for every row.", project_use="Original fields include title, text, subject and date; the project adds label, combined text, word count and text hash.")
    add_definition(doc, "Schema", "The names, types and meaning of dataset fields.")
    add_definition(doc, "Corpus", "A collection of texts used for language analysis.", project_use="The cleaned ISOT articles form the classifier corpus.")
    add_heading(doc, "4.2 ISOT source and labels", 2)
    add_body(
        doc,
        "The ISOT Fake News Dataset provides two CSV files. True.csv is mapped to label 0 (reliable) and Fake.csv to label 1 (misleading). "
        "The title and article body are combined for model input. Subject and date are retained only for audit/EDA where needed, not as classifier features. "
        "The labels reflect the dataset creators' source collection process; they are not fresh human verdicts about every sentence."
    )
    add_table(
        doc,
        ["Stage", "Rows", "Meaning"],
        [
            ("Raw input", f"{int(profile['raw_rows']):,}", "All rows across True.csv and Fake.csv"),
            ("Eligible unique rows", f"{int(profile['clean_rows']):,}", "Rows remaining after length/conflict/duplicate controls"),
            ("Exact duplicates removed", f"{int(profile['duplicates_removed']):,}", "Repeated model-text hashes removed before splitting"),
            ("Short or empty removed", f"{int(profile['short_or_empty_rows_removed']):,}", "Rows below the 40-word eligibility rule"),
            ("Conflicting-label rows removed", f"{int(profile['conflicting_label_rows_removed']):,}", "Same text carrying more than one class label"),
            ("Balanced training sample", f"{int(profile['training_sample_rows']):,}", "12,000 reliable and 12,000 misleading rows"),
        ],
        [1.75, 0.95, 3.8],
        font_size=8.7,
    )
    add_figure(
        doc,
        FIGURES / "class_distribution.png",
        "Figure 4.1. Balanced 24,000-row modelling sample.",
        "Equal class counts make the model comparison easier to interpret, but they do not represent the real-world prevalence of misinformation.",
        page_break=False,
    )
    add_heading(doc, "4.3 Cleaning, eligibility and deduplication", 2)
    add_definition(doc, "Eligibility rule", "A condition an example must satisfy before it is used.", project_use="Model text must contain at least 40 word-like tokens.")
    add_definition(doc, "Deduplication", "Removing repeated examples.", project_use="SHA-256 hashes of model text are compared before train/test splitting.", caution="Duplicates across train and test would make evaluation unrealistically easy.")
    add_definition(doc, "Conflicting labels", "Identical text attached to different answers.", project_use="All rows belonging to a hash with more than one label would be removed.")
    add_definition(doc, "Sampling", "Selecting a subset from a larger collection.", project_use="A seed-42 sample selects up to 12,000 rows per class.")
    add_definition(doc, "Class balance", "Having similar numbers of examples in each class.", project_use="The working sample is exactly balanced 50/50.")
    add_heading(doc, "4.4 Train/test split and stratification", 2)
    add_body(
        doc,
        "The 24,000-row sample begins with 19,200 training rows and a 4,800-row established holdout. A deterministic near-duplicate screen quarantines two contaminated holdout rows. The remaining 4,798 rows are divided into 2,399 validation rows and a 2,399-row untouched final test. Validation is further split into 1,199 calibration rows and 1,200 policy rows. "
        "random_state=42 makes the partitioning reproducible. The final test is locked away from fitting, calibration, model retention and threshold selection, then used once for reporting."
    )
    add_definition(doc, "Training set", "Examples used to learn model parameters and tune settings.")
    add_definition(doc, "Held-out test set", "Examples excluded from learning and used to estimate performance on unseen rows from the same data source.")
    add_definition(doc, "Stratified split", "A split that preserves class proportions.")
    add_definition(doc, "Random seed", "A fixed starting value for pseudo-random operations.", project_use="Seed 42 makes sampling and splitting repeatable; it does not remove all uncertainty.")
    add_heading(doc, "4.5 Data leakage and shortcuts", 2)
    add_definition(doc, "Data leakage", "Information from the test set or target accidentally enters training, producing an overly optimistic score.", project_use="TF-IDF is fitted inside each Pipeline fold and duplicates are removed before splitting.")
    add_definition(doc, "Shortcut feature", "A pattern correlated with the label for accidental reasons rather than the intended concept.", project_use="Publisher markers, subject, period and writing style can reveal ISOT source collection.")
    add_definition(doc, "Source-marker mitigation", "Removing or neutralising obvious byline/source phrases.", project_use="Reuters-style lead text is removed and wire markers become a generic 'wire-service' token.")
    add_bullets(
        doc,
        [
            "Exact duplicates are removed before splitting.",
            "Subject and date/source columns are not classifier features.",
            "Reuters and byline markers are neutralised in model text.",
            "The vectorizer is fitted only inside training folds through a scikit-learn Pipeline.",
            "Remaining outlet, topic, time and style artefacts are disclosed as limitations.",
        ],
    )
    add_note(
        doc,
        "Why the score can still be too high",
        "A random row split from one dataset may place similar publishers, topics and time periods in both train and test. The model can therefore generalise within ISOT while performing less well on a new country, year, publisher or writing style.",
        fill="FFF5E5",
        accent=AMBER,
    )


def chapter_preprocessing(doc) -> None:
    add_chapter(doc, 5, "Text preprocessing and language representation")
    add_note(doc, "In one sentence", "Preprocessing converts messy article text into consistent representations while preserving enough language for readable summaries and classifier features.")
    add_heading(doc, "5.1 Why two text representations exist", 2)
    add_table(
        doc,
        ["Representation", "Purpose", "Main treatment"],
        [
            ("Display-cleaned text", "Readable summary and UI", "Decode HTML, remove controls/URLs, normalise spacing; retain punctuation and source wording"),
            ("Model text", "Classifier and training", "Apply source mitigation, lowercase, map digits to 'number', keep English letters/apostrophes/hyphens, collapse spaces"),
        ],
        [1.45, 1.55, 3.5],
        font_size=8.6,
    )
    add_heading(doc, "5.2 Cleaning operations", 2)
    terms = [
        ("HTML entity decoding", "Converts encodings such as &amp; into readable characters such as &."),
        ("Control-character removal", "Deletes invisible byte/control characters that can disturb parsing or display."),
        ("URL removal", "Removes web addresses from the model/display text because raw links add noise and may identify a source."),
        ("Whitespace normalisation", "Collapses repeated spaces and excessive blank lines into consistent spacing."),
        ("Lowercasing", "Maps 'News' and 'news' to the same model token."),
        ("Digit normalisation", "Maps every number sequence to the token 'number', reducing memorisation of specific dates or values."),
        ("Character filtering", "Keeps English letters, whitespace, apostrophes and hyphens for the packaged English-only classifier."),
    ]
    for term, meaning in terms:
        add_definition(doc, term, meaning)
    add_heading(doc, "5.3 Tokens, vocabulary and n-grams", 2)
    add_definition(doc, "Token", "A basic text unit counted or encoded by an NLP process. Here it is usually a word-like item.")
    add_definition(doc, "Vocabulary", "The set of terms known to the vectorizer after fitting.")
    add_definition(doc, "Unigram", "One token, such as 'election'.")
    add_definition(doc, "Bigram", "Two adjacent tokens, such as 'president trump'.")
    add_definition(doc, "Trigram", "Three adjacent tokens. Trigrams are inspected in EDA but the classifier uses only unigrams and bigrams.")
    add_definition(doc, "Stop word", "A very common word such as 'the' or 'and' that often carries little topic information.", project_use="The English stop-word list is removed by TF-IDF vectorizers.")
    add_heading(doc, "5.4 Sentence segmentation", 2)
    add_body(
        doc,
        "The extractive summarizer and transformer chunker need sentences. split_sentences uses deterministic regular expressions to identify punctuation followed by a likely new sentence. "
        "It avoids downloading NLTK resources, which improves offline reproducibility. If a very long text still appears as one sentence, newlines are used as a fallback."
    )
    add_definition(doc, "Regular expression (regex)", "A compact pattern language for finding or replacing text.")
    add_definition(doc, "Sentence boundary", "The position at which one sentence ends and the next begins.")
    add_heading(doc, "5.5 Deliberately omitted transformations", 2)
    add_definition(doc, "Stemming", "Aggressively shortens words to crude roots, for example 'studies' to 'studi'.")
    add_definition(doc, "Lemmatisation", "Maps grammatical forms to dictionary forms, for example 'was' to 'be'.")
    add_body(
        doc,
        "Neither stemming nor lemmatisation is applied. This keeps features human-readable for explanations and preserves negation/style details. "
        "The trade-off is a larger vocabulary with separate forms of related words."
    )
    add_note(doc, "Training/inference consistency", "The same text_for_model function is called when preparing ISOT and when analysing a new article. A mismatch would be a form of training-serving skew.", fill="E9FAF7", accent=TEAL)


def chapter_summarization(doc, summary_metrics: dict[str, object]) -> None:
    add_chapter(doc, 6, "Summarization methods")
    add_heading(doc, "6.1 Extractive versus abstractive summarization", 2)
    add_table(
        doc,
        ["Aspect", "Extractive TF-IDF centroid", "Optional DistilBART"],
        [
            ("Output", "Selects complete original sentences", "Generates new wording"),
            ("Core dependency", "scikit-learn + NumPy", "Transformers + PyTorch + downloaded model"),
            ("Offline suitability", "Fully available in lightweight setup", "First use needs model download"),
            ("CPU speed", "Very fast", "Much slower"),
            ("Factual risk", "Cannot invent new words, but may omit context", "Can paraphrase well but may hallucinate or distort"),
            ("Long input", "Sentence ranking has no transformer context limit", "Sentence-aware chunks plus hierarchical reduction"),
        ],
        [1.2, 2.65, 2.65],
        font_size=8.3,
    )
    add_heading(doc, "6.2 Extractive TF-IDF centroid method", 2)
    add_numbered(
        doc,
        [
            "Split the article into sentences and discard sentences shorter than four words.",
            "Build a sentence-by-term TF-IDF sparse matrix using English stop-word removal and word 1-2 grams.",
            "Average all sentence vectors to obtain the document centroid, a numerical representation of the article's central vocabulary.",
            "Calculate cosine similarity between each sentence vector and the centroid.",
            "Add a mild lead-position bonus and information-density bonus.",
            "Choose the highest-scoring number of sentences for Short, Medium or Detailed mode.",
            "Restore selected sentences to their original order so the result reads coherently.",
        ],
    )
    add_definition(doc, "Centroid", "The average vector representing the centre of a group.", project_use="The mean of all sentence TF-IDF vectors represents the article's central term pattern.")
    add_definition(doc, "Cosine similarity", "A measure of direction similarity between two vectors, usually from 0 to 1 for non-negative TF-IDF vectors.", project_use="Sentences close to the centroid receive higher relevance scores.")
    add_definition(doc, "Lead bonus", "A small extra score for early sentences.", project_use="News writing often places the main facts near the beginning; the bonus decays exponentially.")
    add_definition(doc, "Density bonus", "A small reward for sentences containing enough content.", project_use="Sentence length contributes at most 0.04 to the score.")
    add_table(
        doc,
        ["User length", "Target fraction of eligible sentences", "Maximum selected sentences"],
        [("Short", "18%", "3"), ("Medium", "30%", "6"), ("Detailed", "45%", "10")],
        [1.4, 2.75, 2.35],
        font_size=8.8,
    )
    add_formula(doc, "Compression ratio", "compression % = (1 - summary_words / original_words) x 100", "A 70% compression ratio means the summary contains about 30% as many words as the original. It does not mean 70% of meaning was preserved.")
    add_heading(doc, "6.3 Optional DistilBART method", 2)
    add_definition(doc, "Transformer", "A neural-network architecture that models relationships between tokens using attention.")
    add_definition(doc, "BART", "A sequence-to-sequence transformer trained to reconstruct corrupted text and then adapted for generation tasks such as summarization.")
    add_definition(doc, "DistilBART", "A smaller, faster BART-derived checkpoint created by knowledge distillation.", project_use="The optional checkpoint is sshleifer/distilbart-cnn-6-6.")
    add_definition(doc, "Tokenizer", "The model-specific component that turns text into token IDs.")
    add_definition(doc, "Context window", "The maximum number of tokens a transformer can process at once.")
    add_definition(doc, "Chunking", "Splitting long input into smaller segments.", project_use="The implementation groups complete sentences up to about 650 words and overlaps one sentence.")
    add_definition(doc, "Hierarchical reduction", "Summarising chunks first and, when needed, summarising the combined partial summaries again.")
    add_definition(doc, "Deterministic decoding", "Generation without random sampling.", project_use="do_sample=False makes repeated outputs more stable.")
    add_body(
        doc,
        "The transformer pipeline runs on CPU with device=-1. Length-specific minimum and maximum generation bounds are calculated from chunk word count. "
        "If several partial summaries together exceed 180 words, a second summarization pass reduces them. The resource should be cached because model loading is expensive."
    )
    add_note(doc, "Fallback", "If optional transformer packages are unavailable, the app gives an actionable message and the complete extractive mode remains usable.", fill="F3F0FF", accent=VIOLET)
    add_heading(doc, "6.4 How summarization was evaluated", 2)
    add_body(
        doc,
        f"The packaged extractive method was evaluated on a fixed seed-42 sample of {int(summary_metrics['sample_size'])} XSum test articles. "
        "XSum reference summaries are highly abstractive one-sentence outputs, so a sentence-copying method is expected to have modest word-overlap scores."
    )
    add_table(
        doc,
        ["Measure", "Measured value", "Plain interpretation"],
        [
            ("ROUGE-1 F1", f"{float(summary_metrics['rouge1_f1']):.6f}", "Overlap of individual words"),
            ("ROUGE-2 F1", f"{float(summary_metrics['rouge2_f1']):.6f}", "Overlap of adjacent word pairs"),
            ("ROUGE-L F1", f"{float(summary_metrics['rougeL_f1']):.6f}", "Longest in-order word-sequence overlap"),
            ("Mean compression", f"{float(summary_metrics['mean_compression_ratio_pct']):.2f}%", "Average reduction in word count"),
            ("Mean latency", f"{float(summary_metrics['mean_latency_ms']):.3f} ms", "Average extractive processing time on the evaluation run"),
        ],
        [1.5, 1.35, 3.65],
        font_size=8.7,
    )
    add_note(doc, "ROUGE limitation", "ROUGE measures textual overlap, not factual consistency, completeness, coherence or usefulness. Human review is still necessary.", fill="FFF5E5", accent=AMBER)


def chapter_classifier(doc, metadata: dict[str, object]) -> None:
    add_chapter(doc, 7, "Credibility-risk classification")
    add_note(doc, "In one sentence", "The classifier turns article words into TF-IDF numbers, multiplies them by learned Logistic Regression weights and converts the resulting score into two probability estimates.")
    add_heading(doc, "7.1 TF-IDF vectorisation", 2)
    add_definition(doc, "Vector", "An ordered list of numbers.", project_use="Each article becomes a vector with one position for every retained unigram/bigram.")
    add_definition(doc, "Vectorizer", "A fitted component that learns a vocabulary and converts text to vectors.")
    add_definition(doc, "Sparse matrix", "A matrix that stores only non-zero entries.", project_use="Most articles use only a tiny fraction of the 40,000 possible features, so sparse storage saves memory.")
    add_definition(doc, "Term Frequency (TF)", "How strongly a term occurs in one document.")
    add_definition(doc, "Inverse Document Frequency (IDF)", "A weight that reduces the importance of terms appearing in many documents.")
    add_definition(doc, "TF-IDF", "TF multiplied by IDF; a term receives a high value when it matters in the current article but is not common everywhere.")
    add_formula(doc, "Simplified TF-IDF", "TF-IDF(term, document) = sublinear_TF x IDF\nsublinear_TF = 1 + log(raw_count)\nIDF ~= log((1 + number_of_documents) / (1 + documents_containing_term)) + 1", "scikit-learn applies smoothing and L2 normalisation by default. The formula conveys the idea; the saved vectorizer performs the exact calculation.")
    add_table(
        doc,
        ["TF-IDF setting", "Project value", "Purpose"],
        [
            ("ngram_range", "(1, 2)", "Use unigrams and bigrams"),
            ("stop_words", "English", "Remove common function words"),
            ("min_df", "3", "Ignore terms appearing in fewer than three training documents"),
            ("max_df", "0.92", "Ignore terms appearing in more than 92% of training documents"),
            ("max_features", "40,000", "Cap vocabulary/memory"),
            ("sublinear_tf", "True", "Use logarithmic rather than raw term counts"),
            ("strip_accents", "Unicode", "Normalise accented characters"),
        ],
        [1.4, 1.25, 3.85],
        font_size=8.5,
    )
    add_heading(doc, "7.2 Logistic Regression", 2)
    add_definition(doc, "Logistic Regression", "A linear classification algorithm that adds weighted feature values and passes the result through a sigmoid function.")
    add_definition(doc, "Coefficient or weight", "A learned number showing the direction and strength of a feature's influence on the linear score.")
    add_definition(doc, "Intercept", "A learned baseline score before article features are added.")
    add_definition(doc, "Log-odds", "The linear score used by logistic regression before conversion to probability.")
    add_definition(doc, "Sigmoid", "A smooth S-shaped function that maps any real score to a number between 0 and 1.")
    add_formula(doc, "Logistic probability", "z = intercept + sum(feature_j x coefficient_j)\nP(misleading) = 1 / (1 + exp(-z))\nP(reliable) = 1 - P(misleading)", "A positive z pushes probability toward class 1 (misleading); a negative z pushes toward class 0 (reliable). The class boundary is 0.50 before communication bands are applied.")
    add_table(
        doc,
        ["Classifier setting", "Value", "Meaning"],
        [
            ("C", str(metadata["best_params"]["classifier__C"]), "Inverse regularisation strength; selected by grid search"),
            ("class_weight", "balanced", "Automatically compensate for class frequency"),
            ("solver", "liblinear", "Optimisation algorithm suited to this linear binary problem"),
            ("max_iter", "1,200", "Maximum optimisation iterations before stopping"),
            ("random_state", "42", "Reproducible solver behaviour where randomness applies"),
        ],
        [1.5, 1.1, 3.9],
        font_size=8.6,
    )
    add_heading(doc, "7.3 Regularisation", 2)
    add_definition(doc, "Regularisation", "A penalty that discourages excessively large coefficients, reducing overfitting.")
    add_definition(doc, "Hyperparameter C", "The inverse of regularisation strength in Logistic Regression and Linear SVM. Larger C usually allows a more complex fit; smaller C imposes a stronger penalty.")
    add_heading(doc, "7.4 Candidate models", 2)
    add_table(
        doc,
        ["Model", "Core idea", "Role in this project"],
        [
            ("Multinomial Naive Bayes", "Combines class-specific term likelihoods under a simplifying independence assumption", "Fast probabilistic baseline; alpha tuned"),
            ("Logistic Regression", "Linear weighted sum plus sigmoid", "Selected champion: native probabilities, speed and direct coefficients"),
            ("Linear SVM", "Finds a large-margin separating hyperplane", "Highest score but not selected because the predefined tolerance favoured probability/XAI usability"),
        ],
        [1.65, 2.35, 2.5],
        font_size=8.4,
    )
    add_definition(doc, "Naive Bayes", "A probabilistic family that assumes features are conditionally independent given the class. The assumption is not literally true for language but often provides a strong text baseline.")
    add_definition(doc, "Linear Support Vector Machine (SVM)", "A classifier that seeks a separating boundary with a large margin between classes.")
    add_definition(doc, "Margin", "Distance from the decision boundary; a larger margin can improve robustness within the training domain.")
    add_definition(doc, "Alpha", "The smoothing hyperparameter of Multinomial Naive Bayes, preventing zero probabilities for unseen feature/class combinations.")


def chapter_training(doc, benchmark: dict[str, object]) -> None:
    add_chapter(doc, 8, "Training, tuning and champion selection")
    add_figure(
        doc,
        DIAGRAMS / "09_ml_training_pipeline.png",
        "Figure 8.1. Offline training and selection pipeline.",
        "The final test follows its own locked path and is used only after calibration, model-retention and threshold-policy decisions are fixed.",
        page_break=False,
    )
    add_heading(doc, "8.1 Fit, transform and predict", 2)
    add_definition(doc, "fit()", "Learn information from training data.", project_use="TF-IDF learns vocabulary/IDF values and the classifier learns coefficients.")
    add_definition(doc, "transform()", "Apply a learned representation without relearning it.", project_use="The fitted vectorizer converts held-out or user text to the existing feature space.")
    add_definition(doc, "predict()", "Return the selected class label for new feature rows.")
    add_definition(doc, "predict_proba()", "Return class probability estimates when supported.", project_use="Logistic Regression provides this directly.")
    add_heading(doc, "8.2 Why a scikit-learn Pipeline matters", 2)
    add_body(
        doc,
        "A Pipeline binds TF-IDF and the classifier. During each cross-validation fold, the vocabulary is learned only from that fold's training portion. "
        "The same sequence is saved in Joblib and applied during inference. This prevents accidental fitting on validation/test text and eliminates preprocessing drift."
    )
    add_heading(doc, "8.3 Validation partitions", 2)
    add_definition(doc, "Calibration partition", "Held-out rows used only to fit the mapping from a model score to a probability.", project_use="NewsLens AI uses 1,199 calibration rows.")
    add_definition(doc, "Policy partition", "Held-out rows used to choose an operational rule after candidate models and calibrators are fixed.", project_use="NewsLens AI uses 1,200 rows for the retention tolerance and 0.59 review threshold.")
    add_definition(doc, "Final test", "Rows excluded from fitting and every policy choice, used once for reporting.", project_use="The controlled release benchmark uses 2,399 final-test rows.")
    add_body(
        doc,
        "All three candidate pipelines fit only the same 19,200 training rows. Platt calibration fits only the calibration subset. The 1,200-row policy subset applies the predeclared retention and review-threshold rules. The 2,399-row final test is excluded from every one of those steps."
    )
    add_heading(doc, "8.4 Fixed candidate settings", 2)
    add_definition(doc, "Hyperparameter", "A setting chosen before fitting rather than learned as an ordinary model coefficient.")
    add_definition(doc, "Controlled comparison", "A comparison in which candidates receive the same training, calibration, policy and final-test partitions.")
    add_table(
        doc,
        ["Candidate", "Controlled value", "Purpose"],
        [
            ("Logistic Regression C", "2.0", "Verified production candidate"),
            ("Linear SVC C", "1.0", "Maximum-margin comparison"),
            ("Naive Bayes alpha", "0.1", "Probabilistic text baseline"),
        ],
        [2.25, 2.25, 2.0],
        font_size=8.8,
    )
    add_heading(doc, "8.5 Champion selection rule", 2)
    add_body(
        doc,
        "The predefined retention rule keeps Logistic Regression when Linear SVC's macro-F1 advantage on the validation-policy partition is below 0.01. Linear SVC reached 0.997500 and Logistic Regression reached 0.995000, a 0.002500 advantage. Logistic Regression therefore remains selected for direct coefficient explanations, compact deployment and preservation of the verified production artifact. The final test is not used for this choice."
    )
    comparison_table = []
    for row in benchmark["models"]:
        comparison_table.append(
            (
                row["model"],
                f"{float(row['policy_metrics']['macro_f1']):.6f}",
                f"{float(row['metrics']['macro_f1']):.6f}",
                f"{float(row['metrics']['roc_auc']):.6f}",
                f"{float(row['mean_inference_ms_per_article']):.3f} ms",
            )
        )
    add_table(doc, ["Model", "Policy macro-F1", "Final-test macro-F1", "Final-test ROC-AUC", "Inference/article"], comparison_table, [1.55, 1.25, 1.25, 1.25, 1.2], font_size=7.9)
    add_figure(
        doc,
        FIGURES / "model_comparison.png",
        "Figure 8.2. Held-out model comparison.",
        "All three classical models perform strongly on the same-source ISOT test set; the small visual differences must be interpreted with the dataset limitations.",
        page_break=False,
    )
    add_heading(doc, "8.6 Persistence and artifact tracking", 2)
    add_definition(doc, "Joblib", "A Python library used to serialise and reload fitted objects efficiently.", project_use="models/fake_news_pipeline.joblib contains the champion Pipeline.")
    add_definition(doc, "Serialisation", "Converting an in-memory object into bytes that can be saved and later reconstructed.")
    add_definition(doc, "Model artifact", "A saved file produced by training and consumed by runtime code.")
    add_definition(doc, "Model metadata", "Human/machine-readable facts about a model, such as artifact ID, data, parameters, seed, label mapping and limitations.")
    add_definition(doc, "Model artifact ID", "A stable identifier for a particular trained artifact.", project_use="isot-tfidf-lr-v1.0.0 is stored in model_metadata.json and result records.")
    add_note(doc, "Security note", "Joblib/pickle files can execute code while loading. Load only the packaged trusted artifact, not an untrusted file downloaded from an unknown source.", fill="FFF1F4", accent=RED)


def chapter_metrics(doc, metrics: dict[str, object]) -> None:
    add_chapter(doc, 9, "Evaluation metrics explained")
    add_note(doc, "In one sentence", "A metric summarises one aspect of performance; no single number proves that a model is safe, fair or generally correct.")
    add_heading(doc, "9.1 Confusion matrix", 2)
    add_body(
        doc,
        "A confusion matrix counts every combination of true and predicted class. In this project the positive class is misleading (1). "
        "The untouched final test contains 1,200 reliable-labelled and 1,199 misleading-labelled rows."
    )
    add_table(
        doc,
        ["True class", "Predicted reliable", "Predicted misleading"],
        [("Reliable", "1,196 true negatives", "4 false positives"), ("Misleading", "15 false negatives", "1,184 true positives")],
        [1.55, 2.35, 2.6],
        font_size=9.0,
    )
    add_definition(doc, "True positive (TP)", "A misleading-labelled article correctly predicted misleading. Count: 1,184.")
    add_definition(doc, "True negative (TN)", "A reliable-labelled article correctly predicted reliable. Count: 1,196.")
    add_definition(doc, "False positive (FP)", "A reliable-labelled article incorrectly predicted misleading. Count: 4.", caution="This can unfairly cast doubt on a credible article.")
    add_definition(doc, "False negative (FN)", "A misleading-labelled article incorrectly predicted reliable. Count: 15.", caution="This can create false reassurance.")
    add_figure(
        doc,
        FIGURES / "confusion_matrix.png",
        "Figure 9.1. Held-out Logistic Regression confusion matrix.",
        "The diagonal cells are correct predictions; off-diagonal cells are errors. There were 19 errors among 2,399 final-test rows.",
        page_break=False,
    )
    add_heading(doc, "9.2 Accuracy, precision, recall and F1", 2)
    add_formula(doc, "Accuracy", "accuracy = (TP + TN) / (TP + TN + FP + FN)", f"The measured accuracy is {float(metrics['accuracy']):.6f}, meaning 2,380 of 2,399 final-test rows were classified correctly.")
    add_formula(doc, "Precision for misleading", "precision = TP / (TP + FP)", f"Measured precision is {float(metrics['misleading_precision']):.6f}. Among rows predicted misleading, about 99.66% carried the misleading label in this test set.")
    add_formula(doc, "Recall for misleading", "recall = TP / (TP + FN)", f"Measured recall is {float(metrics['misleading_recall']):.6f}. The model identified about 98.75% of misleading-labelled final-test rows.")
    add_formula(doc, "F1 score", "F1 = 2 x precision x recall / (precision + recall)", f"Measured misleading-class F1 is {float(metrics['misleading_f1']):.6f}. F1 balances precision and recall through their harmonic mean.")
    add_definition(doc, "Support", "The number of true examples used to calculate a class metric. Reliable support is 1,200 and misleading support is 1,199.")
    add_heading(doc, "9.3 Macro and weighted averages", 2)
    add_definition(doc, "Macro-F1", "Calculate F1 separately for each class, then take an equal average.", project_use=f"The champion test macro-F1 is {float(metrics['macro_f1']):.6f}; it is the primary balanced selection metric.")
    add_definition(doc, "Weighted-F1", "Average class F1 scores using class support as weights.", project_use=f"Because the final test is almost balanced, weighted-F1 is effectively {float(metrics['macro_f1']):.6f}, nearly identical to macro-F1.")
    add_heading(doc, "9.4 ROC and precision-recall curves", 2)
    add_definition(doc, "Threshold", "The probability cut-off used to convert a score into a class. The basic decision threshold is 0.50.")
    add_definition(doc, "True-positive rate (TPR)", "Another name for recall: TP divided by all actual positives.")
    add_definition(doc, "False-positive rate (FPR)", "FP divided by all actual negatives.")
    add_definition(doc, "ROC curve", "A plot of TPR against FPR as the decision threshold moves.")
    add_definition(doc, "ROC-AUC", "Area under the ROC curve; the probability that a random positive is ranked above a random negative under the evaluation distribution.", project_use=f"Measured ROC-AUC is {float(metrics['roc_auc']):.6f}.")
    add_definition(doc, "Precision-recall curve", "A plot of precision against recall as the threshold moves.")
    add_definition(doc, "PR-AUC / Average Precision", "A summary of the precision-recall trade-off, often informative when the positive class is rare.", project_use=f"Measured PR-AUC is {float(metrics['pr_auc']):.6f}.")
    add_figure(
        doc,
        FIGURES / "roc_pr_curves.png",
        "Figure 9.2. ROC and precision-recall curves.",
        "Near-perfect curves show strong ranking on the held-out ISOT distribution. They do not prove cross-domain generalisation or probability calibration.",
        page_break=False,
    )
    add_heading(doc, "9.5 Calibration, policy and latency", 2)
    add_definition(doc, "Brier score", "The mean squared error of predicted probabilities against binary labels; lower is better.", project_use=f"Platt-calibrated final-test Brier score is {float(metrics['brier_score']):.6f}.")
    add_definition(doc, "Expected calibration error (ECE)", "A binned summary of the gap between average confidence and observed label frequency.", project_use=f"The ten-bin final-test ECE is {float(metrics['expected_calibration_error']):.6f}.")
    add_definition(doc, "Latency", "Elapsed time to process an input.", project_use=f"Measured mean classifier inference was {float(metrics['mean_inference_ms_per_article']):.3f} ms per held-out article in the recorded environment.")
    add_definition(doc, "Editorial-review threshold", "A confidence boundary below which the application abstains from an automatic directional outcome.", project_use=f"The validation-policy threshold is {float(metrics['editorial_review_threshold']):.2f}.")
    add_note(doc, "Metric discipline", "Always state the dataset, split, positive class, sample size and limitations with a metric. '99.21% final-test accuracy' alone is incomplete and potentially misleading.", fill="FFF5E5", accent=AMBER)


def chapter_explainability(doc) -> None:
    add_chapter(doc, 10, "Confidence, explainability and interpretation")
    add_heading(doc, "10.1 From class probability to display label", 2)
    add_body(
        doc,
        "The raw predicted class uses a 0.50 misleading-probability boundary. The interface then computes confidence as the larger of the two class probabilities and applies communication bands. "
        "The bands change how certainty is described; they do not retrain the model."
    )
    add_table(
        doc,
        ["Calibrated confidence", "Decision state", "Displayed outcome"],
        [
            ("Below 0.59", "Abstain", "Editorial review required"),
            ("0.59 or above; class 0", "Automatic direction", "Lower misleading-content risk indicated"),
            ("0.59 or above; class 1", "Automatic direction", "Higher misleading-content risk indicated"),
        ],
        [2.1, 1.45, 2.95],
        font_size=8.7,
    )
    add_definition(doc, "Confidence", "The larger of the two Platt-calibrated class probabilities for the selected class.")
    add_definition(doc, "Editorial-review state", "An abstention outcome caused by low calibrated confidence or an input/scope diagnostic.")
    add_definition(doc, "Calibration", "Agreement between predicted probabilities and observed frequencies, for example whether events predicted at 80% occur about 80% of the time.")
    add_note(doc, "Probability caveat", "A private Platt calibration study improves final-test Brier score and ECE, but it remains relative to ISOT labels. Treat the displayed percentage as calibrated model confidence under the evaluated distribution, not a verified chance that the article is true or false.", fill="FFF5E5", accent=AMBER)
    add_heading(doc, "10.2 Explainable AI (XAI)", 2)
    add_definition(doc, "Explainable AI (XAI)", "Methods that make aspects of a model's behaviour understandable to people.")
    add_definition(doc, "Global explanation", "A model-wide view of which vocabulary terms have the largest coefficients.")
    add_definition(doc, "Local explanation", "A view of which terms present in one article contributed most to its score.")
    add_formula(doc, "Local linear contribution", "contribution_j = article_TFIDF_j x classifier_coefficient_j", "A positive value pushes the log-odds toward misleading; a negative value pushes toward reliable. Only terms actually observed in the article are shown.")
    add_definition(doc, "Direction", "Whether a term pushes the current score toward class 1 or class 0.")
    add_definition(doc, "Magnitude", "The absolute size of a contribution or coefficient.")
    add_definition(doc, "Feature importance", "A broad phrase for how influential a feature appears to a model. In this linear project, coefficients provide a direct but correlation-based measure.")
    add_figure(
        doc,
        FIGURES / "feature_importance.png",
        "Figure 10.1. Global coefficient inspection.",
        "Strong source- and style-related terms reveal why excellent same-dataset performance must not be interpreted as universal truth detection.",
        page_break=False,
    )
    add_heading(doc, "10.3 What the explanation does not mean", 2)
    add_bullets(
        doc,
        [
            "A highlighted term is not proof that the article is true or false.",
            "A coefficient is a learned correlation, not a causal explanation.",
            "The model does not understand a term in the human semantic sense.",
            "A globally strong term may reflect a publisher, time period or data-collection shortcut.",
            "Removing every highlighted term can change the score without improving factual accuracy.",
        ],
    )
    add_figure(
        doc,
        SCREENSHOTS / "04_explainability_and_downloads.png",
        "Figure 10.2. Explanation, disclaimer and exports.",
        "The interface places local term contributions next to the responsible-use warning so users do not mistake correlations for evidence.",
        width=6.3,
        page_break=False,
    )


def chapter_architecture(doc) -> None:
    add_chapter(doc, 11, "Software architecture and project organisation")
    add_figure(
        doc,
        DIAGRAMS / "01_overall_system_architecture.png",
        "Figure 11.1. Six-layer project architecture.",
        "A layer groups responsibilities. Arrows show runtime dependency direction; offline evaluation produces artifacts rather than serving user requests.",
        page_break=False,
    )
    layers = [
        ("Presentation", "Streamlit pages, cards, Plotly charts and downloads presented to the user."),
        ("Ingestion", "Direct text, public URL, TXT/PDF parsing and validation."),
        ("NLP processing", "Cleaning, sentence splitting, metadata and statistics."),
        ("AI", "Extractive/abstractive summarization, TF-IDF + classifier and XAI."),
        ("Persistence", "Joblib model, metadata, visitor-isolated SQLite history and configuration."),
        ("Evaluation", "Metrics, ROUGE, error analysis, figures and automated tests."),
    ]
    for term, meaning in layers:
        add_definition(doc, term + " layer", meaning)
    add_heading(doc, "11.1 Why modular design matters", 2)
    add_definition(doc, "Module", "A Python file containing related code.")
    add_definition(doc, "Package", "A folder of modules importable under one namespace.", project_use="src, ui and training organise reusable code.")
    add_definition(doc, "Separation of concerns", "Keeping different responsibilities in different modules.", project_use="Streamlit orchestrates, src contains business logic, training writes offline artifacts and tests verify behaviour.")
    add_definition(doc, "Dependency", "A library or module required by another component.")
    add_definition(doc, "Interface or contract", "A stable way components exchange data or call functions.")
    add_definition(doc, "Local-first", "Designed to run mainly on the user's own computer.", project_use="The model and core extractive mode require no paid API; public hosting keeps SQLite temporary and session-isolated.")
    add_heading(doc, "11.2 Important Python concepts", 2)
    add_definition(doc, "Function", "A named block of reusable behaviour with inputs and a return value.")
    add_definition(doc, "Class", "A template combining data and behaviour. Do not confuse a Python class with an ML class label.")
    add_definition(doc, "Dataclass", "A Python class focused on structured data fields.", project_use="ArticleData, SummaryResult and PredictionResult are dataclasses.")
    add_definition(doc, "Frozen dataclass", "An immutable dataclass whose fields cannot be reassigned after creation.")
    add_definition(doc, "Exception", "A controlled signal that an operation failed.", project_use="User-correctable ArticleExtractionError, FileParseError and ModelLoadError become friendly messages.")
    add_definition(doc, "Bytes", "Raw binary data.", project_use="Uploaded files and generated JSON/PDF downloads are handled as bytes in memory.")
    add_definition(doc, "Path", "A filesystem location.", project_use="src/config.py builds portable paths from PROJECT_ROOT.")
    add_definition(doc, "Environment variable", "An operating-system setting read by a program.", project_use="Model, database and optional transformer locations/names can be overridden without editing code.")
    add_heading(doc, "11.3 Main folder map", 2)
    add_table(
        doc,
        ["Folder/file", "Responsibility"],
        [
            ("app.py", "Native st.Page registry and same-tab st.navigation router"),
            ("pages/", "News Desk, analysis, performance, EDA, history and research/about pages"),
            ("ui/", "Warm editorial tokens, branded masthead and reusable presentation components"),
            ("assets/", "Original local logo/masthead artwork and attribution record"),
            ("src/", "Reusable ingestion, NLP, summarization, prediction, XAI, database, export and chart logic"),
            ("training/", "Dataset download/preparation, model tuning and summarizer evaluation"),
            ("models/", "Packaged Joblib pipeline and model metadata"),
            ("reports/", "Measured JSON/CSV results, figures, diagrams and screenshots"),
            ("database/", "Trusted-local persistent SQLite location; public UI uses temporary session files"),
            ("web/", "Next.js presentation website and responsive Streamlit iframe"),
            ("data/sample/", "Synthetic demonstration articles"),
            ("notebooks/", "Readable development/evaluation walkthroughs"),
            ("tests/", "Unit, integration, export and Streamlit smoke tests"),
            ("docs/", "Report, guides, cards and research matrix"),
            ("scripts/", "Reproducible diagrams, EDA, screenshots and document generation"),
        ],
        [1.55, 4.95],
        font_size=8.5,
    )


def chapter_ingestion_security(doc) -> None:
    add_chapter(doc, 12, "Input handling, extraction and security")
    add_heading(doc, "12.1 Supported inputs", 2)
    add_table(
        doc,
        ["Input", "How it is handled", "Main limitation"],
        [
            ("Paste text", "Use the pasted article body", "At least 40 words"),
            ("Public URL", "Validate, download HTML, extract metadata/body", "Paywalls, bot blocking and unsupported pages may fail"),
            ("TXT", "Decode common encodings and clean text", "Up to 10 MB and at least 40 words"),
            ("Text-based PDF", "Use pypdf page text extraction", "Scanned/image-only PDFs require OCR first"),
        ],
        [1.1, 3.0, 2.4],
        font_size=8.6,
    )
    add_heading(doc, "12.2 URL terminology", 2)
    add_definition(doc, "URL", "A web address containing a scheme, host and optional path/query.")
    add_definition(doc, "HTTP/HTTPS", "Protocols used to request web resources; HTTPS encrypts transport between client and server.")
    add_definition(doc, "Hostname/domain", "The named network location such as example.com.")
    add_definition(doc, "DNS resolution", "Converting a hostname into one or more IP addresses.")
    add_definition(doc, "IP address", "A numeric network address.")
    add_definition(doc, "Redirect", "A server response that sends the requester to another URL.")
    add_definition(doc, "Timeout", "A maximum waiting period.", project_use="URL requests stop after 15 seconds rather than hanging indefinitely.")
    add_definition(doc, "User-Agent", "An HTTP header identifying the requesting software.", project_use="The extractor sends a browser-like identifier with a NewsLens AI suffix.")
    add_definition(doc, "Content-Type", "An HTTP header describing returned media.", project_use="The extractor expects HTML or text beginning with HTML markup.")
    add_heading(doc, "12.3 SSRF protection", 2)
    add_definition(doc, "Server-Side Request Forgery (SSRF)", "An attack that tricks a server into requesting internal or otherwise protected network addresses.")
    add_definition(doc, "Loopback address", "An address referring back to the same machine, such as 127.0.0.1.")
    add_definition(doc, "Private address", "An IP range used inside local networks and not publicly routed.")
    add_definition(doc, "Link-local address", "An address valid only on the local network segment.")
    add_definition(doc, "Reserved address", "An address range set aside for special use.")
    add_bullets(
        doc,
        [
            "Only complete http:// and https:// URLs are accepted.",
            "Embedded usernames/passwords are rejected.",
            "localhost, 0.0.0.0 and .local names are rejected.",
            "DNS results are checked; private, loopback, link-local and reserved addresses are blocked.",
            "Automatic redirects are disabled; every destination is independently parsed, DNS-validated and requested, with loop detection and a strict five-hop maximum.",
            "Every DNS answer must be globally routable; where the transport exposes it, the connected peer must match the validated address set.",
            "Response headers and streamed bytes enforce a 5 MB HTML limit; environment proxy settings are not trusted for article retrieval.",
        ],
    )
    add_heading(doc, "12.4 Article extraction", 2)
    add_definition(doc, "HTML", "The markup language used to structure web pages.")
    add_definition(doc, "Boilerplate", "Navigation, menus, advertisements, comments and other page text that is not the article body.")
    add_definition(doc, "Trafilatura", "A library specialised in extracting main article text and metadata from HTML.", project_use="It is the first extraction method.")
    add_definition(doc, "BeautifulSoup", "A general HTML parser.", project_use="It is the fallback when Trafilatura returns fewer than 40 words.")
    add_definition(doc, "Fallback", "A secondary method used when the preferred method fails.")
    add_definition(doc, "Metadata", "Descriptive information about content.", project_use="Title, author, publication date, source URL/domain and extractor name are recorded when available.")
    add_heading(doc, "12.5 TXT and PDF concepts", 2)
    add_definition(doc, "Character encoding", "A mapping between bytes and characters.", project_use="TXT decoding tries UTF-8 with/without BOM, CP1252 and Latin-1.")
    add_definition(doc, "BOM", "A byte-order mark that can identify a Unicode encoding at the start of a file.")
    add_definition(doc, "PDF text layer", "Machine-readable text stored inside a PDF.")
    add_definition(doc, "OCR", "Optical Character Recognition, which converts text in an image/scan into machine-readable characters.", caution="The project does not perform OCR.")
    add_definition(doc, "Encrypted/password-protected PDF", "A PDF whose content access is restricted. Files requiring a password are not supported.")
    add_note(doc, "Validation philosophy", "User-correctable input failures produce specific messages instead of a page crash. Unexpected public-interface failures use a stable generic message without exposing paths, stack traces or implementation detail.", fill="E9FAF7", accent=TEAL)


def chapter_persistence(doc) -> None:
    add_chapter(doc, 13, "SQLite history, hashing and exports")
    add_figure(
        doc,
        DIAGRAMS / "11_sqlite_er_diagram.png",
        "Figure 13.1. Local SQLite analyses table.",
        "The schema stores structured outputs and a unique article hash, not the uploaded file or full original article body.",
        page_break=False,
    )
    add_heading(doc, "13.1 Database fundamentals", 2)
    add_definition(doc, "Database", "An organised system for storing and retrieving structured data.")
    add_definition(doc, "SQLite", "A zero-administration relational database stored in one local file.", project_use="Public visitors receive separate temporary SQLite files; trusted single-user local use may explicitly select database/analysis_history.db.")
    add_definition(doc, "Table", "A named collection of rows with defined columns.", project_use="The main table is analyses.")
    add_definition(doc, "Primary key", "A unique identifier for each row.", project_use="analysis_id is an auto-incrementing integer.")
    add_definition(doc, "AUTOINCREMENT", "Automatically assigns the next integer identifier to a new row.")
    add_definition(doc, "NOT NULL", "A constraint requiring a value.")
    add_definition(doc, "CHECK constraint", "A rule limiting acceptable values.", project_use="Probabilities must remain between 0 and 1; counts/times cannot be negative.")
    add_definition(doc, "UNIQUE constraint", "A rule preventing duplicate values.", project_use="article_hash must be unique.")
    add_definition(doc, "Index", "An auxiliary structure that speeds common searches/sorts.", project_use="Timestamp and prediction-label indexes support history queries.")
    add_definition(doc, "Schema", "The SQL definitions of tables, columns, types, constraints and indexes.")
    add_heading(doc, "13.2 CRUD and safe SQL", 2)
    add_definition(doc, "CRUD", "Create, Read, Update and Delete, the basic data operations. This project creates, reads and deletes history; it does not expose arbitrary edits.")
    add_definition(doc, "SQL", "Structured Query Language, used to define and query relational databases.")
    add_definition(doc, "Parameterized query", "SQL with values supplied separately from the statement.", project_use="Search, get and delete operations use parameters to reduce injection risk.")
    add_definition(doc, "Transaction", "A group of database operations committed consistently.", project_use="Context managers keep connections short and commit/close safely.")
    add_definition(doc, "IntegrityError", "A database exception raised when a constraint is violated.", project_use="A duplicate article hash returns the existing analysis ID instead of inserting another row.")
    add_heading(doc, "13.3 SHA-256 hashing and duplicate detection", 2)
    add_definition(doc, "Hash function", "A one-way function that maps input data to a fixed-length fingerprint.")
    add_definition(doc, "SHA-256", "A cryptographic hash algorithm producing a 256-bit digest, commonly shown as 64 hexadecimal characters.")
    add_definition(doc, "Deterministic", "The same normalised input always produces the same output.")
    add_body(
        doc,
        "article_hash lowercases the article, collapses whitespace and hashes the UTF-8 bytes. This detects exact normalised repeats without storing the full original article as the duplicate key. "
        "A hash is not encryption and does not prove authorship or truth."
    )
    add_heading(doc, "13.4 Timestamps and privacy", 2)
    add_definition(doc, "UTC", "Coordinated Universal Time, a timezone-neutral reference.")
    add_definition(doc, "ISO 8601", "A standard timestamp format such as 2026-07-17T04:00:00+00:00.")
    add_body(
        doc,
        "History stores title, source information, word count, generated summary, prediction fields, model artifact ID and processing time. It does not retain the uploaded file or complete original article body. "
        "SQLite is not encrypted by default. Public hosting therefore isolates each visitor in a temporary file and makes no durable-history promise; trusted local installations still rely on operating-system access controls."
    )
    add_heading(doc, "13.5 JSON, CSV and PDF export", 2)
    add_definition(doc, "JSON", "JavaScript Object Notation, a portable text format containing key-value objects, arrays, strings, numbers and booleans.")
    add_definition(doc, "UTF-8", "The character encoding used for the JSON export.")
    add_definition(doc, "ReportLab", "A Python library that creates the compact PDF analysis report.", project_use="All user-controlled Paragraph strings are escaped so markup and external-resource tags remain literal text.")
    add_definition(doc, "Spreadsheet formula injection", "A risk where CSV text beginning with =, +, -, @, a tab or carriage return can be interpreted as a formula by spreadsheet software.", project_use="Archive CSV export prefixes those cells with an apostrophe while preserving readable text.")
    add_definition(doc, "In-memory generation", "Creating output in RAM rather than writing a temporary file.", project_use="JSON/PDF/CSV bytes are produced only when the user requests a download.")
    add_definition(doc, "Portable filename", "A filename stripped of characters that are unsafe across operating systems.")


def chapter_ui(doc) -> None:
    add_chapter(doc, 14, "Streamlit interface and runtime behaviour")
    add_note(
        doc,
        "Interface design",
        "NewsLens AI uses one warm beige newsroom design throughout the application, screenshots and documentation.",
        fill="E8EEE8",
        accent=TEAL,
    )
    add_figure(
        doc,
        DIAGRAMS / "12_streamlit_navigation_diagram.png",
        "Figure 14.1. Six-page Streamlit navigation.",
        "Closely related input, summary, detection and combined-result capabilities remain on one Analyse Article page so they share the same state and validation logic.",
        page_break=False,
    )
    add_definition(doc, "Streamlit", "A Python framework that turns scripts into interactive web applications.")
    add_definition(doc, "Multipage application", "An interface organised into separate navigable pages.")
    add_definition(doc, "Widget", "An interactive control such as a text area, select box, button or uploader.")
    add_definition(doc, "Rerun", "Streamlit re-executes a page script when a widget changes or an action occurs.")
    add_definition(doc, "Session state", "Per-user data that survives Streamlit reruns.", project_use="Loaded samples, the latest analysis and duplicate state are retained.")
    add_definition(doc, "Cache", "Stored reusable output that avoids repeating expensive work.")
    add_definition(doc, "st.cache_resource", "A Streamlit cache intended for long-lived resources.", project_use="The Joblib classifier and optional transformer pipeline are loaded once per process.")
    add_definition(doc, "CSS", "Cascading Style Sheets, rules controlling web-page appearance.", project_use="Shared CSS creates warm paper surfaces, editorial typography, focus styles and responsive layouts.")
    add_definition(doc, "HTML", "Markup inserted for small reusable visual components such as the masthead, navigation, verdict and reading panels.")
    add_definition(doc, "Design token", "A named reusable visual value such as a colour, spacing unit, radius or shadow.", project_use="ui/theme.py centralises the paper, charcoal, brown, semantic colour, spacing and radius tokens.")
    add_definition(doc, "Editorial hierarchy", "A deliberate reading order created by overlines, headlines, decks, rules, body text and captions.")
    add_definition(doc, "Serif typeface", "A typeface whose letter strokes have small finishing marks.", project_use="Georgia gives headlines a publication-like voice.")
    add_definition(doc, "Sans-serif typeface", "A typeface without finishing strokes, commonly used for compact interface text.", project_use="The system sans stack keeps controls and body copy clear.")
    add_definition(doc, "st.Page", "A native Streamlit page definition containing a source file, title and optional route path.", project_use="app.py registers all six NewsLens AI sections as st.Page objects.")
    add_definition(doc, "st.navigation", "Streamlit's native router and page selector.", project_use="It keeps internal routes in one browser tab and supports direct URLs, refresh and browser history.")
    add_definition(doc, "st.page_link", "A native link to a registered Streamlit page.", project_use="The News Desk hero uses it for same-tab calls to action.")
    add_definition(doc, "Top navigation", "A persistent row of native route links near the masthead.", project_use="It connects every working page, retains active state and collapses into Streamlit's mobile navigation control.")
    add_definition(doc, "Responsive design", "Layout rules that adapt to different screen widths.", project_use="Multi-column sections stack into one column on narrow screens.")
    add_definition(doc, "Breakpoint", "A viewport width at which responsive CSS changes the layout.")
    add_definition(doc, "Focus indicator", "A visible outline showing which interactive element has keyboard focus.")
    add_definition(doc, "Reduced motion", "An accessibility preference that asks interfaces to minimise non-essential animation.")
    add_definition(doc, "Plotly", "An interactive charting library.", project_use="Probability gauge, feature contributions and model comparison charts are Plotly figures.")
    add_definition(doc, "Gauge chart", "A dial-like display showing misleading-risk probability from 0 to 100.")
    add_definition(doc, "Diverging bar chart", "Bars extending in opposing directions around zero.", project_use="Local term contributions show reliable-side and misleading-side directions.")
    add_heading(doc, "14.1 Page responsibilities", 2)
    add_table(
        doc,
        ["Page", "Purpose"],
        [
            ("News Desk", "Orientation, measured headline values, architecture and scientific boundary"),
            ("Analyse Article", "Input, summarization, classification, XAI, session-history insert and downloads"),
            ("Model Accountability", "Saved model metrics, comparison, confusion matrix, curves, ROUGE and evaluation limits"),
            ("Dataset Analysis", "Dataset quality, distributions, leakage controls and descriptive charts"),
            ("Editorial Archive", "Search, filter, inspect, export and delete local rows"),
            ("Research & About", "Literature matrix, references, methodology and responsible-use context"),
        ],
        [1.75, 4.75],
        font_size=8.6,
    )
    add_heading(doc, "14.2 Error containment", 2)
    add_definition(doc, "User-correctable error", "A known problem the user can fix, such as too little text or a private URL.")
    add_definition(doc, "Unexpected exception", "A failure not anticipated by a specific validation branch.")
    add_definition(doc, "Graceful degradation", "Keeping the rest of the system usable when an optional feature fails.", project_use="Extractive mode remains available without transformer dependencies.")
    add_note(doc, "Accessibility principle", "Colour is reinforced by text labels such as Likely Reliable, Potentially Misleading and Model Uncertain. The result must never depend on red/green colour alone.", fill="E9FAF7", accent=TEAL)


def chapter_testing(doc) -> None:
    add_chapter(doc, 15, "Testing, reproducibility and deployment concepts")
    add_heading(doc, "15.1 Automated testing", 2)
    add_definition(doc, "Test case", "A defined input, expected behaviour and observed outcome.")
    add_definition(doc, "Unit test", "Tests one function or small component in isolation.")
    add_definition(doc, "Integration test", "Tests several components working together.")
    add_definition(doc, "Smoke test", "A broad check that an application/page starts without a major failure.")
    add_definition(doc, "Mock", "A controlled replacement for an external dependency.", project_use="Network responses are mocked so tests do not depend on a live news website.")
    add_definition(doc, "Assertion", "A statement that must be true for a test to pass.")
    add_definition(doc, "Regression test", "A test that prevents a previously fixed behaviour from breaking again.")
    add_definition(doc, "pytest", "The Python testing framework used by the project.")
    add_body(doc, "The package defines 56 checks covering preprocessing, summarization, model loading/prediction, calibration, editorial review, privacy-safe analytics, local explanation shape, redirect-safe URL ingestion, upload safety, SQLite operations, injection-safe JSON/PDF/CSV exports, attribution and release policy, every Streamlit script and the shared editorial interface contract. The suite includes the 29 established checks plus 27 hardening checks. Chromium audits additionally exercise the six sections, same-tab routes, direct navigation, refresh, back/forward, keyboard activation, analysis outputs, exports, responsive widths and cross-visitor archive isolation.")
    add_table(
        doc,
        ["Test area", "Examples checked"],
        [
            ("Preprocessing", "Cleaning, sentence splitting, word/read statistics, hashing and language hint"),
            ("Summarization", "Empty/short/long input, sentence order, length modes and compression"),
            ("Prediction/XAI", "Joblib load, probabilities, display labels, explanation shape and missing model"),
            ("Ingestion", "Private URL rejection, mocked HTML fallback, TXT and upload failures"),
            ("Persistence/export", "CRUD, duplicate hash, clear history, JSON round-trip and valid PDF"),
            ("UI smoke", "News Desk plus all five secondary Streamlit pages"),
            ("UI contract", "Local assets, warm design tokens, native same-tab navigation, shared shell and responsible wording"),
        ],
        [1.55, 4.95],
        font_size=8.6,
    )
    add_note(doc, "Verification evidence", "reports/results/project_verification.json distinguishes recorded checks from the full installed-environment test command. This is better scientific practice than replacing the baseline JUnit file with an unexecuted claim.", fill="F3EBDD", accent=AMBER)
    add_heading(doc, "15.2 Reproducibility", 2)
    add_definition(doc, "Reproducibility", "The ability to repeat a process and obtain sufficiently consistent results from documented inputs, code and settings.")
    add_definition(doc, "Deterministic test", "A test whose result does not depend on unpredictable external state.")
    add_definition(doc, "Dependency release", "The installed release of a software library; incompatible changes can alter behaviour.")
    add_definition(doc, "requirements.txt", "A list of Python dependencies used to recreate the environment.")
    add_definition(doc, "Virtual environment", "An isolated Python installation for one project, preventing package conflicts.")
    add_definition(doc, "Model/data hash", "A fingerprint used to identify exact source bytes.", project_use="Metadata records SHA-256 hashes for the two raw ISOT CSVs used during training.")
    add_bullets(
        doc,
        [
            "Random seed 42 is recorded for sampling, splitting and applicable models.",
            "The exact model artifact ID, parameters, class mapping and source hashes are saved.",
            "Training and evaluation scripts generate CSV/JSON evidence and figures.",
            "Raw datasets are intentionally excluded from the final ZIP and must be downloaded from official sources.",
            "A different operating system or library release can still produce small timing or numerical differences.",
        ],
    )
    add_heading(doc, "15.3 Runtime and deployment", 2)
    add_definition(doc, "CPU", "Central Processing Unit; the general-purpose processor used by the default project.")
    add_definition(doc, "Localhost", "A network name referring to the user's own computer.", project_use="Streamlit normally runs at localhost:8501.")
    add_definition(doc, "Port", "A numbered network endpoint used by a process.", project_use="8501 is Streamlit's default port.")
    add_definition(doc, "Deployment", "Making an application available in an environment where users can run it.")
    add_definition(doc, "Streamlit Community Cloud", "A service that deploys a Streamlit entrypoint from a GitHub repository.", project_use="NewsLens AI deploys branch main and app.py with root Python dependencies.")
    add_definition(doc, "Vercel", "A hosting platform used here only for the lightweight Next.js presentation website.", project_use="web/ supplies the landing page and /app iframe; Python/ML remains in Streamlit.")
    add_definition(doc, "Iframe", "An HTML element that embeds another public page inside the current page.", project_use="The /app route embeds the public Streamlit URL with ?embed=true and provides a same-tab fallback.")
    add_definition(doc, "Persistent storage", "Storage that survives application restarts/redeployments.", caution="NewsLens AI does not claim durable Community Cloud history; its safe default is temporary session storage.")
    add_definition(doc, "API", "A defined way software components communicate. The core project uses local Python calls and no paid external AI API.")
    add_definition(doc, "Rate limit", "A cap on requests over time, useful for protecting an internet deployment.")
    add_note(doc, "Testing boundary", "Passing tests proves that specified behaviours worked in the tested environment. It does not prove the model is factually correct on all future news or secure under every deployment scenario.", fill="FFF5E5", accent=AMBER)


def chapter_ethics(doc) -> None:
    add_chapter(doc, 16, "Responsible AI, ethics and limitations")
    add_note(
        doc,
        "Authorship and academic integrity",
        f"NewsLens AI was designed and developed by {PROJECT_AUTHOR}. {COPYRIGHT_NOTICE} Public visibility supports demonstration and evaluation; it is not permission to submit, copy, modify, redistribute or publicly host the project as another person's work. Academic references should cite the project and author.",
        fill="F3EBDD",
        accent=AMBER,
    )
    add_heading(doc, "16.1 Why responsible wording matters", 2)
    add_body(
        doc,
        "A misinformation label can affect reputation, politics and public trust. The project therefore uses cautious display labels, an uncertainty state, probabilities, local correlations and an always-visible disclaimer. "
        "The user is asked to confirm important claims through trusted independent sources."
    )
    add_definition(doc, "Bias", "Systematic behaviour that disadvantages or misrepresents groups, topics or sources.")
    add_definition(doc, "Dataset bias", "Bias caused by which examples were collected and how they were labelled.")
    add_definition(doc, "Selection bias", "The observed sample differs systematically from the broader population.")
    add_definition(doc, "Domain shift", "New inputs differ from the training distribution.", project_use="New publishers, countries, years, languages or genres may use different patterns.")
    add_definition(doc, "Generalisation", "Performance on genuinely unseen data beyond the specific training sample.")
    add_definition(doc, "Overfitting", "Learning training-specific detail that does not transfer to new data.")
    add_definition(doc, "Underfitting", "A model is too simple or insufficiently trained to capture useful patterns.")
    add_definition(doc, "Distribution", "The pattern of values and frequencies in data.")
    add_definition(doc, "Adversarial wording", "Text deliberately changed to manipulate a model while preserving the intended message.")
    add_definition(doc, "Satire", "Humorous or ironic content that may resemble misleading language without intending deception.")
    add_definition(doc, "Clickbait", "Sensational wording intended to attract attention; it is not automatically false.")
    add_definition(doc, "Multilingual limitation", "The classifier keeps English letters and was trained on English ISOT text, so non-English use is unsupported.")
    add_heading(doc, "16.2 Harm from errors", 2)
    add_table(
        doc,
        ["Error", "Possible harm", "Responsible response"],
        [
            ("False positive", "Credible reporting may be unfairly doubted", "Show probability/XAI/disclaimer; require independent verification"),
            ("False negative", "Misleading content may receive false reassurance", "Never treat Likely Reliable as proof; verify consequential claims"),
            ("High confidence on shifted domain", "Users may overtrust an unfamiliar-input prediction", "Detect/disclose unsupported language/domain and validate externally"),
            ("Shortcut explanation", "A publisher/style token may be mistaken for evidence", "Describe contributions as correlations and audit global features"),
        ],
        [1.35, 2.45, 2.7],
        font_size=8.4,
    )
    add_heading(doc, "16.3 Privacy and retention", 2)
    add_bullets(
        doc,
        [
            "Uploaded files are parsed in memory and not retained by the project.",
            "The complete original article body is not stored in SQLite.",
            "Structured results, title/source details, summary and an article hash may be stored locally.",
            "Public session SQLite is temporary and isolated but not encrypted; it should not be treated as a high-security vault or durable cloud history.",
            "Users can delete individual rows or clear the complete history after confirmation.",
        ],
    )
    add_heading(doc, "16.4 What would make a real fact-checking system", 2)
    add_numbered(
        doc,
        [
            "Extract individual factual claims from the article.",
            "Search trusted, current and independent primary sources.",
            "Match evidence to each claim and assess support, contradiction or insufficiency.",
            "Display citations, source quality and publication dates.",
            "Use human review for ambiguous, harmful or politically sensitive conclusions.",
            "Continuously monitor domain drift, bias, calibration and appeal/correction processes.",
        ],
    )
    add_note(doc, "Ethical conclusion", "The best use of this project is educational triage: shorten text, expose model signals, teach evaluation and encourage verification—not automate accusations.", fill="E9FAF7", accent=TEAL)


def chapter_results(doc, metrics: dict[str, object], summary_metrics: dict[str, object]) -> None:
    add_chapter(doc, 17, "What the measured results mean in plain English")
    add_heading(doc, "17.1 Classification result", 2)
    add_body(
        doc,
        f"The packaged Logistic Regression pipeline correctly classified 2,380 of 2,399 untouched final-test ISOT rows, giving {float(metrics['accuracy']):.2%} accuracy and {float(metrics['macro_f1']):.6f} macro-F1. "
        "It made four false-positive and 15 false-negative errors. Its ROC-AUC and PR-AUC were both near 1.0, showing excellent ranking within this split."
    )
    add_heading(doc, "17.2 Why Linear SVM was not packaged", 2)
    add_body(
        doc,
        "On the validation-policy partition, Linear SVC achieved 0.997500 macro-F1 and Logistic Regression achieved 0.995000. The 0.002500 advantage was below the predefined 0.01 tolerance. "
        "The project therefore retained the verified Logistic Regression artifact for compact deployment and direct signed coefficients. Every candidate received the same Platt-calibration protocol. This is an engineering trade-off, not a claim that Logistic Regression is universally superior."
    )
    add_heading(doc, "17.3 Summarization result", 2)
    add_body(
        doc,
        f"On {int(summary_metrics['sample_size'])} fixed XSum test articles, the Medium extractive summarizer produced ROUGE-1/2/L F1 values of "
        f"{float(summary_metrics['rouge1_f1']):.6f}, {float(summary_metrics['rouge2_f1']):.6f} and {float(summary_metrics['rougeL_f1']):.6f}. "
        f"Average compression was {float(summary_metrics['mean_compression_ratio_pct']):.2f}% and mean measured latency was {float(summary_metrics['mean_latency_ms']):.3f} ms. "
        "Low word-overlap is expected because XSum references paraphrase aggressively while an extractive method can only reuse source sentences."
    )
    add_heading(doc, "17.4 Strong conclusions that are justified", 2)
    add_bullets(
        doc,
        [
            "The end-to-end Streamlit application works with text, public URL, TXT and text-based PDF inputs under defined validation rules.",
            "The saved Logistic Regression pipeline performs extremely well on the packaged same-source ISOT held-out split.",
            "The leakage controls address exact duplicates, obvious source fields/markers and vectorizer fitting leakage.",
            "The extractive summarizer is fast, deterministic and transparent about modest XSum ROUGE overlap.",
            "Local linear contributions, session-private history, export, tests and responsible-use messaging are implemented and documented.",
        ],
    )
    add_heading(doc, "17.5 Conclusions that are not justified", 2)
    add_bullets(
        doc,
        [
            "The model is 99.35% accurate on all news worldwide.",
            "A high misleading-risk probability proves that a claim is false.",
            "A highlighted term caused the article to be misleading.",
            "The model is unbiased across publishers, languages, political groups or years.",
            "ROUGE alone proves that a summary is factually correct and readable.",
            "Passing 56 automated checks proves the application is secure in every public-cloud deployment.",
        ],
    )
    add_note(doc, "Best one-line viva summary", "The project is a reproducible local NLP system that combines transparent summarization with a high-performing but domain-limited linguistic risk classifier, while explicitly preserving independence, explainability and responsible-use boundaries.", fill="F3F0FF", accent=VIOLET)


def chapter_viva(doc) -> None:
    add_chapter(doc, 18, "Viva questions and concise answers")
    questions = [
        ("What is the main objective?", "To help a reader process a news article by generating a concise summary and an independent credibility-risk estimate with explanation, a private session archive and export."),
        ("Why combine summarization and detection?", "They are complementary reading aids, but the implementation keeps them independent because compression and credibility estimation answer different questions."),
        ("Why classify the original cleaned article?", "A summary can omit wording that the classifier needs; using the original prevents summary-induced information loss or distortion."),
        ("What is NLP?", "Computational processing of human language. This project uses NLP for cleaning, sentence splitting, vectorisation, summarization and classification."),
        ("What is TF-IDF?", "A numerical weight that is high when a term is important in one document but not common across every document."),
        ("Why use unigrams and bigrams?", "Unigrams capture individual words; bigrams preserve short phrases and some local context while keeping a sparse linear model practical."),
        ("Why Logistic Regression?", "Linear SVC's validation-policy macro-F1 advantage was below the 0.01 retention tolerance; Logistic Regression preserves the verified compact artifact and direct signed coefficients for XAI."),
        ("Why not Linear SVC?", "It scored slightly higher on the policy partition, but the predeclared tolerance retained Logistic Regression; both candidates used the same private calibration protocol."),
        ("What is the positive class?", "Label 1, misleading. Precision and recall in the main metric JSON refer to this class."),
        ("What prevents data leakage?", "Deduplication before splitting, metadata exclusion, source-marker mitigation and fitting TF-IDF inside each training-only Pipeline fold."),
        ("What is cross-validation?", "Repeated training/validation splits within the training set used to compare hyperparameters without touching the final test set."),
        ("What is macro-F1?", "The equal average of per-class F1 scores, so both reliable and misleading performance matter equally."),
        ("Why is 99.21% final-test accuracy not enough?", "ISOT contains publisher, topic, time and style correlations. A same-source final test can be easier than genuinely new-domain news."),
        ("What is a confusion matrix?", "A table counting correct and incorrect predictions for every true/predicted class combination."),
        ("What is ROC-AUC?", "A threshold-independent ranking measure; it does not by itself prove calibrated probabilities or external generalisation."),
        ("What is local explainability here?", "For each observed term, TF-IDF value is multiplied by the learned linear coefficient to show its signed contribution."),
        ("Are highlighted words evidence?", "No. They are correlations used by the model, not evidence about the real-world claim."),
        ("What is extractive summarization?", "Selecting important original sentences and preserving their wording."),
        ("What is abstractive summarization?", "Generating new wording that paraphrases the source, here through optional DistilBART."),
        ("How are long transformer inputs handled?", "Sentence-aware chunks with one-sentence overlap are summarised, then combined summaries may receive a second reduction pass."),
        ("What is ROUGE?", "A family of reference-summary word-overlap metrics; it does not directly measure factuality or readability."),
        ("What is SSRF?", "An attack that tricks a server into requesting internal addresses; public URL validation blocks local/private/reserved destinations."),
        ("Why SQLite?", "It is portable and serverless; temporary per-session files provide the smallest safe public archive boundary, while trusted local use can opt into persistence."),
        ("How are duplicates detected?", "A UNIQUE SHA-256 fingerprint of case/whitespace-normalised article text returns the existing row rather than inserting a duplicate."),
        ("What is cached?", "The Joblib model and optional transformer resource, because loading them repeatedly on every Streamlit rerun is expensive."),
        ("Does the app require a paid API?", "No paid API is required for the core workflow. URL input contacts the supplied public website and optional first-use DistilBART downloads from Hugging Face; hosting, network, compute and third-party terms may still carry costs."),
        ("What are the main ethical risks?", "False accusations, false reassurance, dataset bias, political/reputational harm and overtrust in probability or explanations."),
        ("How would you improve generalisation?", "Use publisher/event/time group splits, external datasets, calibration analysis and error audits across domains and languages."),
        ("How would you build real fact-checking?", "Add claim extraction, trusted evidence retrieval, entailment/contradiction assessment, citations and human review."),
        ("What is the strongest project contribution?", "Integration: a reproducible local application connecting safe ingestion, two independent NLP branches, transparent evaluation, XAI, persistence, export, tests and responsible boundaries."),
    ]
    for question, answer in questions:
        add_heading(doc, question, 2)
        add_body(doc, answer)


def appendix_formulas(doc) -> None:
    add_chapter(doc, 19, "Appendix A - Formula and symbol sheet")
    add_table(
        doc,
        ["Symbol", "Meaning"],
        [
            ("TP", "True positives: misleading rows predicted misleading"),
            ("TN", "True negatives: reliable rows predicted reliable"),
            ("FP", "False positives: reliable rows predicted misleading"),
            ("FN", "False negatives: misleading rows predicted reliable"),
            ("x_j", "TF-IDF value of feature j in the current article"),
            ("w_j", "Learned Logistic Regression coefficient for feature j"),
            ("z", "Linear log-odds score before sigmoid"),
            ("P", "Model probability estimate"),
            ("df", "Number of documents containing a term"),
            ("n", "Number of documents in the fitted corpus"),
        ],
        [1.25, 5.25],
        font_size=8.8,
    )
    formulas = [
        ("Accuracy", "(TP + TN) / (TP + TN + FP + FN)", "Overall fraction correct."),
        ("Precision", "TP / (TP + FP)", "Purity of misleading predictions."),
        ("Recall", "TP / (TP + FN)", "Coverage of actual misleading rows."),
        ("F1", "2 x precision x recall / (precision + recall)", "Harmonic balance of precision and recall."),
        ("False-positive rate", "FP / (FP + TN)", "Fraction of reliable rows incorrectly flagged."),
        ("Sigmoid", "1 / (1 + exp(-z))", "Maps log-odds to a 0-1 probability estimate."),
        ("Linear score", "z = intercept + sum(x_j x w_j)", "Adds every observed feature contribution."),
        ("Local contribution", "x_j x w_j", "Signed influence of one observed term."),
        ("Compression ratio", "(1 - summary_words / original_words) x 100", "Percentage reduction in word count."),
        ("Cosine similarity", "dot(a, b) / (norm(a) x norm(b))", "Direction similarity between a sentence and centroid vector."),
        ("Sublinear TF", "1 + log(raw term count)", "Dampens repeated-term influence."),
        ("Smoothed IDF", "log((1 + n) / (1 + df)) + 1", "Downweights terms found in many documents."),
    ]
    add_heading(doc, "19.1 Formula reference", 2)
    for name, formula, explanation in formulas:
        add_formula(doc, name, formula, explanation)
    add_note(doc, "Reading formulas", "A formula is a precise recipe, not a separate concept to memorise blindly. First identify what each symbol counts, then substitute the values and interpret the result in words.", fill="E9FAF7", accent=TEAL)


def appendix_file_map(doc) -> None:
    add_chapter(doc, 20, "Appendix B - File-by-file reference")
    rows = [
        ("app.py", "Native st.Page registry and same-tab st.navigation router"),
        ("pages/00_News_Desk.py", "News Desk, measured cards and project boundary"),
        ("pages/01_Analyse_Article.py", "Complete user analysis orchestration"),
        ("pages/02_Model_Performance.py", "Saved classifier and summarizer metrics/figures"),
        ("pages/03_Dataset_EDA.py", "Dataset quality, EDA and leakage disclosure"),
        ("pages/04_Analysis_History.py", "SQLite search, filter, reopen, export and delete"),
        ("pages/05_Research_About.py", "Research matrix, methodology and ethics"),
        ("src/config.py", "Paths, thresholds, model name/artifact ID and disclaimer"),
        ("src/text_preprocessor.py", "Cleaning, model representation, sentence split and language hint"),
        ("src/article_extractor.py", "Public URL validation, HTML extraction and metadata"),
        ("src/file_parser.py", "TXT decoding and PDF text extraction"),
        ("src/extractive_summarizer.py", "TF-IDF centroid sentence ranking"),
        ("src/abstractive_summarizer.py", "Optional DistilBART chunking/generation"),
        ("src/fake_news_predictor.py", "Joblib loading, probability mapping, labels/bands and result contract"),
        ("src/explainability.py", "Local contributions and global coefficients"),
        ("src/database.py", "Path-parameterized SQLite schema and history operations"),
        ("src/session_history.py", "Temporary visitor isolation and trusted-local persistence opt-in"),
        ("src/report_exporter.py", "JSON and in-memory ReportLab PDF exports"),
        ("src/visualizations.py", "Warm publication-style Plotly figures"),
        ("ui/theme.py", "Warm paper tokens, typography, responsive CSS and accessibility states"),
        ("ui/navigation.py", "Reusable branded masthead; native routing remains in app.py"),
        ("ui/components.py", "Editorial hero, section, verdict, reading and evidence components"),
        ("assets/", "Original local logo/masthead SVG artwork and attribution record"),
        ("src/ui.py", "Compatibility re-export for the shared ui package"),
        ("src/utils.py", "Word/read statistics, SHA-256, timestamps, JSON and filenames"),
        ("training/download_data.py", "Official dataset download and placement"),
        ("training/prepare_dataset.py", "ISOT cleaning, hashing, deduplication and sampling"),
        ("training/train_fake_news_models.py", "Candidate tuning, evaluation, persistence and figures"),
        ("training/evaluate_summarizer.py", "Fixed-sample XSum ROUGE/compression/latency evaluation"),
        ("models/fake_news_pipeline.joblib", "Packaged fitted TF-IDF + Logistic Regression Pipeline"),
        ("models/model_metadata.json", "Model identity, settings, data hashes and limitations"),
        ("database/analysis_history.db", "Trusted-local structured history when persistent mode is explicitly selected"),
        ("web/", "Next.js landing page and responsive Streamlit iframe for Vercel"),
        ("reports/results/", "Measured CSV/JSON evidence and test results"),
        ("reports/figures/", "EDA and evaluation charts"),
        ("reports/diagrams/", "System/design diagrams"),
        ("reports/screenshots/", "Warm editorial desktop/mobile captures"),
        ("tests/", "Automated unit, integration, export, smoke and editorial UI contract tests"),
        ("docs/DATASET_CARD.md", "Dataset source, schema, licensing/usage and limitations"),
        ("docs/MODEL_CARD.md", "Model purpose, metrics, intended use and risks"),
        ("requirements-lite.txt", "Complete lightweight dependencies without transformer stack"),
        ("requirements.txt", "Full optional transformer/PyTorch environment"),
    ]
    add_table(doc, ["File or folder", "Meaning"], rows, [2.45, 4.05], font_size=8.15)
    add_note(doc, "Where to start reading code", "Read pages/01_Analyse_Article.py for the user flow, then follow imported functions into src/. Read training/ only after the runtime path is clear.", fill="E9FAF7", accent=TEAL)


def appendix_glossary(doc) -> None:
    add_chapter(doc, 21, "Appendix C - A-Z glossary")
    add_body(doc, "These are short lookup definitions. The main chapters contain the complete contextual explanations.")
    glossary = [
        ("Abstractive summarization", "Generates new wording that paraphrases the source."),
        ("Accuracy", "Fraction of all predictions that are correct."),
        ("AI", "Broad field of systems performing intelligence-associated tasks."),
        ("Algorithm", "Defined computational procedure for solving a problem."),
        ("Alpha", "Naive Bayes smoothing hyperparameter."),
        ("API", "Defined interface through which software components communicate."),
        ("ArticleData", "Dataclass carrying validated article text and metadata."),
        ("Artifact", "Saved output such as a model, metric file, figure or document."),
        ("AUC", "Area under a performance curve, summarising behaviour across thresholds."),
        ("BART", "Transformer architecture used for sequence generation."),
        ("Baseline", "A simpler comparison method establishing a reference performance."),
        ("BeautifulSoup", "Fallback HTML parser used to extract paragraph text."),
        ("Bias", "Systematic distortion caused by data, labels, modelling or use."),
        ("Bigram", "Two adjacent tokens treated as one feature."),
        ("BOM", "Byte-order mark at the beginning of some Unicode text files."),
        ("Boilerplate", "Non-article webpage content such as navigation or adverts."),
        ("Bytes", "Raw binary data used for uploads and downloads."),
        ("Cache", "Stored reusable resource that avoids repeated expensive work."),
        ("Calibration", "Agreement between predicted probabilities and observed frequencies."),
        ("C", "Inverse regularisation strength for Logistic Regression/SVM."),
        ("Centroid", "Average vector representing the centre of sentence vectors."),
        ("Champion model", "Candidate selected for packaging under a defined rule."),
        ("Character encoding", "Mapping between bytes and text characters."),
        ("Chunk", "A bounded segment of a long article processed separately."),
        ("Class", "A prediction category; here reliable or misleading."),
        ("Class balance", "Relative number of examples belonging to each class."),
        ("Classification", "Assigning an input to a defined category."),
        ("Classifier", "Model that performs classification."),
        ("Cleaning", "Removing or normalising unwanted text content."),
        ("Clickbait", "Sensational attention-seeking language, not automatically false."),
        ("Coefficient", "Learned linear feature weight."),
        ("Compression ratio", "Percentage reduction from original to summary word count."),
        ("Confidence", "Largest predicted class probability."),
        ("Confidence band", "Low, Moderate or High category derived from confidence."),
        ("Confusion matrix", "Counts of true/predicted class combinations."),
        ("Content-Type", "HTTP description of returned media such as HTML."),
        ("Context window", "Maximum token sequence a transformer can process at once."),
        ("Corpus", "Collection of texts used for language modelling/analysis."),
        ("Cosine similarity", "Vector direction similarity used in sentence ranking."),
        ("Cross-validation", "Repeated training/validation splits inside the training set."),
        ("CRUD", "Create, Read, Update and Delete data operations."),
        ("CSV", "Comma-separated text file representing tabular data."),
        ("CSS", "Rules controlling web-interface appearance."),
        ("Dataclass", "Python class designed primarily to store structured fields."),
        ("DataFrame", "Pandas in-memory table with labelled rows and columns."),
        ("Data leakage", "Unfair information flow from evaluation data/target into training."),
        ("Dataset", "Organised collection of examples."),
        ("Decision boundary", "Rule separating predicted classes."),
        ("Decision function", "Raw classifier score before probability conversion."),
        ("Deduplication", "Removal of repeated examples."),
        ("Density bonus", "Small extractive score reward for informative sentence length."),
        ("Dependency", "External library/module required by code."),
        ("Deployment", "Making an application available in a runtime environment."),
        ("Deterministic", "Same input/state gives the same output."),
        ("DistilBART", "Smaller BART-derived pretrained summarization checkpoint."),
        ("Distribution", "Pattern and frequency of values in data."),
        ("DNS", "System translating hostnames into IP addresses."),
        ("Domain shift", "Production inputs differ from training data."),
        ("EDA", "Exploratory Data Analysis: descriptive investigation before modelling."),
        ("Encoding", "Representation of characters as bytes."),
        ("Entity", "Named external actor/data store in a system diagram."),
        ("Error analysis", "Detailed inspection of incorrect predictions."),
        ("Estimator", "scikit-learn object that can be fitted and used for prediction/transformation."),
        ("Exception", "Runtime signal indicating an error condition."),
        ("Extractive summarization", "Selects source sentences without generating new wording."),
        ("Fallback", "Secondary method used when the preferred method fails."),
        ("False negative", "Misleading-labelled article predicted reliable."),
        ("False positive", "Reliable-labelled article predicted misleading."),
        ("Feature", "Numerical input used by a model."),
        ("Feature contribution", "Observed TF-IDF value multiplied by a coefficient."),
        ("Feature importance", "Measure of a feature's influence on model behaviour."),
        ("fit", "Learn model/vectorizer state from training data."),
        ("Fold", "Temporary validation portion in cross-validation."),
        ("F1", "Harmonic mean of precision and recall."),
        ("Frozen dataclass", "Structured object whose fields cannot be reassigned."),
        ("Gauge", "Dial-like probability visual."),
        ("Generalisation", "Performance on truly unseen data."),
        ("Grid search", "Trying a defined set of hyperparameter combinations."),
        ("Hash", "Fixed-length deterministic fingerprint of input data."),
        ("Held-out test set", "Final evaluation examples excluded from training/tuning."),
        ("Hierarchical reduction", "Summarise chunks, then optionally summarise their summaries."),
        ("HTML", "Web-page markup language."),
        ("HTTP/HTTPS", "Protocols used to request web content."),
        ("Hyperparameter", "Setting chosen before fitting rather than learned as a coefficient."),
        ("IDF", "Weight reducing importance of terms appearing in many documents."),
        ("Immutable", "Cannot be modified after creation."),
        ("Index (database)", "Structure that accelerates common queries."),
        ("Inference", "Using a trained model on new input."),
        ("Input validation", "Checking input safety, format, size and content."),
        ("Intercept", "Baseline linear model score."),
        ("IP address", "Numeric network location."),
        ("ISOT", "Dataset supplying True.csv and Fake.csv for classifier training."),
        ("Iteration", "One optimisation step; max_iter caps Logistic Regression steps."),
        ("Joblib", "Library used to save/reload the fitted scikit-learn Pipeline."),
        ("JSON", "Portable key-value text data format."),
        ("Label", "Known training answer or class target."),
        ("Language hint", "Simple Latin-script ratio warning, not a full language detector."),
        ("Latency", "Elapsed processing time."),
        ("Lead bonus", "Extra extractive score for early news sentences."),
        ("Lemmatisation", "Mapping word forms to dictionary forms; not used here."),
        ("Linear model", "Model whose score is a weighted sum of features."),
        ("Local explanation", "Feature contributions for one article."),
        ("Local-first", "Runs mainly on the user's machine."),
        ("Logistic Regression", "Linear classifier with sigmoid probability output."),
        ("Log-odds", "Logarithm of probability odds; Logistic Regression's linear score."),
        ("Macro-F1", "Equal mean of class-specific F1 scores."),
        ("Metadata", "Descriptive facts about data/model/content."),
        ("MIME type", "Standard identifier for a file/content media type."),
        ("min_df/max_df", "Vocabulary filters based on document frequency."),
        ("ML", "Machine Learning: learning patterns from examples."),
        ("Mock", "Controlled substitute for an external dependency in tests."),
        ("Model", "Fitted mathematical object making predictions."),
        ("Model artifact", "Saved trained model file."),
        ("Model card", "Document describing intended use, metrics and risks."),
        ("Module", "Python source file containing related logic."),
        ("Multinomial Naive Bayes", "Term-count probability classifier used as a baseline."),
        ("N-gram", "Sequence of n adjacent tokens."),
        ("NLP", "Natural Language Processing."),
        ("Normalisation", "Converting varied representations into a consistent form."),
        ("OCR", "Recognition of text from images/scans; not implemented here."),
        ("Offline training", "Model development outside the user-facing app."),
        ("Overfitting", "Learning dataset-specific details that fail to transfer."),
        ("Package", "Importable collection of Python modules."),
        ("Parameter", "Value learned by a model or supplied to a function."),
        ("Parameterized SQL", "SQL statement with separately bound values."),
        ("PDF", "Portable Document Format."),
        ("Pipeline", "Ordered fitted transformations treated as one model."),
        ("Plotly", "Interactive chart library used by Streamlit pages."),
        ("Positive class", "Class used for precision/recall; misleading (1)."),
        ("PR-AUC", "Summary area for precision-recall trade-off."),
        ("Precision", "Fraction of positive predictions that are correct."),
        ("Preprocessing", "Transformations applied before modelling."),
        ("Primary key", "Unique identifier for a database row."),
        ("Probability estimate", "Model score from 0 to 1, not verified truth."),
        ("Public URL", "Internet-reachable HTTP(S) address, excluding local/private ranges."),
        ("Random seed", "Fixed value making pseudo-random operations repeatable."),
        ("Recall", "Fraction of actual positives correctly found."),
        ("Regex", "Pattern language for text matching/replacement."),
        ("Regularisation", "Penalty discouraging excessively large coefficients."),
        ("Reliable/misleading", "Dataset-derived binary classes, not absolute philosophical truth labels."),
        ("ReportLab", "Python PDF-generation library."),
        ("Reproducibility", "Ability to repeat a documented process/results."),
        ("Rerun", "Streamlit re-execution after interaction."),
        ("ROC-AUC", "Threshold-independent ranking metric."),
        ("ROUGE", "Reference-summary textual-overlap metric family."),
        ("ROUGE-1/2/L", "Unigram, bigram and longest-sequence overlap variants."),
        ("Row/record", "One dataset or database example."),
        ("Sampling", "Selecting a subset from a larger collection."),
        ("Schema", "Defined structure of data fields/tables."),
        ("scikit-learn", "Python ML library used for vectorisation, pipelines, models and metrics."),
        ("Sentence segmentation", "Splitting prose into sentences."),
        ("Serialisation", "Saving an object as reconstructable bytes."),
        ("Session state", "Per-user values retained across Streamlit reruns."),
        ("SHA-256", "Cryptographic fingerprint algorithm."),
        ("Shortcut", "Accidental label-correlated pattern exploited by a model."),
        ("Sigmoid", "Function mapping a score to 0-1."),
        ("Sparse matrix", "Matrix storing mainly non-zero values."),
        ("Split", "Division into training and test portions."),
        ("SQLite", "Serverless relational database in one local file."),
        ("SSRF", "Attack using a server to access protected network locations."),
        ("Stemming", "Crude word-root reduction; not used here."),
        ("Stop word", "Very common word removed from TF-IDF features."),
        ("Stratification", "Preserving class proportions during a split."),
        ("Streamlit", "Python interactive web-app framework."),
        ("Sublinear TF", "Logarithmically dampened term frequency."),
        ("Support", "Number of true examples underlying a class metric."),
        ("SVM", "Support Vector Machine, a margin-based classifier."),
        ("Target", "Training label the model tries to predict."),
        ("Term frequency", "How often a term occurs in one document."),
        ("Test set", "Locked examples used for final evaluation."),
        ("TF-IDF", "Term importance weight combining local and corpus frequency."),
        ("Threshold", "Cut-off converting a score/probability into a category."),
        ("Token", "Basic text unit processed by an NLP component."),
        ("Tokenizer", "Component mapping text to a model's token IDs."),
        ("Topic shortcut", "Label correlation caused by different class topics."),
        ("Trafilatura", "Primary webpage article-extraction library."),
        ("Training", "Learning parameters from examples."),
        ("Training-serving skew", "Training and runtime preprocessing differ."),
        ("Transformer", "Attention-based neural sequence architecture."),
        ("Trigram", "Three adjacent tokens."),
        ("Tuning", "Choosing hyperparameters through validation."),
        ("Uncertainty", "Insufficient model confidence for a strong display label."),
        ("Unicode", "Standard covering characters from many writing systems."),
        ("Unigram", "One-token feature."),
        ("URL", "Web resource address."),
        ("UTC", "Timezone-neutral Coordinated Universal Time."),
        ("Vector", "Ordered numerical representation."),
        ("Vectorizer", "Component converting text into feature vectors."),
        ("Virtual environment", "Isolated Python dependency environment."),
        ("Vocabulary", "Terms retained by a fitted vectorizer."),
        ("Weighted-F1", "Support-weighted average of class F1 scores."),
        ("Widget", "Interactive interface control."),
        ("XAI", "Explainable AI."),
        ("XSum", "Highly abstractive news-summary dataset used for extractive evaluation."),
    ]
    groups = [("A-F", "A", "F"), ("G-M", "G", "M"), ("N-S", "N", "S"), ("T-Z", "T", "Z")]
    for title, start, end in groups:
        add_heading(doc, title, 2)
        rows = [(term, definition) for term, definition in glossary if start <= term[0].upper() <= end]
        add_table(doc, ["Term", "Meaning in this project"], rows, [1.8, 4.7], font_size=8.05)


def appendix_sources(doc) -> None:
    add_chapter(doc, 22, "Appendix D - Evidence basis and further learning")
    add_heading(doc, "22.1 Project evidence used for this guide", 2)
    add_bullets(
        doc,
        [
            "All Python modules under app.py, pages/, src/, training/ and tests/.",
            "models/model_metadata.json and the packaged Joblib model identity.",
            "reports/results/dataset_profile.json, model_metrics.json, model_comparison.csv, classification_report.csv, error_analysis.csv and summarization_metrics.json.",
            "The project diagrams, EDA/evaluation figures and warm editorial interface screenshots.",
            "README.md, MODEL_CARD.md, DATASET_CARD.md, the project report, setup guide and developer guide.",
        ],
    )
    add_heading(doc, "22.2 External conceptual sources named by the project", 2)
    add_table(
        doc,
        ["Resource", "Project relevance"],
        [
            ("ISOT Fake News Dataset - University of Victoria", "Classifier training/evaluation source; URL recorded in model metadata"),
            ("XSum - EdinburghNLP", "Fixed reference dataset for extractive ROUGE evaluation"),
            ("sshleifer/distilbart-cnn-6-6 - Hugging Face", "Optional pretrained abstractive summarizer"),
            ("scikit-learn documentation", "TF-IDF, Pipeline, GridSearchCV, Logistic Regression, SVM, Naive Bayes and metrics"),
            ("Streamlit documentation", "Multipage UI, session state and caching"),
            ("SQLite documentation", "Relational schema, constraints, indexes and queries"),
        ],
        [2.35, 4.15],
        font_size=8.5,
    )
    add_heading(doc, "22.3 Suggested learning sequence", 2)
    add_numbered(
        doc,
        [
            "Run one packaged sample and identify each visible output.",
            "Re-read Chapters 1-3 while following Figure 3.1.",
            "Learn tokens, vectors and TF-IDF before Logistic Regression formulas.",
            "Manually calculate precision and recall from the 2x2 confusion matrix.",
            "Open pages/01_Analyse_Article.py and trace each imported src function.",
            "Read training/prepare_dataset.py, then training/train_fake_news_models.py.",
            "Use the viva chapter aloud without memorising unsupported claims.",
            "Finish by explaining the difference between risk classification and evidence-based fact-checking.",
        ],
        space_after=2,
        line_spacing=1.15,
    )
    add_note(doc, "Final takeaway", "Understanding the limitations is part of understanding the technology. A technically correct explanation always includes what the system measures, the data on which it was measured, and what the result cannot prove.", fill="E9FAF7", accent=TEAL)


def build_guide() -> Path:
    profile = json.loads((RESULTS / "dataset_profile.json").read_text(encoding="utf-8"))
    metrics = json.loads((RESULTS / "model_metrics.json").read_text(encoding="utf-8"))
    summary_metrics = json.loads((RESULTS / "summarization_metrics.json").read_text(encoding="utf-8"))
    metadata = json.loads((ROOT / "models" / "model_metadata.json").read_text(encoding="utf-8"))
    import csv

    benchmark = json.loads((ROOT / "reports" / "model_benchmark_summary.json").read_text(encoding="utf-8"))

    DOCS.mkdir(parents=True, exist_ok=True)
    doc = configure_document(
        "NewsLens AI · Complete Concepts, Methodologies and Terminology Guide",
        "compact_reference_guide",
    )
    doc.core_properties.title = "NewsLens AI - Complete Concepts, Methodologies and Terminology Guide"
    doc.core_properties.subject = "Beginner-friendly technical companion to the NewsLens AI project"
    doc.core_properties.author = PROJECT_AUTHOR
    doc.core_properties.last_modified_by = PROJECT_AUTHOR
    doc.core_properties.comments = COPYRIGHT_NOTICE
    doc.core_properties.keywords = "AI, ML, NLP, TF-IDF, Logistic Regression, summarization, terminology, beginner guide"
    add_cover(doc)
    add_guide_map(doc)
    chapter_big_picture(doc)
    chapter_foundations(doc)
    chapter_system_flow(doc)
    chapter_data(doc, profile)
    chapter_preprocessing(doc)
    chapter_summarization(doc, summary_metrics)
    chapter_classifier(doc, metadata)
    chapter_training(doc, benchmark)
    chapter_metrics(doc, metrics)
    chapter_explainability(doc)
    chapter_architecture(doc)
    chapter_ingestion_security(doc)
    chapter_persistence(doc)
    chapter_ui(doc)
    chapter_testing(doc)
    chapter_ethics(doc)
    chapter_results(doc, metrics, summary_metrics)
    chapter_viva(doc)
    appendix_formulas(doc)
    appendix_file_map(doc)
    appendix_glossary(doc)
    appendix_sources(doc)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_guide()
    print(output)
    print(f"{output.stat().st_size:,} bytes")
